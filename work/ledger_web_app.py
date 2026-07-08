import cgi
import datetime as dt
import hashlib
import hmac
import html
import io
import json
import os
import re
import shutil
import sqlite3
import sys
import secrets
import ast
import operator
import threading
import urllib.parse
import traceback
from http.cookies import SimpleCookie
from http import HTTPStatus
from http.server import ThreadingHTTPServer, BaseHTTPRequestHandler
from pathlib import Path

from openpyxl import load_workbook, Workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(__file__).resolve().parents[1]
DATA_DIR = BASE_DIR / "outputs" / "ledger_system"
DB_PATH = DATA_DIR / "ledger_system.db"
STORAGE_DIR = DATA_DIR / "storage"
DELETED_STORAGE_DIR = DATA_DIR / "deleted_storage"
LOG_PATH = DATA_DIR / "server.log"
DEFAULT_USERNAME = os.environ.get("LEDGER_ADMIN_USER", "ZhanLin2026")
DEFAULT_PASSWORD = os.environ.get("LEDGER_ADMIN_PASSWORD", "Ahcx@ZL2026")
SESSION_COOKIE = "ledger_session"
ROLE_ADMIN = "admin"
ROLE_USER = "user"
PROJECT_OPTIONS = [
    "湖南安化抽水蓄能电站筹建期洞室及道路工程",
    "湖南安化抽水蓄能电站上水库工程",
    "湖南安化抽水蓄能电站引水系统工程",
    "湖南安化抽水蓄能电站地下厂房及尾水系统工程",
    "湖南安化抽水蓄能电站下水库工程",
]
SECTION_OPTIONS = ["Q1标", "Q2标", "C1标", "C2标", "C3标", "C4标"]
SOURCE_TYPE_OPTIONS = [
    "中国水利水电第七工程局有限公司湖南安化抽水蓄能电站筹建期洞室及道路工程项目经理部",
    "中国安能第三工程局湖南安化抽水蓄能电站上水库工程项目部",
    "中铁十一局集团有限公司湖南安化抽水蓄能电站引水系统工程项目经理部",
    "中国水利水电第七工程局有限公司湖南安化抽水蓄能电站地下厂房及尾水系统工程项目经理部",
    "中国水电建设集团十五工程局有限公司湖南安化抽水蓄能电站下水库工程项目部",
    "华东咨询-天津冀水联合体安化抽水蓄能电站工程建设监理中心",
    "湖南安化抽水蓄能有限公司",
]
DETECTION_UNIT_NAME = "湖南安化抽水蓄能电站施工期物探检测项目部"
USER_ORGANIZATION_OPTIONS = SOURCE_TYPE_OPTIONS + [DETECTION_UNIT_NAME]
PROJECT_SECTION_MAP = {
    "湖南安化抽水蓄能电站筹建期洞室及道路工程": "Q2标",
    "湖南安化抽水蓄能电站上水库工程": "C1标",
    "湖南安化抽水蓄能电站引水系统工程": "C2标",
    "湖南安化抽水蓄能电站地下厂房及尾水系统工程": "C3标",
    "湖南安化抽水蓄能电站下水库工程": "C4标",
}
PROJECT_CLIENT_MAP = {
    "湖南安化抽水蓄能电站筹建期洞室及道路工程": "中国水利水电第七工程局有限公司湖南安化抽水蓄能电站筹建期洞室及道路工程项目经理部",
    "湖南安化抽水蓄能电站上水库工程": "中国安能第三工程局湖南安化抽水蓄能电站上水库工程项目部",
    "湖南安化抽水蓄能电站引水系统工程": "中铁十一局集团有限公司湖南安化抽水蓄能电站引水系统工程项目经理部",
    "湖南安化抽水蓄能电站地下厂房及尾水系统工程": "中国水利水电第七工程局有限公司湖南安化抽水蓄能电站地下厂房及尾水系统工程项目经理部",
    "湖南安化抽水蓄能电站下水库工程": "中国水电建设集团十五工程局有限公司湖南安化抽水蓄能电站下水库工程项目部",
}


def log_error(message):
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    with open(LOG_PATH, "a", encoding="utf-8") as f:
        f.write(f"[{now_text()}] {message}\n")


def now_text():
    return dt.datetime.now().isoformat(timespec="seconds")


def safe_name(value):
    text = str(value or "未填写").strip()
    text = re.sub(r'[<>:"/\\|?*]+', "_", text)
    text = re.sub(r"\s+", "_", text)
    return text[:80] or "未填写"


def file_sha256(path):
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def connect():
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    DELETED_STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.execute("PRAGMA foreign_keys = ON")
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = connect()
    try:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS ledger_file (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_name TEXT,
                section_name TEXT,
                source_type TEXT,
                discipline TEXT,
                original_filename TEXT NOT NULL,
                file_key TEXT NOT NULL,
                current_version_id INTEGER,
                created_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS ledger_file_version (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                ledger_file_id INTEGER NOT NULL REFERENCES ledger_file(id) ON DELETE CASCADE,
                version_no INTEGER NOT NULL,
                stored_path TEXT NOT NULL,
                file_size INTEGER NOT NULL,
                file_hash TEXT NOT NULL,
                uploaded_at TEXT NOT NULL,
                remark TEXT,
                workbook_id INTEGER
            );

            CREATE TABLE IF NOT EXISTS template_workbook (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                source_file TEXT NOT NULL,
                file_name TEXT NOT NULL,
                imported_at TEXT NOT NULL,
                project_name TEXT,
                section_name TEXT,
                source_type TEXT,
                discipline TEXT,
                remark TEXT
            );

            CREATE TABLE IF NOT EXISTS template_sheet (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                workbook_id INTEGER NOT NULL REFERENCES template_workbook(id) ON DELETE CASCADE,
                sheet_index INTEGER NOT NULL,
                sheet_name TEXT NOT NULL,
                max_row INTEGER,
                max_column INTEGER,
                freeze_panes TEXT,
                sheet_state TEXT
            );

            CREATE TABLE IF NOT EXISTS template_cell (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                sheet_id INTEGER NOT NULL REFERENCES template_sheet(id) ON DELETE CASCADE,
                row_index INTEGER NOT NULL,
                col_index INTEGER NOT NULL,
                cell_ref TEXT NOT NULL,
                raw_value TEXT,
                formula TEXT,
                data_type TEXT,
                number_format TEXT,
                style_json TEXT
            );

            CREATE TABLE IF NOT EXISTS template_merge (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                sheet_id INTEGER NOT NULL REFERENCES template_sheet(id) ON DELETE CASCADE,
                range_ref TEXT NOT NULL,
                min_row INTEGER,
                min_col INTEGER,
                max_row INTEGER,
                max_col INTEGER
            );

            CREATE TABLE IF NOT EXISTS template_row_dimension (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                sheet_id INTEGER NOT NULL REFERENCES template_sheet(id) ON DELETE CASCADE,
                row_index INTEGER NOT NULL,
                height REAL,
                hidden INTEGER,
                outline_level INTEGER
            );

            CREATE TABLE IF NOT EXISTS template_column_dimension (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                sheet_id INTEGER NOT NULL REFERENCES template_sheet(id) ON DELETE CASCADE,
                col_index INTEGER NOT NULL,
                col_letter TEXT NOT NULL,
                width REAL,
                hidden INTEGER,
                outline_level INTEGER
            );

            CREATE INDEX IF NOT EXISTS idx_ledger_file_key ON ledger_file(file_key);
            CREATE INDEX IF NOT EXISTS idx_template_cell_sheet ON template_cell(sheet_id, row_index, col_index);

            CREATE TABLE IF NOT EXISTS app_user (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                salt TEXT NOT NULL,
                display_name TEXT,
                organization TEXT,
                role TEXT NOT NULL DEFAULT 'user',
                is_active INTEGER NOT NULL DEFAULT 1,
                created_by INTEGER REFERENCES app_user(id) ON DELETE SET NULL,
                created_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS app_session (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL REFERENCES app_user(id) ON DELETE CASCADE,
                token TEXT NOT NULL UNIQUE,
                created_at TEXT NOT NULL,
                expires_at TEXT NOT NULL
            );
            """
        )
        ensure_user_columns(conn)
        ensure_default_user(conn)
        normalize_detection_unit(conn)
        conn.commit()
    finally:
        conn.close()


def hash_password(password, salt):
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 120000)
    return digest.hex()


def ensure_user_columns(conn):
    columns = {row["name"] for row in conn.execute("PRAGMA table_info(app_user)")}
    if "role" not in columns:
        conn.execute("ALTER TABLE app_user ADD COLUMN role TEXT NOT NULL DEFAULT 'user'")
    if "created_by" not in columns:
        conn.execute("ALTER TABLE app_user ADD COLUMN created_by INTEGER REFERENCES app_user(id) ON DELETE SET NULL")
    if "organization" not in columns:
        conn.execute("ALTER TABLE app_user ADD COLUMN organization TEXT")


def ensure_default_user(conn):
    row = conn.execute("SELECT id FROM app_user WHERE username = ?", (DEFAULT_USERNAME,)).fetchone()
    salt = secrets.token_hex(16)
    if row:
        conn.execute(
            """
            UPDATE app_user
            SET password_hash = ?, salt = ?, display_name = ?, role = ?, is_active = 1
            WHERE id = ?
            """,
            (hash_password(DEFAULT_PASSWORD, salt), salt, "管理员", ROLE_ADMIN, row["id"]),
        )
        conn.execute(
            "UPDATE app_user SET is_active = 0 WHERE username = ? AND username <> ?",
            ("admin", DEFAULT_USERNAME),
        )
        return
    conn.execute(
        """
        INSERT INTO app_user (username, password_hash, salt, display_name, role, created_at)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (DEFAULT_USERNAME, hash_password(DEFAULT_PASSWORD, salt), salt, "管理员", ROLE_ADMIN, now_text()),
    )
    conn.execute(
        "UPDATE app_user SET is_active = 0 WHERE username = ? AND username <> ?",
        ("admin", DEFAULT_USERNAME),
    )


def normalize_detection_unit(conn):
    conn.execute(
        "UPDATE ledger_file SET discipline = ? WHERE discipline IS NULL OR discipline = '' OR discipline = '物探'",
        (DETECTION_UNIT_NAME,),
    )
    conn.execute(
        "UPDATE template_workbook SET discipline = ? WHERE discipline IS NULL OR discipline = '' OR discipline = '物探'",
        (DETECTION_UNIT_NAME,),
    )


def authenticate_user(username, password):
    conn = connect()
    try:
        row = conn.execute(
            "SELECT * FROM app_user WHERE username = ? AND is_active = 1",
            (username,),
        ).fetchone()
        if not row:
            return None
        expected = hash_password(password, row["salt"])
        if hmac.compare_digest(expected, row["password_hash"]):
            return row
        return None
    finally:
        conn.close()


def is_admin(user):
    return bool(user and user["role"] == ROLE_ADMIN)


def generate_password(length=10):
    alphabet = "ABCDEFGHJKLMNPQRSTUVWXYZabcdefghijkmnopqrstuvwxyz23456789"
    return "".join(secrets.choice(alphabet) for _ in range(length))


def create_app_user(username, display_name, organization, role, created_by):
    username = (username or "").strip()
    display_name = (display_name or "").strip() or username
    organization = (organization or "").strip()
    role = (role or ROLE_USER).strip()
    if not username:
        raise ValueError("账号不能为空")
    if not re.fullmatch(r"[A-Za-z0-9_@.\-]{3,32}", username):
        raise ValueError("账号需为3-32位，可使用字母、数字、下划线、横线、点或@")
    if organization not in USER_ORGANIZATION_OPTIONS:
        raise ValueError("请选择有效的用户单位")
    if role not in (ROLE_USER, ROLE_ADMIN):
        raise ValueError("请选择有效的账号类型")

    password = generate_password()
    salt = secrets.token_hex(16)
    conn = connect()
    try:
        with conn:
            exists = conn.execute("SELECT id FROM app_user WHERE username = ?", (username,)).fetchone()
            if exists:
                raise ValueError("该账号已存在")
            conn.execute(
                """
                INSERT INTO app_user
                    (username, password_hash, salt, display_name, organization, role, is_active, created_by, created_at)
                VALUES (?, ?, ?, ?, ?, ?, 1, ?, ?)
                """,
                (
                    username,
                    hash_password(password, salt),
                    salt,
                    display_name,
                    organization,
                    role,
                    created_by["id"] if created_by else None,
                    now_text(),
                ),
            )
    finally:
        conn.close()
    return password


def set_user_active(target_user_id, active, operator):
    if not is_admin(operator):
        raise ValueError("仅管理员可操作账号状态")
    active = 1 if active else 0
    if int(target_user_id) == int(operator["id"]) and not active:
        raise ValueError("不能停用当前登录的管理员账号")

    conn = connect()
    try:
        with conn:
            target = conn.execute("SELECT * FROM app_user WHERE id = ?", (target_user_id,)).fetchone()
            if not target:
                raise ValueError("未找到该账号")
            if int(target["is_active"]) == active:
                return
            if not active and target["role"] == ROLE_ADMIN:
                active_admin_count = conn.execute(
                    "SELECT COUNT(*) FROM app_user WHERE role = ? AND is_active = 1",
                    (ROLE_ADMIN,),
                ).fetchone()[0]
                if active_admin_count <= 1:
                    raise ValueError("至少需要保留一个启用状态的管理员账号")
            conn.execute("UPDATE app_user SET is_active = ? WHERE id = ?", (active, target_user_id))
            if not active:
                conn.execute("DELETE FROM app_session WHERE user_id = ?", (target_user_id,))
    finally:
        conn.close()


def update_user_organization(target_user_id, organization, operator):
    if not is_admin(operator):
        raise ValueError("仅管理员可修改用户单位")
    organization = (organization or "").strip()
    if organization not in USER_ORGANIZATION_OPTIONS:
        raise ValueError("请选择有效的用户单位")

    conn = connect()
    try:
        with conn:
            target = conn.execute("SELECT id FROM app_user WHERE id = ?", (target_user_id,)).fetchone()
            if not target:
                raise ValueError("未找到该账号")
            conn.execute("UPDATE app_user SET organization = ? WHERE id = ?", (organization, target_user_id))
    finally:
        conn.close()


def update_user_role(target_user_id, role, operator):
    if not is_admin(operator):
        raise ValueError("仅管理员可修改账号角色")
    role = (role or "").strip()
    if role not in (ROLE_USER, ROLE_ADMIN):
        raise ValueError("请选择有效的账号角色")
    if int(target_user_id) == int(operator["id"]) and role != ROLE_ADMIN:
        raise ValueError("不能将当前登录管理员改为普通用户")

    conn = connect()
    try:
        with conn:
            target = conn.execute("SELECT * FROM app_user WHERE id = ?", (target_user_id,)).fetchone()
            if not target:
                raise ValueError("未找到该账号")
            if target["role"] == ROLE_ADMIN and role != ROLE_ADMIN and target["is_active"]:
                active_admin_count = conn.execute(
                    "SELECT COUNT(*) FROM app_user WHERE role = ? AND is_active = 1",
                    (ROLE_ADMIN,),
                ).fetchone()[0]
                if active_admin_count <= 1:
                    raise ValueError("至少需要保留一个启用状态的管理员账号")
            conn.execute("UPDATE app_user SET role = ? WHERE id = ?", (role, target_user_id))
            if role != target["role"]:
                conn.execute("DELETE FROM app_session WHERE user_id = ?", (target_user_id,))
    finally:
        conn.close()


def reset_user_password(target_user_id, operator):
    if not is_admin(operator):
        raise ValueError("仅管理员可重置密码")
    password = generate_password()
    salt = secrets.token_hex(16)
    conn = connect()
    try:
        with conn:
            target = conn.execute("SELECT id, username FROM app_user WHERE id = ?", (target_user_id,)).fetchone()
            if not target:
                raise ValueError("未找到该账号")
            conn.execute(
                "UPDATE app_user SET password_hash = ?, salt = ? WHERE id = ?",
                (hash_password(password, salt), salt, target_user_id),
            )
            conn.execute("DELETE FROM app_session WHERE user_id = ?", (target_user_id,))
            return target["username"], password
    finally:
        conn.close()


def save_sheet_changes(sheet_id, changes):
    if not isinstance(changes, list):
        raise ValueError("保存数据格式不正确")
    conn = connect()
    try:
        with conn:
            sheet = conn.execute("SELECT id FROM template_sheet WHERE id = ?", (sheet_id,)).fetchone()
            if not sheet:
                raise ValueError("未找到该工作表")
            existing_values = {
                (row["row_index"], row["col_index"]): row["raw_value"]
                for row in conn.execute(
                    "SELECT row_index, col_index, raw_value FROM template_cell WHERE sheet_id = ?",
                    (sheet_id,),
                )
            }
            header_row = detect_header_row(existing_values)
            data_start_row = header_row + 1
            ratio_cols = {
                col
                for (row, col), cell_value in existing_values.items()
                if row == header_row and is_ratio_header(cell_value)
            }
            saved = 0
            for item in changes:
                if not isinstance(item, dict):
                    continue
                row_index = int(item.get("row") or 0)
                col_index = int(item.get("col") or 0)
                if row_index <= 0 or col_index <= 0:
                    continue
                value = normalize_cell_text(str(item.get("value") or ""))
                formula = value if value.startswith("=") else None
                display_value = evaluate_db_formula(formula, existing_values) if formula else value
                if formula and display_value is None:
                    display_value = ""
                display_value = normalize_cell_text(display_value)
                if row_index >= data_start_row and col_index in ratio_cols:
                    display_value = display_ratio_text(display_value)
                existing = conn.execute(
                    """
                    SELECT id
                    FROM template_cell
                    WHERE sheet_id = ? AND row_index = ? AND col_index = ?
                    ORDER BY id
                    LIMIT 1
                    """,
                    (sheet_id, row_index, col_index),
                ).fetchone()
                if existing:
                    conn.execute(
                        """
                        UPDATE template_cell
                        SET raw_value = ?, formula = ?, data_type = 's'
                        WHERE id = ?
                        """,
                        (str(display_value), formula, existing["id"]),
                    )
                else:
                    conn.execute(
                        """
                        INSERT INTO template_cell
                            (sheet_id, row_index, col_index, cell_ref, raw_value, formula, data_type, number_format, style_json)
                        VALUES (?, ?, ?, ?, ?, ?, 's', 'General', NULL)
                        """,
                        (sheet_id, row_index, col_index, f"{col_letter(col_index)}{row_index}", str(display_value), formula),
                    )
                existing_values[(row_index, col_index)] = display_value
                saved += 1
    finally:
        conn.close()
    return saved


def delete_sheet_axis(sheet_id, axis, index):
    axis = str(axis or "")
    index = int(index or 0)
    if axis not in {"row", "col"} or index <= 0:
        raise ValueError("删除参数不正确")
    conn = connect()
    try:
        with conn:
            sheet = conn.execute("SELECT * FROM template_sheet WHERE id = ?", (sheet_id,)).fetchone()
            if not sheet:
                raise ValueError("未找到该工作表")
            if axis == "row":
                max_row = int(sheet["max_row"] or 0)
                if index > max_row:
                    raise ValueError("行号超出范围")
                conn.execute("DELETE FROM template_cell WHERE sheet_id = ? AND row_index = ?", (sheet_id, index))
                conn.execute("UPDATE template_cell SET row_index = row_index - 1 WHERE sheet_id = ? AND row_index > ?", (sheet_id, index))
                for row in conn.execute("SELECT id, row_index, col_index FROM template_cell WHERE sheet_id = ?", (sheet_id,)):
                    conn.execute("UPDATE template_cell SET cell_ref = ? WHERE id = ?", (f"{col_letter(row['col_index'])}{row['row_index']}", row["id"]))
                conn.execute("DELETE FROM template_row_dimension WHERE sheet_id = ? AND row_index = ?", (sheet_id, index))
                conn.execute("UPDATE template_row_dimension SET row_index = row_index - 1 WHERE sheet_id = ? AND row_index > ?", (sheet_id, index))
                conn.execute("UPDATE template_sheet SET max_row = CASE WHEN max_row > 0 THEN max_row - 1 ELSE 0 END WHERE id = ?", (sheet_id,))
            else:
                max_col = int(sheet["max_column"] or 0)
                if index > max_col:
                    raise ValueError("列号超出范围")
                conn.execute("DELETE FROM template_cell WHERE sheet_id = ? AND col_index = ?", (sheet_id, index))
                conn.execute("UPDATE template_cell SET col_index = col_index - 1 WHERE sheet_id = ? AND col_index > ?", (sheet_id, index))
                for row in conn.execute("SELECT id, row_index, col_index FROM template_cell WHERE sheet_id = ?", (sheet_id,)):
                    conn.execute("UPDATE template_cell SET cell_ref = ? WHERE id = ?", (f"{col_letter(row['col_index'])}{row['row_index']}", row["id"]))
                conn.execute("DELETE FROM template_column_dimension WHERE sheet_id = ? AND col_index = ?", (sheet_id, index))
                conn.execute("UPDATE template_column_dimension SET col_index = col_index - 1 WHERE sheet_id = ? AND col_index > ?", (sheet_id, index))
                for row in conn.execute("SELECT id, col_index FROM template_column_dimension WHERE sheet_id = ? AND col_index >= ?", (sheet_id, index)):
                    conn.execute(
                        "UPDATE template_column_dimension SET col_letter = ? WHERE id = ?",
                        (col_letter(row["col_index"]), row["id"]),
                    )
                conn.execute("UPDATE template_sheet SET max_column = CASE WHEN max_column > 0 THEN max_column - 1 ELSE 0 END WHERE id = ?", (sheet_id,))
    finally:
        conn.close()


def create_session(user_id):
    token = secrets.token_urlsafe(32)
    created = dt.datetime.now()
    expires = created + dt.timedelta(hours=12)
    conn = connect()
    try:
        with conn:
            conn.execute(
                "INSERT INTO app_session (user_id, token, created_at, expires_at) VALUES (?, ?, ?, ?)",
                (user_id, token, created.isoformat(timespec="seconds"), expires.isoformat(timespec="seconds")),
            )
    finally:
        conn.close()
    return token, expires


def get_session_user(token):
    if not token:
        return None
    conn = connect()
    try:
        row = conn.execute(
            """
            SELECT u.*
            FROM app_session s
            JOIN app_user u ON u.id = s.user_id
            WHERE s.token = ? AND s.expires_at > ? AND u.is_active = 1
            """,
            (token, now_text()),
        ).fetchone()
        return row
    finally:
        conn.close()


def delete_session(token):
    if not token:
        return
    conn = connect()
    try:
        with conn:
            conn.execute("DELETE FROM app_session WHERE token = ?", (token,))
    finally:
        conn.close()


def to_text(value):
    if value is None:
        return None
    if isinstance(value, dt.datetime):
        return value.date().isoformat()
    if isinstance(value, dt.date):
        return value.isoformat()
    return str(value)


def display_cell_text(value):
    text = "" if value is None else str(value)
    match = re.fullmatch(r"(\d{4}-\d{1,2}-\d{1,2})[ T]\d{1,2}:\d{2}:\d{2}(?:\.\d+)?", text.strip())
    if match:
        return match.group(1)
    return text


def normalize_cell_text(value):
    return display_cell_text(value)


def display_ratio_text(value):
    text = display_cell_text(value).strip()
    if not text:
        return text
    has_percent = text.endswith("%")
    number_text = text[:-1].strip() if has_percent else text
    try:
        number = float(number_text.replace(",", ""))
    except ValueError:
        return text
    return f"{number:.2f}%" if has_percent else f"{number:.2f}"


def is_ratio_header(value):
    text = display_cell_text(value)
    return "比例" in text or "占比" in text


def detect_header_row(cells, row_limit=20):
    header_keywords = [
        "检测比例",
        "占比",
        "合格率",
        "工程部位",
        "施工数量",
        "检测数量",
        "一类",
        "二类",
        "三类",
        "四类",
        "桩号",
        "日期",
    ]
    best_row = 4
    best_score = 0
    for row in range(1, min(row_limit, 20) + 1):
        values = [display_cell_text(value) for (r, _), value in cells.items() if r == row]
        if not values:
            continue
        joined = " ".join(values)
        score = sum(1 for keyword in header_keywords if keyword in joined)
        non_empty = sum(1 for value in values if value.strip())
        if score > best_score or (score == best_score and score > 0 and non_empty > 3):
            best_row = row
            best_score = score
    return best_row


def parse_date_value(value):
    text = display_cell_text(value).strip()
    match = re.search(r"(\d{4})[-/\.](\d{1,2})[-/\.](\d{1,2})", text)
    if not match:
        return None
    try:
        return dt.date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
    except ValueError:
        return None


def parse_number_value(value):
    text = display_cell_text(value).strip().replace(",", "")
    if not text:
        return 0
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        return 0
    try:
        return float(match.group(0))
    except ValueError:
        return 0


def unique_ordered_values(values):
    result = []
    seen = set()
    for value in values:
        text = str(value).strip()
        if text and text not in seen:
            seen.add(text)
            result.append(text)
    return result


def normalize_design_length_text(value):
    text = display_cell_text(value).strip()
    if not text:
        return ""
    normalized = (
        text.replace("Φ", "φ")
        .replace("Ф", "φ")
        .replace("￠", "φ")
        .replace("，", "/")
        .replace("、", "/")
        .replace("；", "/")
        .replace(";", "/")
        .replace(" ", "")
    )
    without_diameter = re.sub(r"φ\s*\d+(?:\.\d+)?", "", normalized, flags=re.IGNORECASE)
    numbers = re.findall(r"(\d+(?:\.\d+)?)\s*[mM米]", without_diameter)
    if not numbers:
        numbers = re.findall(r"\d+(?:\.\d+)?", without_diameter)
    return "/".join(unique_ordered_values(numbers)) if numbers else text


def normalize_design_diameter_text(value):
    text = display_cell_text(value).strip()
    if not text:
        return ""
    normalized = (
        text.replace("Φ", "φ")
        .replace("Ф", "φ")
        .replace("￠", "φ")
        .replace("，", "/")
        .replace("、", "/")
        .replace("；", "/")
        .replace(";", "/")
        .replace(" ", "")
    )
    phi_numbers = re.findall(r"φ\s*(\d+(?:\.\d+)?)", normalized, flags=re.IGNORECASE)
    if phi_numbers:
        return "/".join(unique_ordered_values(phi_numbers))

    without_material_grade = re.sub(r"HRB\d+[A-Z]?", "", normalized, flags=re.IGNORECASE)
    candidates = []
    for number in re.findall(r"\d+(?:\.\d+)?", without_material_grade):
        try:
            numeric_value = float(number)
        except ValueError:
            continue
        if 10 <= numeric_value <= 80:
            candidates.append(str(int(numeric_value)) if numeric_value.is_integer() else number)
    return "/".join(unique_ordered_values(candidates)) if candidates else text


def find_col_by_keywords(header_map, keywords, min_col=1):
    best_col = None
    best_score = 0
    for col, header in header_map.items():
        if col < min_col:
            continue
        text = display_cell_text(header)
        score = sum(1 for keyword in keywords if keyword in text)
        if score > best_score:
            best_col = col
            best_score = score
    return best_col


def find_col_by_keywords_excluding(header_map, keywords, excluded_keywords=None, min_col=1):
    excluded_keywords = excluded_keywords or []
    best_col = None
    best_score = 0
    for col, header in header_map.items():
        if col < min_col:
            continue
        text = display_cell_text(header)
        if any(keyword in text for keyword in excluded_keywords):
            continue
        score = sum(1 for keyword in keywords if keyword in text)
        if score > best_score:
            best_col = col
            best_score = score
    return best_col


def client_short_type(source_type):
    text = source_type or ""
    if "监理" in text or "冀水" in text or "咨询" in text:
        return "监理"
    if "业主" in text or "有限公司" == text[-4:]:
        return "业主"
    return "施工"


def monthly_report_table_type(sheet_name, client_type):
    name = clean_sheet_display_name(sheet_name)
    suffix = f"[{client_type}]" if client_type in ("施工", "监理") else f"[{client_type}委托]"
    if "锚杆无损" in name:
        return f"锚杆无损检测完成情况{suffix}"
    if "锚杆拉拔" in name:
        return f"锚杆拉拔检测完成情况{suffix}"
    if "钻孔摄像" in name or "钻孔成像" in name:
        return f"钻孔摄像检测完成情况{suffix}"
    if "回填灌浆" in name:
        return f"回填灌浆质量单孔注浆试验检测{suffix}"
    if "锚索" in name and "张拉" in name:
        return f"锚索多循环张拉试验完成情况{suffix}"
    if "预应力锚杆" in name or ("锚杆" in name and "张拉" in name):
        return f"预应力锚杆张拉试验完成情况{suffix}"
    if "桩身完整性" in name:
        return f"桩身完整性检测完成情况{suffix}"
    if "松弛圈" in name:
        return f"松弛圈单孔声波检测完成情况{suffix}"
    if "弹性波" in name:
        return f"围岩弹性波检测完成情况{suffix}"
    return f"{name}完成情况{suffix}"


def default_period(report_type):
    today = dt.date.today()
    if report_type == "week":
        start = today - dt.timedelta(days=today.weekday())
        end = start + dt.timedelta(days=6)
    elif report_type == "year":
        start = dt.date(today.year, 1, 1)
        end = dt.date(today.year, 12, 31)
    elif report_type == "quarter":
        quarter_start_month = ((today.month - 1) // 3) * 3 + 1
        start = dt.date(today.year, quarter_start_month, 1)
        next_month = quarter_start_month + 3
        end = dt.date(today.year + (1 if next_month > 12 else 0), ((next_month - 1) % 12) + 1, 1) - dt.timedelta(days=1)
    else:
        start = dt.date(today.year, today.month, 1)
        end = dt.date(today.year + (1 if today.month == 12 else 0), 1 if today.month == 12 else today.month + 1, 1) - dt.timedelta(days=1)
    return start, end


def parse_statistics_params(params):
    def query_value(key, default=""):
        return params.get(key, [default])[0]

    def query_values(key):
        return [str(value).strip() for value in params.get(key, []) if str(value).strip()]

    report_type = query_value("report_type", "month")
    default_start, default_end = default_period(report_type)
    start_text = query_value("start_date", default_start.isoformat())
    end_text = query_value("end_date", default_end.isoformat())
    start_date = parse_date_value(start_text) or default_start
    end_date = parse_date_value(end_text) or default_end
    if start_date > end_date:
        start_date, end_date = end_date, start_date
    return {
        "report_type": report_type,
        "start_date": start_date,
        "end_date": end_date,
        "source_filters": query_values("source_type"),
        "sheet_filters": query_values("sheet_name"),
        "unit_filters": query_values("unit_name"),
    }


def collect_stat_records(start_date, end_date, source_filters=None, sheet_filters=None):
    source_filters = [value for value in (source_filters or []) if value]
    sheet_filters = [value for value in (sheet_filters or []) if value]
    conn = connect()
    records = []
    try:
        where_parts = ["1 = 1"]
        args = []
        if source_filters:
            where_parts.append(f"f.source_type IN ({','.join('?' for _ in source_filters)})")
            args.extend(source_filters)
        if sheet_filters:
            where_parts.append(f"s.sheet_name IN ({','.join('?' for _ in sheet_filters)})")
            args.extend(sheet_filters)
        sheets = conn.execute(
            f"""
            SELECT s.id AS sheet_id, s.sheet_name, s.max_row, s.max_column,
                   f.project_name, f.section_name, f.source_type, f.original_filename
            FROM template_sheet s
            JOIN template_workbook w ON w.id = s.workbook_id
            JOIN ledger_file_version v ON v.workbook_id = w.id
            JOIN ledger_file f ON f.current_version_id = v.id
            WHERE {' AND '.join(where_parts)}
            ORDER BY f.id ASC, s.sheet_index ASC, s.id ASC
            """,
            args,
        ).fetchall()
        for sheet in sheets:
            rows = conn.execute(
                """
                SELECT row_index, col_index, raw_value
                FROM template_cell
                WHERE sheet_id = ?
                """,
                (sheet["sheet_id"],),
            ).fetchall()
            cells = {(row["row_index"], row["col_index"]): row["raw_value"] for row in rows}
            if not cells:
                continue
            header_row = detect_header_row(cells, sheet["max_row"] or 20)
            data_start = header_row + 1
            header_map = {
                col: cells.get((header_row, col), "")
                for col in range(1, (sheet["max_column"] or 0) + 1)
            }
            sheet_display_name = clean_sheet_display_name(sheet["sheet_name"])
            client_type = client_short_type(sheet["source_type"] or "")
            report_table_type = monthly_report_table_type(sheet["sheet_name"], client_type)
            summary_start_col = 20
            date_col = find_col_by_keywords(header_map, ["日期", "检测日期", "试验日期"])
            unit_col = find_col_by_keywords(header_map, ["单位工程", "单元工程"])
            sub_unit_col = find_col_by_keywords(header_map, ["分部工程"])
            location_col = find_col_by_keywords(header_map, ["工程部位", "部位", "施工部位"], summary_start_col)
            construction_col = find_col_by_keywords(header_map, ["施工数量", "施工量"], summary_start_col)
            testing_col = find_col_by_keywords(header_map, ["检测数量", "检测量"], summary_start_col)
            ratio_col = find_col_by_keywords(header_map, ["抽检比例", "检测比例"], summary_start_col)
            design_force_col = find_col_by_keywords(header_map, ["设计拉拔力", "合格抗拔力"], summary_start_col)
            min_value_col = find_col_by_keywords(header_map, ["最小值"], summary_start_col)
            max_value_col = find_col_by_keywords(header_map, ["最大值"], summary_start_col)
            pass_rate_col = find_col_by_keywords(header_map, ["合格率"], summary_start_col)
            class_one_col = find_col_by_keywords(header_map, ["一类杆", "一类", "Ⅰ", "Ⅰ类"], summary_start_col)
            class_two_col = find_col_by_keywords(header_map, ["二类杆", "二类", "Ⅱ", "Ⅱ类"], summary_start_col)
            class_three_col = find_col_by_keywords(header_map, ["三类杆", "三类", "Ⅲ", "Ⅲ类"], summary_start_col)
            class_four_col = find_col_by_keywords(header_map, ["四类杆", "四类", "Ⅳ", "Ⅳ类"], summary_start_col)
            grout_hole_col = find_col_by_keywords(header_map, ["检查孔编号"], summary_start_col)
            grout_station_col = find_col_by_keywords(header_map, ["检查桩号"], summary_start_col)
            grout_elevation_col = find_col_by_keywords(header_map, ["孔口高程"], summary_start_col)
            grout_pressure_col = find_col_by_keywords(header_map, ["设计压力"], summary_start_col)
            grout_measured_pressure_col = find_col_by_keywords(header_map, ["实测压力"], summary_start_col)
            grout_requirement_col = find_col_by_keywords(header_map, ["设计要求"], summary_start_col)
            grout_volume_col = find_col_by_keywords_excluding(header_map, ["初始10min内注浆量", "注浆量"], ["设计要求", "规定压力"], summary_start_col)
            grout_result_col = find_col_by_keywords(header_map, ["检测结果"], summary_start_col)
            report_no_col = find_col_by_keywords(header_map, ["报告编号", "报告号"])
            entrust_no_col = find_col_by_keywords(header_map, ["委托编号", "委托号"])
            result_col = find_col_by_keywords(header_map, ["检测结果", "试验结果", "综合评判类别", "评判类别", "结论"])
            hole_no_col = find_col_by_keywords(header_map, ["钻孔编号", "锚索孔号", "检查孔编号", "孔号", "锚索编号"])
            depth_start_1_col = find_col_by_keywords(header_map, ["起始深度1", "起始深度"])
            depth_end_1_col = find_col_by_keywords(header_map, ["终点深度1", "终点深度"])
            depth_start_2_col = find_col_by_keywords(header_map, ["起始深度2"])
            depth_end_2_col = find_col_by_keywords(header_map, ["终点深度2"])
            mileage_col = find_col_by_keywords(header_map, ["检测里程", "里程"])
            anchor_no_col = find_col_by_keywords(header_map, ["锚索编号", "锚杆编号"])
            spec_col = find_col_by_keywords(header_map, ["规格/型号", "锚索规格", "锚杆规格", "规格"])
            length_col = find_col_by_keywords(header_map, ["锚杆长度", "锚索长度", "杆长", "桩    长", "桩长", "长度"])
            diameter_col = find_col_by_keywords(header_map, ["锚杆直径", "锚索直径", "直径"])
            diameter_source_col = diameter_col or spec_col
            anchor_length_col = find_col_by_keywords(header_map, ["锚索长度", "锚杆长度", "杆长", "桩    长", "桩长"])
            anchorage_length_col = find_col_by_keywords(header_map, ["锚固长度"])
            free_length_col = find_col_by_keywords(header_map, ["自由段长度"])
            elastic_modulus_col = find_col_by_keywords(header_map, ["弹性模量"])
            load_requirement_col = find_col_by_keywords(header_map, ["分级荷载设计要求"])
            design_anchor_force_col = find_col_by_keywords(header_map, ["设计锚固力", "设计拉拔力", "合格抗拔力"])
            pile_type_col = find_col_by_keywords(header_map, ["桩    型", "桩型"])
            pile_diameter_col = find_col_by_keywords(header_map, ["桩    径", "桩径"])
            pile_top_col = find_col_by_keywords(header_map, ["桩顶高程"])
            category_col = find_col_by_keywords(header_map, ["综合评判类别", "评判类别"])
            hole_count_col = find_col_by_keywords(header_map, ["检测孔数量", "检查孔数量"])
            group_count_col = find_col_by_keywords(header_map, ["组数"])
            detection_type_col = find_col_by_keywords(header_map, ["检测类型", "试验类型"])
            remark_col = find_col_by_keywords(header_map, ["备注"])
            if not construction_col:
                construction_col = find_col_by_keywords(header_map, ["施工数量", "施工量"])
            if not testing_col:
                testing_col = find_col_by_keywords(header_map, ["检测数量", "检测量", "检测里程", "检测孔数量", "组数"])
            if not location_col:
                location_col = find_col_by_keywords(header_map, ["工程部位", "部位", "施工部位"])
            if not date_col:
                continue
            for row_index in range(data_start, (sheet["max_row"] or 0) + 1):
                detect_date = parse_date_value(cells.get((row_index, date_col), ""))
                if not detect_date or detect_date < start_date or detect_date > end_date:
                    continue
                unit_name = display_cell_text(cells.get((row_index, unit_col), "")) if unit_col else "未识别单元工程"
                sub_unit_name = display_cell_text(cells.get((row_index, sub_unit_col), "")) if sub_unit_col else ""
                location_name = display_cell_text(cells.get((row_index, location_col), "")) if location_col else ""
                part_name = location_name or unit_name or "未填写工程部位"
                records.append(
                    {
                        "source_type": sheet["source_type"] or "未填写委托单位",
                        "client_type": client_type,
                        "section_name": sheet["section_name"] or "",
                        "original_filename": sheet["original_filename"] or "",
                        "sheet_name": sheet_display_name,
                        "report_table_type": report_table_type,
                        "unit_name": unit_name or "",
                        "sub_unit_name": sub_unit_name,
                        "location_name": location_name,
                        "part_name": part_name,
                        "report_no": display_cell_text(cells.get((row_index, report_no_col), "")) if report_no_col else "",
                        "entrust_no": display_cell_text(cells.get((row_index, entrust_no_col), "")) if entrust_no_col else "",
                        "result_text": display_cell_text(cells.get((row_index, result_col), "")) if result_col else "",
                        "construction_qty": parse_number_value(cells.get((row_index, construction_col), "")) if construction_col else 0,
                        "testing_qty": parse_number_value(cells.get((row_index, testing_col), "")) if testing_col else 0,
                        "ratio_value": parse_number_value(cells.get((row_index, ratio_col), "")) if ratio_col else 0,
                        "length_value": normalize_design_length_text(cells.get((row_index, length_col), "")) if length_col else "",
                        "diameter_value": normalize_design_diameter_text(cells.get((row_index, diameter_source_col), "")) if diameter_source_col else "",
                        "design_force_value": display_cell_text(cells.get((row_index, design_force_col), "")) if design_force_col else "",
                        "hole_no": display_cell_text(cells.get((row_index, hole_no_col), "")) if hole_no_col else "",
                        "depth_start_1": display_cell_text(cells.get((row_index, depth_start_1_col), "")) if depth_start_1_col else "",
                        "depth_end_1": display_cell_text(cells.get((row_index, depth_end_1_col), "")) if depth_end_1_col else "",
                        "depth_start_2": display_cell_text(cells.get((row_index, depth_start_2_col), "")) if depth_start_2_col else "",
                        "depth_end_2": display_cell_text(cells.get((row_index, depth_end_2_col), "")) if depth_end_2_col else "",
                        "mileage_value": parse_number_value(cells.get((row_index, mileage_col), "")) if mileage_col else 0,
                        "mileage_text": display_cell_text(cells.get((row_index, mileage_col), "")) if mileage_col else "",
                        "anchor_no": display_cell_text(cells.get((row_index, anchor_no_col), "")) if anchor_no_col else "",
                        "spec_text": display_cell_text(cells.get((row_index, spec_col), "")) if spec_col else "",
                        "anchor_length_text": display_cell_text(cells.get((row_index, anchor_length_col), "")) if anchor_length_col else "",
                        "anchorage_length_text": display_cell_text(cells.get((row_index, anchorage_length_col), "")) if anchorage_length_col else "",
                        "free_length_text": display_cell_text(cells.get((row_index, free_length_col), "")) if free_length_col else "",
                        "elastic_modulus_text": display_cell_text(cells.get((row_index, elastic_modulus_col), "")) if elastic_modulus_col else "",
                        "load_requirement_text": display_cell_text(cells.get((row_index, load_requirement_col), "")) if load_requirement_col else "",
                        "design_anchor_force_text": display_cell_text(cells.get((row_index, design_anchor_force_col), "")) if design_anchor_force_col else "",
                        "pile_type_text": display_cell_text(cells.get((row_index, pile_type_col), "")) if pile_type_col else "",
                        "pile_diameter_text": display_cell_text(cells.get((row_index, pile_diameter_col), "")) if pile_diameter_col else "",
                        "pile_top_text": display_cell_text(cells.get((row_index, pile_top_col), "")) if pile_top_col else "",
                        "category_text": display_cell_text(cells.get((row_index, category_col), "")) if category_col else "",
                        "hole_count": parse_number_value(cells.get((row_index, hole_count_col), "")) if hole_count_col else 0,
                        "group_count": parse_number_value(cells.get((row_index, group_count_col), "")) if group_count_col else 0,
                        "detection_type_text": display_cell_text(cells.get((row_index, detection_type_col), "")) if detection_type_col else "",
                        "remark_text": display_cell_text(cells.get((row_index, remark_col), "")) if remark_col else "",
                        "min_value": parse_number_value(cells.get((row_index, min_value_col), "")) if min_value_col else 0,
                        "max_value": parse_number_value(cells.get((row_index, max_value_col), "")) if max_value_col else 0,
                        "min_value_text": display_cell_text(cells.get((row_index, min_value_col), "")) if min_value_col else "",
                        "max_value_text": display_cell_text(cells.get((row_index, max_value_col), "")) if max_value_col else "",
                        "pass_rate_value": parse_number_value(cells.get((row_index, pass_rate_col), "")) if pass_rate_col else 0,
                        "class_one_qty": parse_number_value(cells.get((row_index, class_one_col), "")) if class_one_col else 0,
                        "class_two_qty": parse_number_value(cells.get((row_index, class_two_col), "")) if class_two_col else 0,
                        "class_three_qty": parse_number_value(cells.get((row_index, class_three_col), "")) if class_three_col else 0,
                        "class_four_qty": parse_number_value(cells.get((row_index, class_four_col), "")) if class_four_col else 0,
                        "grout_hole_1": display_cell_text(cells.get((row_index, grout_hole_col), "")) if grout_hole_col else "",
                        "grout_hole_2": display_cell_text(cells.get((row_index, grout_hole_col + 1), "")) if grout_hole_col else "",
                        "grout_station_1": display_cell_text(cells.get((row_index, grout_station_col), "")) if grout_station_col else "",
                        "grout_station_2": display_cell_text(cells.get((row_index, grout_station_col + 1), "")) if grout_station_col else "",
                        "grout_elevation_1": display_cell_text(cells.get((row_index, grout_elevation_col), "")) if grout_elevation_col else "",
                        "grout_elevation_2": display_cell_text(cells.get((row_index, grout_elevation_col + 1), "")) if grout_elevation_col else "",
                        "grout_pressure": display_cell_text(cells.get((row_index, grout_pressure_col), "")) if grout_pressure_col else "",
                        "grout_measured_pressure_1": display_cell_text(cells.get((row_index, grout_measured_pressure_col), "")) if grout_measured_pressure_col else "",
                        "grout_measured_pressure_2": display_cell_text(cells.get((row_index, grout_measured_pressure_col + 1), "")) if grout_measured_pressure_col else "",
                        "grout_requirement": display_cell_text(cells.get((row_index, grout_requirement_col), "")) if grout_requirement_col else "",
                        "grout_volume_1": display_cell_text(cells.get((row_index, grout_volume_col), "")) if grout_volume_col else "",
                        "grout_volume_2": display_cell_text(cells.get((row_index, grout_volume_col + 1), "")) if grout_volume_col else "",
                        "grout_result": display_cell_text(cells.get((row_index, grout_result_col), "")) if grout_result_col else "",
                        "row_count": 1,
                        "ledger_order": len(records),
                    }
                )
    finally:
        conn.close()
    return records


def filter_records_by_units(records, unit_filters=None):
    unit_filters = {value for value in (unit_filters or []) if value}
    if not unit_filters:
        return records
    return [record for record in records if record.get("unit_name") in unit_filters]


def unit_options_from_records(records):
    seen = set()
    options = []
    for record in records:
        unit_name = (record.get("unit_name") or "").strip()
        if not unit_name or unit_name == "未识别单元工程" or unit_name in seen:
            continue
        seen.add(unit_name)
        options.append((unit_name, unit_name))
    return options


def statistics_unit_options(source_filters=None, sheet_filters=None):
    records = collect_stat_records(dt.date(1900, 1, 1), dt.date(2999, 12, 31), source_filters, sheet_filters)
    return unit_options_from_records(records)


def stat_depth_range_text(record):
    ranges = []
    if record.get("depth_start_1") or record.get("depth_end_1"):
        ranges.append(f"{record.get('depth_start_1', '')}-{record.get('depth_end_1', '')}".strip("-"))
    if record.get("depth_start_2") or record.get("depth_end_2"):
        ranges.append(f"{record.get('depth_start_2', '')}-{record.get('depth_end_2', '')}".strip("-"))
    return "；".join(value for value in ranges if value)


def aggregate_stat_records(records):
    grouped = {}
    for record in records:
        key = (
            record["source_type"],
            record["section_name"],
            record["report_table_type"],
            record["part_name"],
        )
        item = grouped.setdefault(
            key,
            {
                "source_type": record["source_type"],
                "client_type": record["client_type"],
                "section_name": record["section_name"],
                "sheet_name": record["sheet_name"],
                "report_table_type": record["report_table_type"],
                "unit_name": record["unit_name"],
                "sub_unit_name": record.get("sub_unit_name", ""),
                "location_name": record["location_name"],
                "part_name": record["part_name"],
                "row_count": 0,
                "construction_qty": 0,
                "testing_qty": 0,
                "mileage_qty": 0,
                "hole_count": 0,
                "group_count": 0,
                "ratio_value": 0,
                "report_nos": [],
                "entrust_nos": [],
                "sub_unit_names": [],
                "result_texts": [],
                "hole_nos": [],
                "depth_ranges": [],
                "mileage_texts": [],
                "anchor_nos": [],
                "spec_texts": [],
                "anchor_length_texts": [],
                "anchorage_length_texts": [],
                "free_length_texts": [],
                "elastic_modulus_texts": [],
                "load_requirement_texts": [],
                "design_anchor_force_texts": [],
                "pile_type_texts": [],
                "pile_diameter_texts": [],
                "pile_top_texts": [],
                "category_texts": [],
                "detection_type_texts": [],
                "remark_texts": [],
                "length_values": [],
                "diameter_values": [],
                "design_force_values": [],
                "min_value_texts": [],
                "max_value_texts": [],
                "min_value": None,
                "max_value": None,
                "pass_rate_value": 0,
                "class_one_qty": 0,
                "class_two_qty": 0,
                "class_three_qty": 0,
                "class_four_qty": 0,
                "grout_hole_1": record.get("grout_hole_1", ""),
                "grout_hole_2": record.get("grout_hole_2", ""),
                "grout_station_1": record.get("grout_station_1", ""),
                "grout_station_2": record.get("grout_station_2", ""),
                "grout_elevation_1": record.get("grout_elevation_1", ""),
                "grout_elevation_2": record.get("grout_elevation_2", ""),
                "grout_pressure": record.get("grout_pressure", ""),
                "grout_measured_pressure_1": record.get("grout_measured_pressure_1", ""),
                "grout_measured_pressure_2": record.get("grout_measured_pressure_2", ""),
                "grout_requirement": record.get("grout_requirement", ""),
                "grout_volume_1": record.get("grout_volume_1", ""),
                "grout_volume_2": record.get("grout_volume_2", ""),
                "grout_result": record.get("grout_result", ""),
            },
        )
        item["row_count"] += record["row_count"]
        item["construction_qty"] += record["construction_qty"]
        item["testing_qty"] += record["testing_qty"]
        item["mileage_qty"] += record.get("mileage_value", 0)
        item["hole_count"] += record.get("hole_count", 0)
        item["group_count"] += record.get("group_count", 0)
        item["ratio_value"] = record["ratio_value"] or item["ratio_value"]
        append_unique_text(item["report_nos"], record.get("report_no", ""))
        append_unique_text(item["entrust_nos"], record.get("entrust_no", ""))
        append_unique_text(item["sub_unit_names"], record.get("sub_unit_name", ""))
        append_unique_text(item["result_texts"], record.get("result_text", ""))
        append_unique_text(item["hole_nos"], record.get("hole_no", ""))
        append_unique_text(item["depth_ranges"], stat_depth_range_text(record))
        append_unique_text(item["mileage_texts"], record.get("mileage_text", ""))
        append_unique_text(item["anchor_nos"], record.get("anchor_no", ""))
        append_unique_text(item["spec_texts"], record.get("spec_text", ""))
        append_unique_text(item["anchor_length_texts"], record.get("anchor_length_text", ""))
        append_unique_text(item["anchorage_length_texts"], record.get("anchorage_length_text", ""))
        append_unique_text(item["free_length_texts"], record.get("free_length_text", ""))
        append_unique_text(item["elastic_modulus_texts"], record.get("elastic_modulus_text", ""))
        append_unique_text(item["load_requirement_texts"], record.get("load_requirement_text", ""))
        append_unique_text(item["design_anchor_force_texts"], record.get("design_anchor_force_text", ""))
        append_unique_text(item["pile_type_texts"], record.get("pile_type_text", ""))
        append_unique_text(item["pile_diameter_texts"], record.get("pile_diameter_text", ""))
        append_unique_text(item["pile_top_texts"], record.get("pile_top_text", ""))
        append_unique_text(item["category_texts"], record.get("category_text", ""))
        append_unique_text(item["detection_type_texts"], record.get("detection_type_text", ""))
        append_unique_text(item["remark_texts"], record.get("remark_text", ""))
        append_unique_text(item["length_values"], record.get("length_value", ""))
        if record["diameter_value"]:
            if record["diameter_value"] not in item["diameter_values"]:
                item["diameter_values"].append(record["diameter_value"])
        if record["design_force_value"]:
            if record["design_force_value"] not in item["design_force_values"]:
                item["design_force_values"].append(record["design_force_value"])
        append_unique_text(item["min_value_texts"], record.get("min_value_text", ""))
        append_unique_text(item["max_value_texts"], record.get("max_value_text", ""))
        if record["min_value"]:
            item["min_value"] = record["min_value"] if item["min_value"] is None else min(item["min_value"], record["min_value"])
        if record["max_value"]:
            item["max_value"] = record["max_value"] if item["max_value"] is None else max(item["max_value"], record["max_value"])
        item["pass_rate_value"] = record["pass_rate_value"] or item["pass_rate_value"]
        item["class_one_qty"] += record["class_one_qty"]
        item["class_two_qty"] += record["class_two_qty"]
        item["class_three_qty"] += record["class_three_qty"]
        item["class_four_qty"] += record["class_four_qty"]
    result = []
    for item in grouped.values():
        item["report_no_text"] = "\n".join(item.pop("report_nos"))
        item["entrust_no_text"] = "\n".join(item.pop("entrust_nos"))
        sub_unit_names = item.pop("sub_unit_names")
        if sub_unit_names:
            item["sub_unit_name"] = "\n".join(sub_unit_names)
        item["result_text"] = "\n".join(item.pop("result_texts"))
        item["hole_no_text"] = "\n".join(item.pop("hole_nos"))
        item["depth_range_text"] = "\n".join(item.pop("depth_ranges"))
        item["mileage_text"] = "\n".join(item.pop("mileage_texts"))
        item["anchor_no_text"] = "\n".join(item.pop("anchor_nos"))
        item["spec_text"] = "\n".join(item.pop("spec_texts"))
        item["anchor_length_text"] = "\n".join(item.pop("anchor_length_texts"))
        item["anchorage_length_text"] = "\n".join(item.pop("anchorage_length_texts"))
        item["free_length_text"] = "\n".join(item.pop("free_length_texts"))
        item["elastic_modulus_text"] = "\n".join(item.pop("elastic_modulus_texts"))
        item["load_requirement_text"] = "\n".join(item.pop("load_requirement_texts"))
        item["design_anchor_force_text"] = "\n".join(item.pop("design_anchor_force_texts"))
        item["pile_type_text"] = "\n".join(item.pop("pile_type_texts"))
        item["pile_diameter_text"] = "\n".join(item.pop("pile_diameter_texts"))
        item["pile_top_text"] = "\n".join(item.pop("pile_top_texts"))
        item["category_text"] = "\n".join(item.pop("category_texts"))
        item["detection_type_text"] = "\n".join(item.pop("detection_type_texts"))
        item["remark_text"] = "\n".join(item.pop("remark_texts"))
        item["length_text"] = "\n".join(item.pop("length_values"))
        item["diameter_text"] = "\n".join(item.pop("diameter_values"))
        item["design_force_text"] = "\n".join(item.pop("design_force_values"))
        item["min_value_text"] = "\n".join(item.pop("min_value_texts"))
        item["max_value_text"] = "\n".join(item.pop("max_value_texts"))
        if not item["ratio_value"] and item["construction_qty"]:
            item["ratio_value"] = item["testing_qty"] / item["construction_qty"] * 100
        result.append(item)
    return result


def format_stat_number(value):
    try:
        number = float(value or 0)
    except (TypeError, ValueError):
        return "0"
    if abs(number - round(number)) < 0.000001:
        return str(int(round(number)))
    return f"{number:.2f}"


def format_significant_number(value, digits=3):
    text = display_cell_text(value).strip()
    if not text:
        return ""
    try:
        number = float(text.replace(",", ""))
    except ValueError:
        return text
    formatted = f"{number:.{digits}g}"
    if "e" in formatted or "E" in formatted:
        formatted = f"{number:.{digits}f}".rstrip("0").rstrip(".")
    return formatted


def format_decimal_number(value, places=3):
    text = display_cell_text(value).strip()
    if not text:
        return ""
    try:
        number = float(text.replace(",", ""))
    except ValueError:
        return text
    return f"{number:.{places}f}"


def append_unique_text(target, value):
    text = display_cell_text(value).strip()
    if text and text not in target:
        target.append(text)


def stat_value_text(item, text_key, value_key):
    text = item.get(text_key, "")
    if text:
        return html.escape(text).replace("\n", "<br>")
    return html.escape(format_stat_number(item[value_key])) if item.get(value_key) is not None else ""


def plain_stat_value_text(item, text_key, value_key):
    text = item.get(text_key, "")
    if text:
        return text
    return format_stat_number(item[value_key]) if item.get(value_key) is not None else ""


def stat_text(value):
    return html.escape(value or "").replace("\n", "<br>")


def statistics_sheet_options():
    conn = connect()
    try:
        rows = conn.execute(
            """
            SELECT DISTINCT s.sheet_name
            FROM template_sheet s
            JOIN template_workbook w ON w.id = s.workbook_id
            JOIN ledger_file_version v ON v.workbook_id = w.id
            JOIN ledger_file f ON f.current_version_id = v.id
            ORDER BY s.sheet_name
            """
        ).fetchall()
    finally:
        conn.close()
    return [(row["sheet_name"], clean_sheet_display_name(row["sheet_name"])) for row in rows]




def query_value(params, key, default=""):
    return params.get(key, [default])[0]


def query_values(params, key):
    return [str(value).strip() for value in params.get(key, []) if str(value).strip()]


def ledger_distinct_options(column):
    allowed_columns = {"section_name": "section_name", "source_type": "source_type", "discipline": "discipline"}
    if column not in allowed_columns:
        return []
    conn = connect()
    try:
        rows = conn.execute(
            f"""
            SELECT DISTINCT {allowed_columns[column]} AS value
            FROM ledger_file
            WHERE {allowed_columns[column]} IS NOT NULL AND TRIM({allowed_columns[column]}) <> ''
            ORDER BY {allowed_columns[column]}
            """
        ).fetchall()
    finally:
        conn.close()
    return [row["value"] for row in rows]


def field_text(cells, row_index, col_index):
    return display_cell_text(cells.get((row_index, col_index), "")).strip() if col_index else ""


def collect_ledger_row_records(section_filters=None, source_filters=None, sheet_filters=None):
    section_filters = [value for value in (section_filters or []) if value]
    source_filters = [value for value in (source_filters or []) if value]
    sheet_filters = [value for value in (sheet_filters or []) if value]
    conn = connect()
    records = []
    try:
        where_parts = ["1 = 1"]
        args = []
        if section_filters:
            where_parts.append(f"f.section_name IN ({','.join('?' for _ in section_filters)})")
            args.extend(section_filters)
        if source_filters:
            where_parts.append(f"f.source_type IN ({','.join('?' for _ in source_filters)})")
            args.extend(source_filters)
        if sheet_filters:
            where_parts.append(f"s.sheet_name IN ({','.join('?' for _ in sheet_filters)})")
            args.extend(sheet_filters)
        sheets = conn.execute(
            f"""
            SELECT s.id AS sheet_id, s.sheet_name, s.max_row, s.max_column,
                   f.id AS file_id, f.project_name, f.section_name, f.source_type,
                   f.original_filename
            FROM template_sheet s
            JOIN template_workbook w ON w.id = s.workbook_id
            JOIN ledger_file_version v ON v.workbook_id = w.id
            JOIN ledger_file f ON f.current_version_id = v.id
            WHERE {' AND '.join(where_parts)}
            ORDER BY f.id ASC, s.sheet_index ASC, s.id ASC
            """,
            args,
        ).fetchall()
        for sheet in sheets:
            rows = conn.execute(
                """
                SELECT row_index, col_index, raw_value
                FROM template_cell
                WHERE sheet_id = ?
                """,
                (sheet["sheet_id"],),
            ).fetchall()
            cells = {(row["row_index"], row["col_index"]): row["raw_value"] for row in rows}
            if not cells:
                continue
            max_column = sheet["max_column"] or 0
            header_row = detect_header_row(cells, sheet["max_row"] or 20)
            header_map = {col: cells.get((header_row, col), "") for col in range(1, max_column + 1)}
            summary_start_col = 20
            date_col = find_col_by_keywords(header_map, ["检测日期", "试验日期", "检测时间"])
            report_col = find_col_by_keywords(header_map, ["报告编号", "报告号"])
            entrust_col = find_col_by_keywords(header_map, ["委托编号", "委托号"])
            result_col = find_col_by_keywords(header_map, ["检测结果", "试验结果", "评定结果", "综合评判类别", "结论"])
            unit_col = find_col_by_keywords(header_map, ["单位工程", "单元工程"])
            sub_unit_col = find_col_by_keywords(header_map, ["分部工程"])
            item_col = find_col_by_keywords(header_map, ["单元工程"])
            location_col = find_col_by_keywords(header_map, ["工程部位", "部位", "施工部位"], summary_start_col)
            construction_col = find_col_by_keywords(header_map, ["施工数量", "施工量"], summary_start_col)
            testing_col = find_col_by_keywords(header_map, ["检测数量", "检测量", "检测里程"], summary_start_col)
            ratio_col = find_col_by_keywords(header_map, ["抽检比例", "检测比例"], summary_start_col)
            if not location_col:
                location_col = find_col_by_keywords(header_map, ["工程部位", "部位", "施工部位"])
            if not construction_col:
                construction_col = find_col_by_keywords(header_map, ["施工数量", "施工量"])
            if not testing_col:
                testing_col = find_col_by_keywords(header_map, ["检测数量", "检测量", "检测里程"])
            sheet_display_name = clean_sheet_display_name(sheet["sheet_name"])
            for row_index in range(header_row + 1, (sheet["max_row"] or 0) + 1):
                values = [display_cell_text(cells.get((row_index, col), "")).strip() for col in range(1, max_column + 1)]
                if not any(values):
                    continue
                report_no = field_text(cells, row_index, report_col)
                entrust_no = field_text(cells, row_index, entrust_col)
                date_text = field_text(cells, row_index, date_col)
                unit_name = field_text(cells, row_index, unit_col)
                sub_unit_name = field_text(cells, row_index, sub_unit_col)
                item_name = field_text(cells, row_index, item_col)
                location_name = field_text(cells, row_index, location_col)
                if not any([report_no, entrust_no, date_text, unit_name, location_name]):
                    continue
                construction_qty = parse_number_value(cells.get((row_index, construction_col), "")) if construction_col else 0
                testing_qty = parse_number_value(cells.get((row_index, testing_col), "")) if testing_col else 0
                records.append({
                    "file_id": sheet["file_id"],
                    "sheet_id": sheet["sheet_id"],
                    "row_index": row_index,
                    "project_name": sheet["project_name"] or "",
                    "section_name": sheet["section_name"] or "",
                    "source_type": sheet["source_type"] or "",
                    "original_filename": sheet["original_filename"] or "",
                    "sheet_name": sheet_display_name,
                    "raw_sheet_name": sheet["sheet_name"] or "",
                    "report_no": report_no,
                    "entrust_no": entrust_no,
                    "detect_date": parse_date_value(date_text),
                    "detect_date_text": date_text,
                    "unit_name": unit_name,
                    "sub_unit_name": sub_unit_name,
                    "item_name": item_name,
                    "location_name": location_name,
                    "result_text": field_text(cells, row_index, result_col),
                    "construction_qty": construction_qty,
                    "testing_qty": testing_qty,
                    "ratio_value": parse_number_value(cells.get((row_index, ratio_col), "")) if ratio_col else 0,
                    "result_column_found": bool(result_col),
                    "construction_column_found": bool(construction_col),
                    "testing_column_found": bool(testing_col),
                })
    finally:
        conn.close()
    return records


def advanced_query_params(params):
    start_text = query_value(params, "start_date", "")
    end_text = query_value(params, "end_date", "")
    start_date = parse_date_value(start_text) if start_text else None
    end_date = parse_date_value(end_text) if end_text else None
    if start_date and end_date and start_date > end_date:
        start_date, end_date = end_date, start_date
        start_text, end_text = end_text, start_text
    try:
        limit = int(query_value(params, "limit", "500") or "500")
    except ValueError:
        limit = 500
    limit = max(50, min(limit, 5000))
    return {
        "section_filters": query_values(params, "section_name"),
        "source_filters": query_values(params, "source_type"),
        "sheet_filters": query_values(params, "sheet_name"),
        "start_date": start_date,
        "end_date": end_date,
        "start_text": start_text,
        "end_text": end_text,
        "report_no": query_value(params, "report_no", "").strip(),
        "entrust_no": query_value(params, "entrust_no", "").strip(),
        "unit_name": query_value(params, "unit_name", "").strip(),
        "location_name": query_value(params, "location_name", "").strip(),
        "result_text": query_value(params, "result_text", "").strip(),
        "filename": query_value(params, "filename", "").strip(),
        "limit": limit,
    }


def text_contains(value, keyword):
    if not keyword:
        return True
    return keyword.lower() in str(value or "").lower()


def filter_advanced_query_records(records, query):
    filtered = []
    for record in records:
        detect_date = record.get("detect_date")
        if query["start_date"] and (not detect_date or detect_date < query["start_date"]):
            continue
        if query["end_date"] and (not detect_date or detect_date > query["end_date"]):
            continue
        if not text_contains(record.get("report_no"), query["report_no"]):
            continue
        if not text_contains(record.get("entrust_no"), query["entrust_no"]):
            continue
        if not text_contains(record.get("unit_name"), query["unit_name"]):
            continue
        location_blob = " ".join([record.get("location_name", ""), record.get("sub_unit_name", ""), record.get("item_name", "")])
        if not text_contains(location_blob, query["location_name"]):
            continue
        if not text_contains(record.get("result_text"), query["result_text"]):
            continue
        if not text_contains(record.get("original_filename"), query["filename"]):
            continue
        filtered.append(record)
    return filtered


def advanced_query_records(params):
    query = advanced_query_params(params)
    records = collect_ledger_row_records(query["section_filters"], query["source_filters"], query["sheet_filters"])
    return query, filter_advanced_query_records(records, query)


def advanced_query_export_url(params):
    pairs = []
    for key, values in params.items():
        if key == "limit":
            continue
        for value in values:
            if str(value).strip():
                pairs.append((key, value))
    query = urllib.parse.urlencode(pairs)
    return "/advanced_query_export" + (f"?{query}" if query else "")


def query_record_date_text(record):
    detect_date = record.get("detect_date")
    return detect_date.isoformat() if detect_date else record.get("detect_date_text", "")


def render_advanced_query_table(records):
    if not records:
        return "<tr><td colspan='13'>没有符合条件的数据</td></tr>"
    rows = []
    for record in records:
        rows.append(
            "<tr>"
            f"<td>{html.escape(record.get('section_name', ''))}</td>"
            f"<td>{html.escape(record.get('source_type', ''))}</td>"
            f"<td>{html.escape(record.get('original_filename', ''))}</td>"
            f"<td>{html.escape(record.get('sheet_name', ''))}</td>"
            f"<td>{html.escape(str(record.get('row_index', '')))}</td>"
            f"<td>{html.escape(query_record_date_text(record))}</td>"
            f"<td>{html.escape(record.get('report_no', ''))}</td>"
            f"<td>{html.escape(record.get('entrust_no', ''))}</td>"
            f"<td>{html.escape(record.get('unit_name', ''))}</td>"
            f"<td>{html.escape(record.get('location_name', ''))}</td>"
            f"<td>{html.escape(record.get('result_text', ''))}</td>"
            f"<td>{html.escape(format_stat_number(record.get('construction_qty', 0)))}</td>"
            f"<td>{html.escape(format_stat_number(record.get('testing_qty', 0)))}</td>"
            "</tr>"
        )
    return "".join(rows)


def build_advanced_query_xlsx(records):
    wb = Workbook()
    ws = wb.active
    ws.title = "高级查询结果"
    headers = ["标段", "委托单位", "原文件名", "工作表", "Excel行号", "检测日期", "报告编号", "委托编号", "单位工程", "分部工程", "单元工程", "工程部位", "检测结果", "施工数量", "检测数量", "检测比例"]
    ws.append(headers)
    for record in records:
        ws.append([
            record.get("section_name", ""), record.get("source_type", ""), record.get("original_filename", ""),
            record.get("sheet_name", ""), record.get("row_index", ""), query_record_date_text(record),
            record.get("report_no", ""), record.get("entrust_no", ""), record.get("unit_name", ""),
            record.get("sub_unit_name", ""), record.get("item_name", ""), record.get("location_name", ""),
            record.get("result_text", ""), record.get("construction_qty", 0), record.get("testing_qty", 0),
            record.get("ratio_value", 0),
        ])
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for index, width in enumerate([12, 36, 24, 18, 10, 14, 24, 24, 24, 24, 24, 36, 14, 12, 12, 12], start=1):
        ws.column_dimensions[col_letter(index)].width = width
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def record_location_text(record):
    return f"{record.get('original_filename', '')} / {record.get('sheet_name', '')} / 第{record.get('row_index', '')}行"


def duplicate_context_summaries(records):
    grouped = {}
    order = []
    for record in records:
        key = (record.get("file_id"), record.get("sheet_id"))
        if key not in grouped:
            grouped[key] = {
                "record": record,
                "section_name": record.get("section_name", ""),
                "sheet_name": record.get("sheet_name", ""),
                "dates": [],
                "construction_qty": 0,
                "testing_qty": 0,
                "locations": [],
            }
            order.append(key)
        item = grouped[key]
        append_unique_text(item["dates"], query_record_date_text(record))
        item["construction_qty"] += record.get("construction_qty", 0) or 0
        item["testing_qty"] += record.get("testing_qty", 0) or 0
        append_unique_text(item["locations"], record_location_text(record))
    summaries = []
    for key in order:
        item = grouped[key]
        summaries.append({
            "record": item["record"],
            "section_name": item["section_name"],
            "sheet_name": item["sheet_name"],
            "date_text": "/".join(item["dates"]),
            "quantity_text": f"施工{format_stat_number(item['construction_qty'])}/检测{format_stat_number(item['testing_qty'])}",
            "location_text": item["locations"][0] if item["locations"] else record_location_text(item["record"]),
        })
    return summaries


def conflict_value_details(summaries, key):
    values = []
    for item in summaries:
        value = display_cell_text(item.get(key, "")).strip()
        if value and value not in values:
            values.append(value)
    return values


def duplicate_conflict_details(records):
    summaries = duplicate_context_summaries(records)
    if len(summaries) <= 1:
        return [], summaries
    checks = [
        ("标段", "section_name"),
        ("检测类型", "sheet_name"),
        ("日期", "date_text"),
        ("数量", "quantity_text"),
    ]
    conflicts = []
    for label, key in checks:
        values = conflict_value_details(summaries, key)
        if len(values) > 1:
            conflicts.append(f"{label}不一致（{'、'.join(values[:6])}）")
    return conflicts, summaries


def add_quality_issue(issues, issue_type, level, message, record=None, file_row=None):
    source = record or file_row or {}
    issues.append({
        "issue_type": issue_type,
        "level": level,
        "message": message,
        "file": source.get("original_filename", ""),
        "sheet": (record or {}).get("sheet_name", ""),
        "row_index": (record or {}).get("row_index", ""),
        "report_no": (record or {}).get("report_no", ""),
        "entrust_no": (record or {}).get("entrust_no", ""),
    })


def filename_expected_section(filename):
    upper_name = str(filename or "").upper()
    for section in ("Q1", "Q2", "C1", "C2", "C3", "C4"):
        if section in upper_name:
            return section
    return ""


def filename_expected_client_type(filename):
    upper_name = str(filename or "").upper()
    if "JW" in upper_name:
        return "监理"
    if "SW" in upper_name:
        return "施工"
    return ""


def collect_file_quality_issues(issues):
    conn = connect()
    try:
        rows = conn.execute(
            """
            SELECT f.id, f.original_filename, f.section_name, f.source_type, COUNT(s.id) AS sheet_count
            FROM ledger_file f
            LEFT JOIN ledger_file_version v ON v.id = f.current_version_id
            LEFT JOIN template_workbook w ON w.id = v.workbook_id
            LEFT JOIN template_sheet s ON s.workbook_id = w.id
            GROUP BY f.id, f.original_filename, f.section_name, f.source_type
            ORDER BY f.id ASC
            """
        ).fetchall()
    finally:
        conn.close()
    for row in rows:
        file_row = dict(row)
        expected_section = filename_expected_section(row["original_filename"])
        if expected_section and expected_section not in (row["section_name"] or ""):
            add_quality_issue(issues, "metadata_mismatch", "高", f"文件名显示为 {expected_section}，但台账标段为 {row['section_name'] or '未填写'}。", file_row=file_row)
        expected_client = filename_expected_client_type(row["original_filename"])
        actual_client = client_short_type(row["source_type"] or "")
        if expected_client and actual_client != expected_client:
            add_quality_issue(issues, "metadata_mismatch", "中", f"文件名显示为{expected_client}台账，但委托单位识别为{actual_client}。", file_row=file_row)
        if not row["sheet_count"]:
            add_quality_issue(issues, "sheet_missing", "高", "文件没有解析出任何工作表。", file_row=file_row)


def build_quality_issues():
    records = collect_ledger_row_records()
    issues = []
    collect_file_quality_issues(issues)
    report_map = {}
    entrust_map = {}
    for record in records:
        report_no = (record.get("report_no") or "").strip()
        entrust_no = (record.get("entrust_no") or "").strip()
        if report_no:
            report_map.setdefault(report_no, []).append(record)
        if entrust_no:
            entrust_map.setdefault(entrust_no, []).append(record)
        if not report_no:
            add_quality_issue(issues, "missing_report", "高", "报告编号为空，后续归档和追溯会受影响。", record=record)
        if not entrust_no:
            add_quality_issue(issues, "missing_entrust", "中", "委托编号为空，无法完整关联委托资料。", record=record)
        if not record.get("detect_date"):
            add_quality_issue(issues, "missing_date", "高", "检测日期为空或无法识别。", record=record)
        if record.get("result_column_found") and not (record.get("result_text") or "").strip():
            add_quality_issue(issues, "missing_result", "中", "检测结果为空。", record=record)
        if not (record.get("unit_name") or record.get("location_name")):
            add_quality_issue(issues, "missing_location", "中", "单位工程和工程部位均为空。", record=record)
        construction_qty = record.get("construction_qty", 0) or 0
        testing_qty = record.get("testing_qty", 0) or 0
        if construction_qty < 0 or testing_qty < 0:
            add_quality_issue(issues, "quantity_abnormal", "高", "数量字段出现负数。", record=record)
        if record.get("construction_column_found") and record.get("testing_column_found"):
            if construction_qty > 0 and testing_qty > construction_qty:
                add_quality_issue(issues, "quantity_abnormal", "中", "检测数量大于施工数量，请核对数量单位或录入值。", record=record)
            if construction_qty > 0 and testing_qty == 0:
                add_quality_issue(issues, "quantity_abnormal", "中", "有施工数量但检测数量为 0。", record=record)
    for report_no, items in report_map.items():
        conflicts, summaries = duplicate_conflict_details(items)
        if conflicts:
            locations = "；".join(summary["location_text"] for summary in summaries[:5])
            add_quality_issue(issues, "report_no_conflict", "高", f"报告编号 {report_no} 存在冲突：{'；'.join(conflicts)}。涉及位置：{locations}", record=items[0])
    for entrust_no, items in entrust_map.items():
        conflicts, summaries = duplicate_conflict_details(items)
        if conflicts:
            locations = "；".join(summary["location_text"] for summary in summaries[:5])
            add_quality_issue(issues, "entrust_no_conflict", "中", f"委托编号 {entrust_no} 存在冲突：{'；'.join(conflicts)}。涉及位置：{locations}", record=items[0])
    return records, issues


def issue_type_label(issue_type):
    return {
        "metadata_mismatch": "文件信息不一致", "sheet_missing": "工作表缺失", "missing_report": "报告编号缺失",
        "missing_entrust": "委托编号缺失", "missing_date": "检测日期缺失", "missing_result": "检测结果缺失",
        "missing_location": "工程部位缺失", "quantity_abnormal": "数量异常", "report_no_conflict": "报告编号冲突",
        "entrust_no_conflict": "委托编号冲突",
    }.get(issue_type, issue_type)


def render_quality_issue_table(issues):
    if not issues:
        return "<p class='muted'>当前未发现异常。</p>"
    grouped = {}
    for issue in issues:
        grouped.setdefault(issue["issue_type"], []).append(issue)
    sections = []
    for issue_type, items in grouped.items():
        rows = []
        for issue in items[:200]:
            rows.append(
                "<tr>"
                f"<td>{html.escape(issue.get('level', ''))}</td>"
                f"<td>{html.escape(issue.get('message', ''))}</td>"
                f"<td>{html.escape(issue.get('file', ''))}</td>"
                f"<td>{html.escape(issue.get('sheet', ''))}</td>"
                f"<td>{html.escape(str(issue.get('row_index', '')))}</td>"
                f"<td>{html.escape(issue.get('report_no', ''))}</td>"
                f"<td>{html.escape(issue.get('entrust_no', ''))}</td>"
                "</tr>"
            )
        more = f"<p class='muted'>仅显示前 200 条，共 {len(items)} 条。</p>" if len(items) > 200 else ""
        sections.append(
            f"<details class='stat-section' open><summary>{html.escape(issue_type_label(issue_type))}（{len(items)}）</summary>"
            "<table class='ledger-table'><thead><tr><th>级别</th><th>问题说明</th><th>文件</th><th>工作表</th><th>行号</th><th>报告编号</th><th>委托编号</th></tr></thead>"
            f"<tbody>{''.join(rows)}</tbody></table>{more}</details>"
        )
    return "".join(sections)

def stat_ratio_text(item):
    return f"{item['ratio_value']:.2f}%" if item.get("ratio_value") else ""


def stat_min_text(item):
    return stat_value_text(item, "min_value_text", "min_value")


def stat_max_text(item):
    return stat_value_text(item, "max_value_text", "max_value")


def plain_stat_min_text(item):
    return plain_stat_value_text(item, "min_value_text", "min_value")


def plain_stat_max_text(item):
    return plain_stat_value_text(item, "max_value_text", "max_value")


def stat_pass_rate_text(item):
    return f"{item['pass_rate_value']:.2f}%" if item.get("pass_rate_value") else ""


def render_stat_detail_row(item):
    return (
        "<tr>"
        f"<td>{html.escape(item['source_type'])}</td>"
        f"<td>{html.escape(item['section_name'])}</td>"
        f"<td>{html.escape(item['client_type'])}</td>"
        f"<td>{html.escape(item['sheet_name'])}</td>"
        f"<td>{html.escape(item['report_table_type'])}</td>"
        f"<td>{html.escape(item['part_name'])}</td>"
        f"<td>{item['row_count']}</td>"
        f"<td>{format_stat_number(item['construction_qty'])}</td>"
        f"<td>{format_stat_number(item['testing_qty'])}</td>"
        f"<td>{stat_ratio_text(item)}</td>"
        f"<td>{stat_text(item['diameter_text'])}</td>"
        f"<td>{stat_text(item['design_force_text'])}</td>"
        f"<td>{stat_min_text(item)}</td>"
        f"<td>{stat_max_text(item)}</td>"
        f"<td>{stat_pass_rate_text(item)}</td>"
        f"<td>{format_stat_number(item['class_one_qty'])}</td>"
        f"<td>{format_stat_number(item['class_two_qty'])}</td>"
        f"<td>{format_stat_number(item['class_three_qty'])}</td>"
        f"<td>{format_stat_number(item['class_four_qty'])}</td>"
        "</tr>"
    )


STAT_DETAIL_HEADER = (
    "<tr><th>委托单位</th><th>标段</th><th>委托类别</th><th>检测项目</th><th>第五章表格类型</th>"
    "<th>工程部位/单元工程</th><th>检测组数</th><th>施工数量</th><th>检测数量</th><th>抽检/检测比例</th>"
    "<th>锚杆直径</th><th>设计拉拔力</th><th>最小值</th><th>最大值</th><th>合格率</th>"
    "<th>一类</th><th>二类</th><th>三类</th><th>四类</th></tr>"
)


NONDESTRUCTIVE_HEADER = (
    "<tr><th rowspan='2'>工程部位</th><th rowspan='2'>施工数量(根)</th><th rowspan='2'>检测数量(根)</th>"
    "<th rowspan='2'>抽检比例(%)</th><th colspan='4'>锚杆分级</th><th rowspan='2'>合格率(%)</th></tr>"
    "<tr><th>I</th><th>II</th><th>III</th><th>IV</th></tr>"
)


PULLOUT_HEADER = (
    "<tr><th rowspan='2'>工程部位</th><th rowspan='2'>施工数量(根)</th><th rowspan='2'>检测数量(根)</th>"
    "<th rowspan='2'>抽检比例(%)</th><th rowspan='2'>锚杆直径(mm)</th><th rowspan='2'>设计拉拔力值(KN)</th>"
    "<th colspan='2'>实测拉拔力值(KN)</th><th rowspan='2'>合格率(%)</th></tr>"
    "<tr><th>最小值</th><th>最大值</th></tr>"
)


GROUT_HEADER = (
    "<tr><th rowspan='2'>工程部位</th><th colspan='2'>检查孔编号</th><th colspan='2'>检查桩号</th>"
    "<th colspan='2'>孔口高程</th><th rowspan='2'>设计压力<br>(Mpa）</th>"
    "<th colspan='2'>实测压力<br>(Mpa）</th>"
    "<th rowspan='2'>设计要求（规定压力下，初始10min内注浆量）(L)</th>"
    "<th colspan='2'>初始10min内注浆量（L)</th><th rowspan='2'>检测结果</th></tr>"
    "<tr><th>1</th><th>2</th><th>1</th><th>2</th><th>1</th><th>2</th><th>1</th><th>2</th><th>1</th><th>2</th></tr>"
)


def render_nondestructive_row(item):
    ratio = f"{item['ratio_value']:.2f}" if item["ratio_value"] else ""
    pass_rate = f"{item['pass_rate_value']:.2f}" if item["pass_rate_value"] else "100.00"
    return (
        "<tr>"
        f"<td>{html.escape(item['part_name'])}</td>"
        f"<td>{format_stat_number(item['construction_qty'])}</td>"
        f"<td>{format_stat_number(item['testing_qty'])}</td>"
        f"<td>{ratio}</td>"
        f"<td>{format_stat_number(item['class_one_qty'])}</td>"
        f"<td>{format_stat_number(item['class_two_qty'])}</td>"
        f"<td>{format_stat_number(item['class_three_qty'])}</td>"
        f"<td>{format_stat_number(item['class_four_qty'])}</td>"
        f"<td>{pass_rate}</td>"
        "</tr>"
    )


def render_pullout_row(item):
    ratio = f"{item['ratio_value']:.2f}" if item["ratio_value"] else ""
    pass_rate = format_stat_number(item["pass_rate_value"]) if item.get("pass_rate_value") else "100"
    return (
        "<tr>"
        f"<td>{html.escape(item['part_name'])}</td>"
        f"<td>{format_stat_number(item['construction_qty'])}</td>"
        f"<td>{format_stat_number(item['testing_qty'])}</td>"
        f"<td>{ratio}</td>"
        f"<td>{stat_text(item['diameter_text'])}</td>"
        f"<td>{stat_text(item['design_force_text'])}</td>"
        f"<td>{stat_min_text(item)}</td>"
        f"<td>{stat_max_text(item)}</td>"
        f"<td>{pass_rate}</td>"
        "</tr>"
    )


def render_grout_row(item):
    return (
        "<tr>"
        f"<td>{html.escape(item['part_name'])}</td>"
        f"<td>{stat_text(item.get('grout_hole_1', ''))}</td>"
        f"<td>{stat_text(item.get('grout_hole_2', ''))}</td>"
        f"<td>{stat_text(item.get('grout_station_1', ''))}</td>"
        f"<td>{stat_text(item.get('grout_station_2', ''))}</td>"
        f"<td>{stat_text(item.get('grout_elevation_1', ''))}</td>"
        f"<td>{stat_text(item.get('grout_elevation_2', ''))}</td>"
        f"<td>{stat_text(item.get('grout_pressure', ''))}</td>"
        f"<td>{stat_text(item.get('grout_measured_pressure_1', ''))}</td>"
        f"<td>{stat_text(item.get('grout_measured_pressure_2', ''))}</td>"
        f"<td>{stat_text(item.get('grout_requirement', ''))}</td>"
        f"<td>{stat_text(format_decimal_number(item.get('grout_volume_1', ''), 3))}</td>"
        f"<td>{stat_text(format_decimal_number(item.get('grout_volume_2', ''), 3))}</td>"
        f"<td>{stat_text(item.get('grout_result', ''))}</td>"
        "</tr>"
    )


def render_pullout_total_row(items):
    construction = sum(item["construction_qty"] for item in items)
    testing = sum(item["testing_qty"] for item in items)
    ratio = f"{testing / construction * 100:.2f}" if construction else ""
    pass_rate_values = [item["pass_rate_value"] for item in items if item.get("pass_rate_value")]
    pass_rate = format_stat_number(min(pass_rate_values)) if pass_rate_values else ("100" if testing else "")
    return (
        "<tr class='stat-total'>"
        "<td>合计</td>"
        f"<td>{format_stat_number(construction)}</td>"
        f"<td>{format_stat_number(testing)}</td>"
        f"<td>{ratio}</td>"
        "<td>/</td>"
        "<td>/</td>"
        "<td>/</td>"
        "<td>/</td>"
        f"<td>{pass_rate}</td>"
        "</tr>"
    )


def render_nondestructive_total_row(items):
    construction = sum(item["construction_qty"] for item in items)
    testing = sum(item["testing_qty"] for item in items)
    class_one = sum(item["class_one_qty"] for item in items)
    class_two = sum(item["class_two_qty"] for item in items)
    class_three = sum(item["class_three_qty"] for item in items)
    class_four = sum(item["class_four_qty"] for item in items)
    ratio = f"{testing / construction * 100:.2f}" if construction else ""
    pass_rate = "100.00" if testing and (class_three + class_four == 0) else ""
    return (
        "<tr class='stat-total'>"
        "<td>合计</td>"
        f"<td>{format_stat_number(construction)}</td>"
        f"<td>{format_stat_number(testing)}</td>"
        f"<td>{ratio}</td>"
        f"<td>{format_stat_number(class_one)}</td>"
        f"<td>{format_stat_number(class_two)}</td>"
        f"<td>{format_stat_number(class_three)}</td>"
        f"<td>{format_stat_number(class_four)}</td>"
        f"<td>{pass_rate}</td>"
        "</tr>"
    )


def stat_table_kind(items):
    text = " ".join(
        f"{item.get('sheet_name', '')} {item.get('report_table_type', '')}"
        for item in items
    )
    if "锚杆无损" in text:
        return "nondestructive"
    if "锚杆拉拔" in text:
        return "pullout"
    if "回填灌浆" in text:
        return "grout"
    if "钻孔摄像" in text or "钻孔成像" in text:
        return "borehole_imaging"
    if "锚索" in text and "张拉" in text:
        return "anchor_cable_tension"
    if "预应力锚杆" in text or ("锚杆" in text and "张拉" in text):
        return "prestressed_anchor"
    if "桩身完整性" in text:
        return "pile_integrity"
    if "松弛圈" in text:
        return "relaxation_circle"
    if "弹性波" in text:
        return "elastic_wave"
    return "general"


def ratio_text(testing, construction):
    return f"{testing / construction * 100:.2f}" if construction else ""


def preferred_qty(item, *keys):
    for key in keys:
        value = item.get(key, 0)
        if value:
            return value
    return 0


def build_total_row(items, column_count, construction_key="construction_qty", testing_key="testing_qty"):
    construction = sum(item.get(construction_key, 0) for item in items)
    testing = sum(item.get(testing_key, 0) for item in items)
    row = ["合计", format_stat_number(construction), format_stat_number(testing), ratio_text(testing, construction)]
    return row + ["/"] * max(0, column_count - len(row))


def build_stat_table_data(items):
    kind = stat_table_kind(items)
    if kind == "nondestructive":
        headers = ["工程部位", "施工数量(根)", "检测数量(根)", "抽检比例(%)", "I", "II", "III", "IV", "合格率(%)"]
        rows = [
            [
                item["part_name"],
                format_stat_number(item["construction_qty"]),
                format_stat_number(item["testing_qty"]),
                f"{item['ratio_value']:.2f}" if item["ratio_value"] else "",
                format_stat_number(item["class_one_qty"]),
                format_stat_number(item["class_two_qty"]),
                format_stat_number(item["class_three_qty"]),
                format_stat_number(item["class_four_qty"]),
                f"{item['pass_rate_value']:.2f}" if item.get("pass_rate_value") else "100.00",
            ]
            for item in items
        ]
        construction = sum(item["construction_qty"] for item in items)
        testing = sum(item["testing_qty"] for item in items)
        class_one = sum(item["class_one_qty"] for item in items)
        class_two = sum(item["class_two_qty"] for item in items)
        class_three = sum(item["class_three_qty"] for item in items)
        class_four = sum(item["class_four_qty"] for item in items)
        rows.append([
            "合计",
            format_stat_number(construction),
            format_stat_number(testing),
            ratio_text(testing, construction),
            format_stat_number(class_one),
            format_stat_number(class_two),
            format_stat_number(class_three),
            format_stat_number(class_four),
            "100.00" if testing and (class_three + class_four == 0) else "",
        ])
        return kind, headers, rows
    if kind == "pullout":
        headers = ["工程部位", "施工数量(根)", "检测数量(根)", "抽检比例(%)", "锚杆直径(mm)", "设计拉拔力值(KN)", "最小值", "最大值", "合格率(%)"]
        rows = [
            [
                item["part_name"],
                format_stat_number(item["construction_qty"]),
                format_stat_number(item["testing_qty"]),
                f"{item['ratio_value']:.2f}" if item["ratio_value"] else "",
                item["diameter_text"],
                item["design_force_text"],
                plain_stat_min_text(item),
                plain_stat_max_text(item),
                format_stat_number(item["pass_rate_value"]) if item.get("pass_rate_value") else "100",
            ]
            for item in items
        ]
        rows.append(build_total_row(items, len(headers)))
        return kind, headers, rows
    if kind == "grout":
        headers = ["工程部位", "检查孔编号1", "检查孔编号2", "检查桩号1", "检查桩号2", "孔口高程1", "孔口高程2", "设计压力(Mpa）", "实测压力1(Mpa）", "实测压力2(Mpa）", "设计要求", "初始10min注浆量1(L)", "初始10min注浆量2(L)", "检测结果"]
        rows = [
            [
                item["part_name"],
                item.get("grout_hole_1", ""),
                item.get("grout_hole_2", ""),
                item.get("grout_station_1", ""),
                item.get("grout_station_2", ""),
                item.get("grout_elevation_1", ""),
                item.get("grout_elevation_2", ""),
                item.get("grout_pressure", ""),
                item.get("grout_measured_pressure_1", ""),
                item.get("grout_measured_pressure_2", ""),
                item.get("grout_requirement", ""),
                format_decimal_number(item.get("grout_volume_1", ""), 3),
                format_decimal_number(item.get("grout_volume_2", ""), 3),
                item.get("grout_result", "") or item.get("result_text", ""),
            ]
            for item in items
        ]
        return kind, headers, rows
    if kind == "borehole_imaging":
        headers = ["工程部位", "孔号", "检测孔数", "检测里程(m)", "起止深度(m)", "报告数", "报告编号"]
        rows = [
            [
                item["part_name"],
                item.get("hole_no_text", ""),
                format_stat_number(item.get("row_count", 0)),
                format_stat_number(preferred_qty(item, "mileage_qty", "testing_qty")),
                item.get("depth_range_text", ""),
                str(item.get("row_count", 0)),
                item.get("report_no_text", ""),
            ]
            for item in items
        ]
        total_mileage = sum(preferred_qty(item, "mileage_qty", "testing_qty") for item in items)
        rows.append(["合计", "/", format_stat_number(sum(item.get("row_count", 0) for item in items)), format_stat_number(total_mileage), "/", str(sum(item.get("row_count", 0) for item in items)), "/"])
        return kind, headers, rows
    if kind == "anchor_cable_tension":
        headers = ["工程部位", "施工数量(根)", "检测数量(根)", "检测比例(%)", "锚索编号", "规格/型号", "锚索长度(m)", "锚固长度(m)", "自由段长度(m)", "设计锚固力", "检测结果"]
        rows = [
            [
                item["part_name"],
                format_stat_number(item["construction_qty"]),
                format_stat_number(item["testing_qty"]),
                ratio_text(item["testing_qty"], item["construction_qty"]),
                item.get("anchor_no_text", ""),
                item.get("spec_text", ""),
                item.get("anchor_length_text", ""),
                item.get("anchorage_length_text", ""),
                item.get("free_length_text", ""),
                item.get("design_anchor_force_text", "") or item.get("design_force_text", ""),
                item.get("result_text", ""),
            ]
            for item in items
        ]
        rows.append(build_total_row(items, len(headers)))
        return kind, headers, rows
    if kind == "prestressed_anchor":
        headers = ["工程部位", "施工数量(根)", "检测数量(根)", "检测比例(%)", "锚杆直径", "锚杆长度(m)", "锚固长度(m)", "自由段长度(m)", "弹性模量", "设计锚固力", "检测结果"]
        rows = [
            [
                item["part_name"],
                format_stat_number(item["construction_qty"]),
                format_stat_number(item["testing_qty"]),
                ratio_text(item["testing_qty"], item["construction_qty"]),
                item.get("diameter_text", ""),
                item.get("anchor_length_text", ""),
                item.get("anchorage_length_text", ""),
                item.get("free_length_text", ""),
                item.get("elastic_modulus_text", ""),
                item.get("design_anchor_force_text", "") or item.get("design_force_text", ""),
                item.get("result_text", ""),
            ]
            for item in items
        ]
        rows.append(build_total_row(items, len(headers)))
        return kind, headers, rows
    if kind == "pile_integrity":
        headers = ["工程部位", "检测数量(根)", "桩型", "桩长", "桩径", "桩顶高程", "综合评判类别", "检测结果", "报告数"]
        rows = [
            [
                item["part_name"],
                format_stat_number(item["testing_qty"]),
                item.get("pile_type_text", ""),
                item.get("anchor_length_text", ""),
                item.get("pile_diameter_text", ""),
                item.get("pile_top_text", ""),
                item.get("category_text", ""),
                item.get("result_text", ""),
                str(item.get("row_count", 0)),
            ]
            for item in items
        ]
        rows.append(["合计", format_stat_number(sum(item.get("testing_qty", 0) for item in items)), "/", "/", "/", "/", "/", "/", str(sum(item.get("row_count", 0) for item in items))])
        return kind, headers, rows
    if kind == "relaxation_circle":
        headers = ["工程部位", "检测孔数量", "报告数", "检测结果", "报告编号"]
        rows = [
            [
                item["part_name"],
                format_stat_number(preferred_qty(item, "hole_count", "testing_qty", "row_count")),
                str(item.get("row_count", 0)),
                item.get("result_text", ""),
                item.get("report_no_text", ""),
            ]
            for item in items
        ]
        rows.append(["合计", format_stat_number(sum(preferred_qty(item, "hole_count", "testing_qty", "row_count") for item in items)), str(sum(item.get("row_count", 0) for item in items)), "/", "/"])
        return kind, headers, rows
    if kind == "elastic_wave":
        headers = ["工程部位", "组数", "检测类型", "检测结果", "报告数", "报告编号"]
        rows = [
            [
                item["part_name"],
                format_stat_number(preferred_qty(item, "group_count", "testing_qty", "row_count")),
                item.get("detection_type_text", ""),
                item.get("result_text", ""),
                str(item.get("row_count", 0)),
                item.get("report_no_text", ""),
            ]
            for item in items
        ]
        rows.append(["合计", format_stat_number(sum(preferred_qty(item, "group_count", "testing_qty", "row_count") for item in items)), "/", "/", str(sum(item.get("row_count", 0) for item in items)), "/"])
        return kind, headers, rows
    headers = ["委托单位", "标段", "委托类别", "检测项目", "统计表类型", "工程部位/单元工程", "检测组数", "施工数量", "检测数量", "检测比例", "检测结果", "报告编号"]
    rows = [
        [
            item["source_type"],
            item["section_name"],
            item["client_type"],
            item["sheet_name"],
            item["report_table_type"],
            item["part_name"],
            str(item["row_count"]),
            format_stat_number(item["construction_qty"]),
            format_stat_number(item["testing_qty"]),
            ratio_text(item["testing_qty"], item["construction_qty"]),
            item.get("result_text", ""),
            item.get("report_no_text", ""),
        ]
        for item in items
    ]
    rows.append(build_total_row(items, len(headers)))
    return kind, headers, rows


def render_simple_stat_table(headers, rows, table_type="general"):
    header_html = "".join(f"<th>{html.escape(header)}</th>" for header in headers)
    body_rows = []
    for row_index, row in enumerate(rows):
        row_class = " class='stat-total'" if row and row[0] == "合计" else ""
        cells = "".join(f"<td>{stat_text(value)}</td>" for value in row)
        body_rows.append(f"<tr{row_class}>{cells}</tr>")
    return (
        f"<table class='ledger-table {html.escape(table_type)}-table'>"
        f"<thead><tr>{header_html}</tr></thead>"
        f"<tbody>{''.join(body_rows)}</tbody>"
        "</table>"
    )


def build_item_stat_tables(grouped):
    section_map = {}
    section_order = []
    for item in grouped:
        key = (item["section_name"], item["sheet_name"], item["report_table_type"])
        if key not in section_map:
            title = f"标段：{item['section_name'] or '未填写标段'}｜{item['sheet_name']} - {item['report_table_type']}"
            section_map[key] = {"title": title, "items": [], "client_type": item["client_type"]}
            section_order.append(key)
        section_map[key]["items"].append(item)

    client_rank = {"施工": 0, "监理": 1, "业主": 2}
    arranged_keys = []
    paired_keys = {}
    paired_order = []
    for key in section_order:
        section_name, sheet_name, report_table_type = key
        base_type = re.sub(r"\[(施工|监理|业主委托|业主|.+?委托)\]$", "", report_table_type)
        pair_key = (section_name, sheet_name, base_type)
        if pair_key not in paired_keys:
            paired_keys[pair_key] = []
            paired_order.append(pair_key)
        paired_keys[pair_key].append(key)
    for pair_key in paired_order:
        arranged_keys.extend(
            sorted(
                paired_keys[pair_key],
                key=lambda k: (client_rank.get(section_map[k]["client_type"], 99), section_order.index(k)),
            )
        )

    tables = []
    for index, key in enumerate(arranged_keys, 1):
        section = section_map[key]
        items = section["items"]
        table_type, headers, rows = build_stat_table_data(items)
        first_item = items[0] if items else {}
        tables.append({
            "id": f"detail_{index}",
            "group": "各检测项分部位统计",
            "title": section["title"],
            "headers": headers,
            "rows": rows,
            "type": table_type,
            "sheet_name": first_item.get("sheet_name", "未填写检测项目"),
            "section_name": first_item.get("section_name", "未填写标段") or "未填写标段",
            "client_type": first_item.get("client_type", ""),
        })
        continue
        is_nondestructive = any("锚杆无损" in item["sheet_name"] or "锚杆无损" in item["report_table_type"] for item in items)
        is_pullout = any("锚杆拉拔" in item["sheet_name"] or "锚杆拉拔" in item["report_table_type"] for item in items)
        is_grout = any("回填灌浆" in item["sheet_name"] or "回填灌浆" in item["report_table_type"] for item in items)
        if is_nondestructive:
            rows = [
                [
                    item["part_name"],
                    format_stat_number(item["construction_qty"]),
                    format_stat_number(item["testing_qty"]),
                    f"{item['ratio_value']:.2f}" if item["ratio_value"] else "",
                    format_stat_number(item["class_one_qty"]),
                    format_stat_number(item["class_two_qty"]),
                    format_stat_number(item["class_three_qty"]),
                    format_stat_number(item["class_four_qty"]),
                    f"{item['pass_rate_value']:.2f}" if item.get("pass_rate_value") else "100.00",
                ]
                for item in items
            ]
            construction = sum(item["construction_qty"] for item in items)
            testing = sum(item["testing_qty"] for item in items)
            class_one = sum(item["class_one_qty"] for item in items)
            class_two = sum(item["class_two_qty"] for item in items)
            class_three = sum(item["class_three_qty"] for item in items)
            class_four = sum(item["class_four_qty"] for item in items)
            rows.append([
                "合计",
                format_stat_number(construction),
                format_stat_number(testing),
                f"{testing / construction * 100:.2f}" if construction else "",
                format_stat_number(class_one),
                format_stat_number(class_two),
                format_stat_number(class_three),
                format_stat_number(class_four),
                "100.00" if testing and (class_three + class_four == 0) else "",
            ])
            headers = ["工程部位", "施工数量(根)", "检测数量(根)", "抽检比例(%)", "I", "II", "III", "IV", "合格率(%)"]
        elif is_pullout:
            rows = [
                [
                    item["part_name"],
                    format_stat_number(item["construction_qty"]),
                    format_stat_number(item["testing_qty"]),
                    f"{item['ratio_value']:.2f}" if item["ratio_value"] else "",
                    item["diameter_text"],
                    item["design_force_text"],
                    plain_stat_min_text(item),
                    plain_stat_max_text(item),
                    format_stat_number(item["pass_rate_value"]) if item.get("pass_rate_value") else "100",
                ]
                for item in items
            ]
            construction = sum(item["construction_qty"] for item in items)
            testing = sum(item["testing_qty"] for item in items)
            pass_rate_values = [item["pass_rate_value"] for item in items if item.get("pass_rate_value")]
            rows.append([
                "合计",
                format_stat_number(construction),
                format_stat_number(testing),
                f"{testing / construction * 100:.2f}" if construction else "",
                "/",
                "/",
                "/",
                "/",
                format_stat_number(min(pass_rate_values)) if pass_rate_values else ("100" if testing else ""),
            ])
            headers = ["工程部位", "施工数量(根)", "检测数量(根)", "抽检比例(%)", "锚杆直径(mm)", "设计拉拔力值(KN)", "最小值", "最大值", "合格率(%)"]
        elif is_grout:
            rows = [
                [
                    item["part_name"],
                    item.get("grout_hole_1", ""),
                    item.get("grout_hole_2", ""),
                    item.get("grout_station_1", ""),
                    item.get("grout_station_2", ""),
                    item.get("grout_elevation_1", ""),
                    item.get("grout_elevation_2", ""),
                    item.get("grout_pressure", ""),
                    item.get("grout_requirement", ""),
                    format_decimal_number(item.get("grout_volume_1", ""), 3),
                    format_decimal_number(item.get("grout_volume_2", ""), 3),
                    item.get("grout_result", ""),
                ]
                for item in items
            ]
            headers = ["工程部位", "检查孔编号1", "检查孔编号2", "检查桩号1", "检查桩号2", "孔口高程1", "孔口高程2", "设计压力(Mpa）", "设计要求（规定压力下，初始10min内注浆量）(L)", "初始10min内注浆量（L)1", "初始10min内注浆量（L)2", "检测结果"]
        else:
            rows = [
                [
                    item["source_type"],
                    item["section_name"],
                    item["client_type"],
                    item["sheet_name"],
                    item["report_table_type"],
                    item["part_name"],
                    str(item["row_count"]),
                    format_stat_number(item["construction_qty"]),
                    format_stat_number(item["testing_qty"]),
                    f"{item['ratio_value']:.2f}%" if item.get("ratio_value") else "",
                    item["diameter_text"],
                    item["design_force_text"],
                    plain_stat_min_text(item),
                    plain_stat_max_text(item),
                    f"{item['pass_rate_value']:.2f}%" if item.get("pass_rate_value") else "",
                    format_stat_number(item["class_one_qty"]),
                    format_stat_number(item["class_two_qty"]),
                    format_stat_number(item["class_three_qty"]),
                    format_stat_number(item["class_four_qty"]),
                ]
                for item in items
            ]
            headers = ["委托单位", "标段", "委托类别", "检测项目", "第五章表格类型", "工程部位/单元工程", "检测组数", "施工数量", "检测数量", "抽检/检测比例", "锚杆直径", "设计拉拔力", "最小值", "最大值", "合格率", "一类", "二类", "三类", "四类"]
        table_type = "nondestructive" if is_nondestructive else "pullout" if is_pullout else "grout" if is_grout else "general"
        first_item = items[0] if items else {}
        tables.append({
            "id": f"detail_{index}",
            "group": "各检测项分部位统计",
            "title": section["title"],
            "headers": headers,
            "rows": rows,
            "type": table_type,
            "sheet_name": first_item.get("sheet_name", "未填写检测项目"),
            "section_name": first_item.get("section_name", "未填写标段") or "未填写标段",
            "client_type": first_item.get("client_type", ""),
        })
    return tables


def build_section_total_tables(grouped):
    totals = {}
    for item in grouped:
        section_name = item["section_name"] or "未填写标段"
        client_type = item["client_type"] or "未填写委托类别"
        item_name = item["report_table_type"] or item["sheet_name"] or "未填写检测项目"
        key = (section_name, client_type, item_name)
        total = totals.setdefault(
            key,
            {
                "section_name": section_name,
                "client_type": client_type,
                "item_name": item_name,
                "row_count": 0,
                "construction_qty": 0,
                "testing_qty": 0,
                "class_one_qty": 0,
                "class_two_qty": 0,
                "class_three_qty": 0,
                "class_four_qty": 0,
            },
        )
        total["row_count"] += item["row_count"]
        total["construction_qty"] += item["construction_qty"]
        total["testing_qty"] += item["testing_qty"]
        total["class_one_qty"] += item["class_one_qty"]
        total["class_two_qty"] += item["class_two_qty"]
        total["class_three_qty"] += item["class_three_qty"]
        total["class_four_qty"] += item["class_four_qty"]

    headers = ["标段", "检测项目", "检测组数", "施工总数", "检测数量"]
    if not totals:
        return [{"id": "summary_1", "group": "各标段总量汇总统计", "title": "施工委托各标段总量汇总统计", "headers": headers, "rows": [["当前筛选条件下暂无可统计数据。"] + [""] * 4], "type": "summary", "sheet_name": "总量汇总", "section_name": "全部标段", "client_type": "施工"}]

    by_client = {}
    for total in totals.values():
        by_client.setdefault(total["client_type"], []).append(total)

    tables = []
    for index, (client_type, client_totals) in enumerate(by_client.items(), 1):
        rows = []
        for total in client_totals:
            rows.append([
                total["section_name"],
                total["item_name"],
                str(total["row_count"]),
                format_stat_number(total["construction_qty"]),
                format_stat_number(total["testing_qty"]),
            ])
        tables.append({"id": f"summary_{index}", "group": "各标段总量汇总统计", "title": f"{client_type}委托各标段总量汇总统计", "headers": headers, "rows": rows, "type": "summary", "sheet_name": "总量汇总", "section_name": "全部标段", "client_type": client_type})
    return tables


def build_statistics_export_tables(grouped):
    return build_item_stat_tables(grouped) + build_section_total_tables(grouped)


def is_period_report_type(report_type):
    return report_type in {"quarter", "year"}


def period_completion_header(report_type):
    if report_type == "year":
        return "年完成量"
    if report_type == "quarter":
        return "本季度完成量"
    return "本期完成量"


def pass_rate_from_counts(item):
    testing = item.get("testing_qty", 0)
    if item.get("pass_rate_value"):
        return format_stat_number(item["pass_rate_value"])
    if not testing:
        return ""
    if item.get("class_three_qty", 0) + item.get("class_four_qty", 0) == 0:
        return "100"
    return ""


def item_report_kind(item):
    return stat_table_kind([item])


def report_unit_name(item):
    return item.get("unit_name") or item.get("part_name") or "未填写单位工程"


def report_sub_unit_name(item):
    return item.get("sub_unit_name") or item.get("location_name") or item.get("part_name") or ""


def period_table_title(stat_params, name):
    report_label = {"quarter": "季报", "year": "年报"}.get(stat_params.get("report_type"), "本期")
    return f"{report_label}{name}"


def build_period_contract_completion_table(grouped, cumulative_grouped, stat_params):
    period_index = {}
    for item in grouped:
        key = (
            item.get("section_name") or "未填写标段",
            item.get("client_type") or "未填写委托类别",
            item.get("report_table_type") or item.get("sheet_name") or "未填写检测项目",
        )
        total = period_index.setdefault(key, {"testing_qty": 0, "row_count": 0})
        total["testing_qty"] += item.get("testing_qty", 0)
        total["row_count"] += item.get("row_count", 0)

    cumulative_index = {}
    for item in cumulative_grouped:
        key = (
            item.get("section_name") or "未填写标段",
            item.get("client_type") or "未填写委托类别",
            item.get("report_table_type") or item.get("sheet_name") or "未填写检测项目",
        )
        total = cumulative_index.setdefault(key, {"construction_qty": 0, "testing_qty": 0, "row_count": 0})
        total["construction_qty"] += item.get("construction_qty", 0)
        total["testing_qty"] += item.get("testing_qty", 0)
        total["row_count"] += item.get("row_count", 0)

    keys = list(cumulative_index)
    for key in period_index:
        if key not in cumulative_index:
            keys.append(key)

    rows = []
    for index, key in enumerate(keys, 1):
        section_name, client_type, item_name = key
        cumulative = cumulative_index.get(key, {})
        period = period_index.get(key, {})
        construction = cumulative.get("construction_qty", 0)
        cumulative_done = cumulative.get("testing_qty", 0)
        period_done = period.get("testing_qty", 0)
        rows.append([
            str(index),
            section_name,
            client_type,
            item_name,
            "项",
            format_stat_number(construction),
            format_stat_number(cumulative_done),
            format_stat_number(period_done),
            ratio_text(cumulative_done, construction),
        ])

    headers = ["编号", "标段", "委托类别", "项目名称", "单位", "工程总量", "累计完成量", period_completion_header(stat_params["report_type"]), "累计完成比例(%)"]
    return {
        "id": "period_contract_completion",
        "group": "季报/年报综合统计",
        "title": period_table_title(stat_params, "已开展工作完成合同量情况"),
        "headers": headers,
        "rows": rows or [["暂无数据"] + [""] * (len(headers) - 1)],
        "type": "summary",
        "sheet_name": "季报年报综合统计",
        "section_name": "全部标段",
        "client_type": "全部委托",
    }


def build_period_unit_coverage_table(grouped, stat_params):
    totals = {}
    order = []
    for item in grouped:
        key = (
            item.get("section_name") or "未填写标段",
            item.get("client_type") or "未填写委托类别",
            item.get("report_table_type") or item.get("sheet_name") or "未填写检测项目",
            report_unit_name(item),
        )
        if key not in totals:
            totals[key] = {
                "row_count": 0,
                "construction_qty": 0,
                "testing_qty": 0,
                "class_one_qty": 0,
                "class_two_qty": 0,
                "class_three_qty": 0,
                "class_four_qty": 0,
            }
            order.append(key)
        total = totals[key]
        total["row_count"] += item.get("row_count", 0)
        total["construction_qty"] += item.get("construction_qty", 0)
        total["testing_qty"] += item.get("testing_qty", 0)
        total["class_one_qty"] += item.get("class_one_qty", 0)
        total["class_two_qty"] += item.get("class_two_qty", 0)
        total["class_three_qty"] += item.get("class_three_qty", 0)
        total["class_four_qty"] += item.get("class_four_qty", 0)

    rows = []
    for index, key in enumerate(order, 1):
        section_name, client_type, item_name, unit_name = key
        total = totals[key]
        rows.append([
            str(index),
            section_name,
            client_type,
            item_name,
            unit_name,
            str(total["row_count"]),
            format_stat_number(total["construction_qty"]),
            format_stat_number(total["testing_qty"]),
            ratio_text(total["testing_qty"], total["construction_qty"]),
            format_stat_number(total["class_one_qty"]),
            format_stat_number(total["class_two_qty"]),
            format_stat_number(total["class_three_qty"]),
            format_stat_number(total["class_four_qty"]),
        ])

    headers = ["序号", "标段", "委托类别", "检测项目", "单位工程/单元工程", "记录数", "施工数量", "检测数量", "检测覆盖率(%)", "I类", "II类", "III类", "IV类"]
    return {
        "id": "period_unit_coverage",
        "group": "季报/年报综合统计",
        "title": period_table_title(stat_params, "单位工程检测覆盖情况"),
        "headers": headers,
        "rows": rows or [["暂无数据"] + [""] * (len(headers) - 1)],
        "type": "summary",
        "sheet_name": "季报年报综合统计",
        "section_name": "全部标段",
        "client_type": "全部委托",
    }


def period_detail_headers(kind):
    if kind == "nondestructive":
        return ["单位工程", "分部工程", "杆长(m)", "直径(mm)", "代表数量(根)", "检测数量(根)", "检测频率(%)", "I", "II", "III", "IV", "合格率(%)"]
    if kind == "pullout":
        return ["工程部位", "施工数量(根)", "检测数量(根)", "检测比例(%)", "锚杆直径(mm)", "设计拉拔力值(KN)", "最小值", "最大值", "合格率(%)"]
    if kind == "grout":
        return ["工程部位", "检查孔编号1", "检查孔编号2", "检查桩号1", "检查桩号2", "孔口高程1", "孔口高程2", "设计压力(Mpa)", "实测压力1(Mpa)", "实测压力2(Mpa)", "设计要求", "初始10min注浆量1(L)", "初始10min注浆量2(L)", "检测结果"]
    if kind == "borehole_imaging":
        return ["工程单位", "分部工程", "钻孔编号", "检测里程(m)", "检测情况", "报告编号"]
    if kind in {"anchor_cable_tension", "prestressed_anchor"}:
        return ["单元工程", "分部工程", "施工数量", "检测数量", "锚索/锚杆编号", "长度(m)", "检测结果"]
    if kind == "pile_integrity":
        return ["单元工程", "分部工程", "工程部位", "检测根数", "综合评判类别", "检测结果"]
    if kind == "relaxation_circle":
        return ["工程部位", "检测孔数量", "检测结果", "报告编号"]
    if kind == "elastic_wave":
        return ["工程部位", "组数", "检测类型", "检测结果", "报告编号"]
    return ["工程部位", "检测组数", "施工数量", "检测数量", "检测比例(%)", "检测结果", "报告编号"]


def period_detail_row(item, kind):
    if kind == "nondestructive":
        return [
            report_unit_name(item),
            report_sub_unit_name(item),
            item.get("anchor_length_text", "") or item.get("length_text", ""),
            item.get("diameter_text", ""),
            format_stat_number(item.get("construction_qty", 0)),
            format_stat_number(item.get("testing_qty", 0)),
            ratio_text(item.get("testing_qty", 0), item.get("construction_qty", 0)),
            format_stat_number(item.get("class_one_qty", 0)),
            format_stat_number(item.get("class_two_qty", 0)),
            format_stat_number(item.get("class_three_qty", 0)),
            format_stat_number(item.get("class_four_qty", 0)),
            pass_rate_from_counts(item),
        ]
    if kind == "pullout":
        return [
            item.get("part_name", ""),
            format_stat_number(item.get("construction_qty", 0)),
            format_stat_number(item.get("testing_qty", 0)),
            ratio_text(item.get("testing_qty", 0), item.get("construction_qty", 0)),
            item.get("diameter_text", ""),
            item.get("design_force_text", ""),
            plain_stat_min_text(item),
            plain_stat_max_text(item),
            pass_rate_from_counts(item),
        ]
    if kind == "grout":
        return [
            item.get("part_name", ""),
            item.get("grout_hole_1", ""),
            item.get("grout_hole_2", ""),
            item.get("grout_station_1", ""),
            item.get("grout_station_2", ""),
            item.get("grout_elevation_1", ""),
            item.get("grout_elevation_2", ""),
            item.get("grout_pressure", ""),
            item.get("grout_measured_pressure_1", ""),
            item.get("grout_measured_pressure_2", ""),
            item.get("grout_requirement", ""),
            format_decimal_number(item.get("grout_volume_1", ""), 3),
            format_decimal_number(item.get("grout_volume_2", ""), 3),
            item.get("grout_result", "") or item.get("result_text", ""),
        ]
    if kind == "borehole_imaging":
        return [
            report_unit_name(item),
            report_sub_unit_name(item),
            item.get("hole_no_text", ""),
            format_stat_number(preferred_qty(item, "mileage_qty", "testing_qty")),
            item.get("result_text", "") or item.get("remark_text", ""),
            item.get("report_no_text", ""),
        ]
    if kind in {"anchor_cable_tension", "prestressed_anchor"}:
        return [
            report_unit_name(item),
            report_sub_unit_name(item),
            format_stat_number(item.get("construction_qty", 0)),
            format_stat_number(item.get("testing_qty", 0)),
            item.get("anchor_no_text", ""),
            item.get("anchor_length_text", ""),
            item.get("result_text", ""),
        ]
    if kind == "pile_integrity":
        return [
            report_unit_name(item),
            report_sub_unit_name(item),
            item.get("part_name", ""),
            format_stat_number(item.get("testing_qty", 0)),
            item.get("category_text", ""),
            item.get("result_text", ""),
        ]
    if kind == "relaxation_circle":
        return [
            item.get("part_name", ""),
            format_stat_number(preferred_qty(item, "hole_count", "testing_qty", "row_count")),
            item.get("result_text", ""),
            item.get("report_no_text", ""),
        ]
    if kind == "elastic_wave":
        return [
            item.get("part_name", ""),
            format_stat_number(preferred_qty(item, "group_count", "testing_qty", "row_count")),
            item.get("detection_type_text", ""),
            item.get("result_text", ""),
            item.get("report_no_text", ""),
        ]
    return [
        item.get("part_name", ""),
        str(item.get("row_count", 0)),
        format_stat_number(item.get("construction_qty", 0)),
        format_stat_number(item.get("testing_qty", 0)),
        ratio_text(item.get("testing_qty", 0), item.get("construction_qty", 0)),
        item.get("result_text", ""),
        item.get("report_no_text", ""),
    ]


def build_period_detail_tables(grouped, stat_params):
    table_map = {}
    order = []
    for item in grouped:
        kind = item_report_kind(item)
        key = (
            item.get("section_name") or "未填写标段",
            item.get("client_type") or "未填写委托类别",
            item.get("sheet_name") or item.get("report_table_type") or "未填写检测项目",
            item.get("report_table_type") or item.get("sheet_name") or "未填写检测项目",
            kind,
        )
        if key not in table_map:
            section_name, client_type, sheet_name, item_name, _ = key
            table_map[key] = {
                "section_name": section_name,
                "client_type": client_type,
                "sheet_name": sheet_name,
                "item_name": item_name,
                "kind": kind,
                "items": [],
            }
            order.append(key)
        table_map[key]["items"].append(item)

    tables = []
    for index, key in enumerate(order, 1):
        table = table_map[key]
        headers = period_detail_headers(table["kind"])
        rows = [period_detail_row(item, table["kind"]) for item in table["items"]]
        title = period_table_title(stat_params, f"{table['section_name']} {table['item_name']}[{table['client_type']}]")
        tables.append({
            "id": f"period_detail_{index}",
            "group": "季报/年报正文统计表",
            "title": title,
            "headers": headers,
            "rows": rows or [["暂无数据"] + [""] * (len(headers) - 1)],
            "type": "grout" if table["kind"] == "grout" else "summary",
            "sheet_name": table["sheet_name"],
            "section_name": table["section_name"],
            "client_type": table["client_type"],
        })
    return tables


def build_period_report_tables(grouped, stat_params, source_filters=None, sheet_filters=None, unit_filters=None):
    if not is_period_report_type(stat_params.get("report_type")):
        return []
    cumulative_records = collect_stat_records(dt.date(1900, 1, 1), stat_params["end_date"], source_filters, sheet_filters)
    cumulative_records = filter_records_by_units(cumulative_records, unit_filters)
    cumulative_grouped = aggregate_stat_records(cumulative_records)
    return [
        build_period_contract_completion_table(grouped, cumulative_grouped, stat_params),
        build_period_unit_coverage_table(grouped, stat_params),
    ] + build_period_detail_tables(grouped, stat_params)


def render_period_report_sections(tables):
    if not tables:
        return ""
    blocks = []
    for index, table in enumerate(tables, 1):
        table_html = render_simple_stat_table(table.get("headers", []), table.get("rows", []), table.get("type", "summary"))
        open_attr = " open" if index <= 2 else ""
        blocks.append(
            f"""
            <details class="stat-section"{open_attr}>
              <summary>{html.escape(table.get("title", ""))}</summary>
              {table_html}
            </details>
            """
        )
    return f"""
    <div class="panel">
      <h3>季报/年报报表统计</h3>
      <p class="muted">优先显示合同量完成情况和单位工程覆盖情况；各检测项目正文表可按需展开。</p>
      {''.join(blocks)}
    </div>
    """


def render_export_options(tables):
    if not tables:
        return "<p class='muted'>当前筛选条件下暂无可导出的统计表。</p>"

    grouped = {}
    project_order = []
    for table in tables:
        sheet_name = table.get("sheet_name") or "未填写检测项目"
        section_name = table.get("section_name") or "未填写标段"
        if sheet_name not in grouped:
            grouped[sheet_name] = {}
            project_order.append(sheet_name)
        if section_name not in grouped[sheet_name]:
            grouped[sheet_name][section_name] = []
        grouped[sheet_name][section_name].append(table)

    blocks = []
    for sheet_name in project_order:
        section_blocks = []
        for section_name, section_tables in grouped[sheet_name].items():
            options = "".join(
                f"<label class='check-line'><input type='checkbox' name='table_id' value='{html.escape(table['id'])}' checked> "
                f"{html.escape(table['group'])}｜{html.escape(table.get('client_type') or '')}｜{html.escape(table['title'])}</label>"
                for table in section_tables
            )
            section_blocks.append(
                f"""
                <details class="export-section" open>
                  <summary>标段：{html.escape(section_name)}</summary>
                  <div class="export-options">{options}</div>
                </details>
                """
            )
        blocks.append(
            f"""
            <details class="export-project" open>
              <summary>检测项目：{html.escape(sheet_name)}</summary>
              {''.join(section_blocks)}
            </details>
            """
        )
    return "".join(blocks)


def set_doc_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    for line_index, line in enumerate(str(text or "").split("\n")):
        paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")
        run.font.size = Pt(10.5)


def set_doc_table_cell_margins(table, top=60, start=60, bottom=60, end=60):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_doc_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4" if edge in ("insideH", "insideV") else "12")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def fill_doc_table_headers(doc_table, table_type, headers):
    if table_type == "nondestructive":
        for index in (0, 1, 2, 3, 8):
            doc_table.cell(0, index).merge(doc_table.cell(1, index))
        doc_table.cell(0, 4).merge(doc_table.cell(0, 7))
        for index, value in ((0, "工程部位"), (1, "施工数量(根)"), (2, "检测数量(根)"), (3, "抽检比例(%)"), (4, "锚杆分级"), (8, "合格率(%)")):
            set_doc_cell_text(doc_table.cell(0, index), value, bold=True)
        for index, value in ((4, "I"), (5, "II"), (6, "III"), (7, "IV")):
            set_doc_cell_text(doc_table.cell(1, index), value, bold=True)
        return 2
    if table_type == "pullout":
        for index in (0, 1, 2, 3, 4, 5, 8):
            doc_table.cell(0, index).merge(doc_table.cell(1, index))
        doc_table.cell(0, 6).merge(doc_table.cell(0, 7))
        for index, value in ((0, "工程部位"), (1, "施工数量(根)"), (2, "检测数量(根)"), (3, "抽检比例(%)"), (4, "锚杆直径(mm)"), (5, "设计拉拔力值(KN)"), (6, "实测拉拔力值(KN)"), (8, "合格率(%)")):
            set_doc_cell_text(doc_table.cell(0, index), value, bold=True)
        for index, value in ((6, "最小值"), (7, "最大值")):
            set_doc_cell_text(doc_table.cell(1, index), value, bold=True)
        return 2
    if table_type == "grout":
        for index in (0, 7, 10, 13):
            doc_table.cell(0, index).merge(doc_table.cell(1, index))
        for start, end in ((1, 2), (3, 4), (5, 6), (8, 9), (11, 12)):
            doc_table.cell(0, start).merge(doc_table.cell(0, end))
        for index, value in ((0, "工程部位"), (1, "检查孔编号"), (3, "检查桩号"), (5, "孔口高程"), (7, "设计压力\n(Mpa）"), (8, "实测压力\n(Mpa）"), (10, "设计要求（规定压力下，初始10min内注浆量）(L)"), (11, "初始10min内注浆量（L)"), (13, "检测结果")):
            set_doc_cell_text(doc_table.cell(0, index), value, bold=True)
        for index, value in ((1, "1"), (2, "2"), (3, "1"), (4, "2"), (5, "1"), (6, "2"), (8, "1"), (9, "2"), (11, "1"), (12, "2")):
            set_doc_cell_text(doc_table.cell(1, index), value, bold=True)
        return 2
    for col_index, header in enumerate(headers):
        set_doc_cell_text(doc_table.cell(0, col_index), header, bold=True)
    return 1


def build_statistics_docx(tables, stat_params):
    document = Document()
    if any(table.get("type") == "grout" for table in tables):
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")
    normal.font.size = Pt(10.5)

    title = document.add_heading("检测数据统计结果", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period = document.add_paragraph(f"统计时间：{stat_params['start_date'].isoformat()} 至 {stat_params['end_date'].isoformat()}")
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not tables:
        document.add_paragraph("未选择需要导出的统计表。")
    for table_index, table_data in enumerate(tables):
        if table_index:
            document.add_paragraph()
        heading = document.add_heading(table_data["title"], level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rows = table_data["rows"] or [[""] * len(table_data["headers"])]
        table_type = table_data.get("type", "general")
        header_rows = 2 if table_type in ("nondestructive", "pullout", "grout") else 1
        doc_table = document.add_table(rows=len(rows) + header_rows, cols=len(table_data["headers"]))
        doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc_table.style = "Table Grid"
        set_doc_table_borders(doc_table)
        header_rows = fill_doc_table_headers(doc_table, table_type, table_data["headers"])
        for row_index, row in enumerate(rows, header_rows):
            padded = list(row) + [""] * (len(table_data["headers"]) - len(row))
            for col_index, value in enumerate(padded[: len(table_data["headers"])]):
                set_doc_cell_text(doc_table.cell(row_index, col_index), value)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def safe_excel_sheet_title(title, used_titles):
    cleaned = re.sub(r"[\[\]\:\*\?\/\\]", "_", str(title or "统计表")).strip() or "统计表"
    cleaned = cleaned[:31]
    candidate = cleaned
    counter = 1
    while candidate in used_titles:
        suffix = f"_{counter}"
        candidate = f"{cleaned[:31 - len(suffix)]}{suffix}"
        counter += 1
    used_titles.add(candidate)
    return candidate


def build_statistics_xlsx(tables, stat_params):
    workbook = Workbook()
    default_sheet = workbook.active
    workbook.remove(default_sheet)
    used_titles = set()

    summary = workbook.create_sheet(safe_excel_sheet_title("导出说明", used_titles))
    summary.append(["统计时间", f"{stat_params['start_date'].isoformat()} 至 {stat_params['end_date'].isoformat()}"])
    summary.append(["统计表数量", len(tables)])
    summary.append([])
    summary.append(["序号", "分组", "标题", "类型", "行数"])
    for index, table in enumerate(tables, 1):
        summary.append([index, table.get("group", ""), table.get("title", ""), table.get("type", ""), len(table.get("rows") or [])])

    for index, table in enumerate(tables, 1):
        sheet_title = safe_excel_sheet_title(f"{index}_{table.get('sheet_name') or table.get('type')}", used_titles)
        sheet = workbook.create_sheet(sheet_title)
        headers = table.get("headers") or []
        rows = table.get("rows") or []
        sheet.append([table.get("title", "")])
        sheet.append(["统计时间", f"{stat_params['start_date'].isoformat()} 至 {stat_params['end_date'].isoformat()}"])
        sheet.append([])
        sheet.append(headers)
        for row in rows:
            padded = list(row) + [""] * (len(headers) - len(row))
            sheet.append(padded[: len(headers)])
        sheet.freeze_panes = "A5"
        if headers:
            sheet.auto_filter.ref = f"A4:{col_letter(len(headers))}{max(4, len(rows) + 4)}"
        for col_index in range(1, max(1, len(headers)) + 1):
            sheet.column_dimensions[col_letter(col_index)].width = 18

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def render_item_stat_sections(grouped):
    if not grouped:
        return "<div class='panel'><h3>各检测项分部位统计</h3><p>当前筛选条件下暂无可统计数据。</p></div>"
    section_map = {}
    section_order = []
    for item in grouped:
        key = (item["section_name"], item["sheet_name"], item["report_table_type"])
        if key not in section_map:
            title = f"标段：{item['section_name'] or '未填写标段'}｜{item['sheet_name']} - {item['report_table_type']}"
            section_map[key] = {"title": title, "items": [], "client_type": item["client_type"]}
            section_order.append(key)
        section_map[key]["items"].append(item)

    client_rank = {"施工": 0, "监理": 1, "业主": 2}
    arranged_keys = []
    paired_keys = {}
    paired_order = []
    for key in section_order:
        section_name, sheet_name, report_table_type = key
        base_type = re.sub(r"\[(施工|监理|业主委托|业主|.+?委托)\]$", "", report_table_type)
        pair_key = (section_name, sheet_name, base_type)
        if pair_key not in paired_keys:
            paired_keys[pair_key] = []
            paired_order.append(pair_key)
        paired_keys[pair_key].append(key)
    for pair_key in paired_order:
        arranged_keys.extend(
            sorted(
                paired_keys[pair_key],
                key=lambda k: (client_rank.get(section_map[k]["client_type"], 99), section_order.index(k)),
            )
        )
    sections = [(section_map[key]["title"], section_map[key]["items"]) for key in arranged_keys]

    blocks = []
    for index, (title, items) in enumerate(sections, 1):
        if any("锚杆无损" in item["sheet_name"] or "锚杆无损" in item["report_table_type"] for item in items):
            rows = "".join(render_nondestructive_row(item) for item in items)
            rows += render_nondestructive_total_row(items)
            table_html = f"""
              <table class="ledger-table nondestructive-table">
                <thead>{NONDESTRUCTIVE_HEADER}</thead>
                <tbody>{rows}</tbody>
              </table>
            """
        elif any("锚杆拉拔" in item["sheet_name"] or "锚杆拉拔" in item["report_table_type"] for item in items):
            rows = "".join(render_pullout_row(item) for item in items)
            rows += render_pullout_total_row(items)
            table_html = f"""
              <table class="ledger-table pullout-table">
                <thead>{PULLOUT_HEADER}</thead>
                <tbody>{rows}</tbody>
              </table>
            """
        elif any("回填灌浆" in item["sheet_name"] or "回填灌浆" in item["report_table_type"] for item in items):
            rows = "".join(render_grout_row(item) for item in items)
            table_html = f"""
              <table class="ledger-table grout-table">
                <thead>{GROUT_HEADER}</thead>
                <tbody>{rows}</tbody>
              </table>
            """
        else:
            table_type, headers, rows = build_stat_table_data(items)
            table_html = render_simple_stat_table(headers, rows, table_type)
        open_attr = " open" if index <= 3 else ""
        blocks.append(
            f"""
            <details class="stat-section"{open_attr}>
              <summary>{html.escape(title)}</summary>
              {table_html}
            </details>
            """
        )
    return f"""
    <div class="panel">
      <h3>检测项目明细统计</h3>
      <p class="muted">按标段、检测项目和委托类别分组，表内按工程部位列出统计结果。</p>
      {''.join(blocks)}
    </div>
    """


def render_section_total_summary(grouped):
    totals = {}
    for item in grouped:
        section_name = item["section_name"] or "未填写标段"
        client_type = item["client_type"] or "未填写委托类别"
        item_name = item["report_table_type"] or item["sheet_name"] or "未填写检测项目"
        key = (section_name, client_type, item_name)
        total = totals.setdefault(
            key,
            {
                "section_name": section_name,
                "client_type": client_type,
                "item_name": item_name,
                "row_count": 0,
                "construction_qty": 0,
                "testing_qty": 0,
                "class_one_qty": 0,
                "class_two_qty": 0,
                "class_three_qty": 0,
                "class_four_qty": 0,
            },
        )
        total["row_count"] += item["row_count"]
        total["construction_qty"] += item["construction_qty"]
        total["testing_qty"] += item["testing_qty"]
        total["class_one_qty"] += item["class_one_qty"]
        total["class_two_qty"] += item["class_two_qty"]
        total["class_three_qty"] += item["class_three_qty"]
        total["class_four_qty"] += item["class_four_qty"]

    table_header = (
        "<thead><tr><th>标段</th><th>检测项目</th><th>检测组数</th><th>施工总数</th>"
        "<th>检测数量</th></tr></thead>"
    )
    tables = []
    if not totals:
        rows = "<tr><td colspan='5'>当前筛选条件下暂无可统计数据。</td></tr>"
        tables.append(
            f"""
            <h4>施工委托各标段总量汇总统计</h4>
            <table class="ledger-table">{table_header}<tbody>{rows}</tbody></table>
            """
        )
    else:
        by_client = {}
        for total in totals.values():
            by_client.setdefault(total["client_type"], []).append(total)
        for client_type, client_totals in by_client.items():
            rows = []
            for total in client_totals:
                rows.append(
                    "<tr>"
                    f"<td>{html.escape(total['section_name'])}</td>"
                    f"<td>{html.escape(total['item_name'])}</td>"
                    f"<td>{total['row_count']}</td>"
                    f"<td>{format_stat_number(total['construction_qty'])}</td>"
                    f"<td>{format_stat_number(total['testing_qty'])}</td>"
                    "</tr>"
                )
            title = f"{client_type}委托各标段总量汇总统计"
            tables.append(
                f"""
                <h4>{html.escape(title)}</h4>
                <table class="ledger-table">{table_header}<tbody>{''.join(rows)}</tbody></table>
                """
            )
    return f"""
    <div class="panel">
      <h3>检测项目总量汇总</h3>
      <p class="muted">按委托类别汇总各标段、各检测项目的检测组数、施工总数和检测数量。</p>
      {''.join(tables)}
    </div>
    """


def render_unit_project_summary(records, unit_filters=None):
    def unit_ratio_text(testing, construction):
        if not construction:
            return ""
        return f"{(testing / construction) * 100:.2f}"

    def is_nondestructive_item(text):
        return "锚杆无损" in (text or "")

    def pass_rate_text(class_three, class_four, testing):
        if not testing:
            return ""
        if class_three + class_four == 0:
            return "100"
        return ""

    def render_unit_nondestructive_table(client_records):
        part_totals = {}
        part_order = []
        for record in client_records:
            unit_name = (record.get("unit_name") or "").strip()
            if not unit_name or unit_name == "未识别单元工程":
                unit_name = "未填写单位工程"
            part_name = (record.get("part_name") or record.get("location_name") or "").strip() or "未填写分部工程"
            length = (record.get("length_value") or "").strip()
            diameter = (record.get("diameter_value") or "").strip()
            key = (unit_name, part_name, length, diameter)
            if key not in part_totals:
                part_totals[key] = {
                    "unit_name": unit_name,
                    "part_name": part_name,
                    "length": length,
                    "diameter": diameter,
                    "construction_qty": 0,
                    "testing_qty": 0,
                    "class_one_qty": 0,
                    "class_two_qty": 0,
                    "class_three_qty": 0,
                    "class_four_qty": 0,
                }
                part_order.append(key)
            total = part_totals[key]
            total["construction_qty"] += record.get("construction_qty", 0)
            total["testing_qty"] += record.get("testing_qty", 0)
            total["class_one_qty"] += record.get("class_one_qty", 0)
            total["class_two_qty"] += record.get("class_two_qty", 0)
            total["class_three_qty"] += record.get("class_three_qty", 0)
            total["class_four_qty"] += record.get("class_four_qty", 0)

        run_lengths = {}
        run_start = 0
        while run_start < len(part_order):
            unit_name = part_totals[part_order[run_start]]["unit_name"]
            run_end = run_start + 1
            while run_end < len(part_order) and part_totals[part_order[run_end]]["unit_name"] == unit_name:
                run_end += 1
            run_lengths[run_start] = run_end - run_start
            run_start = run_end

        rows = []
        construction = testing = class_one = class_two = class_three = class_four = 0
        for index, key in enumerate(part_order):
            total = part_totals[key]
            unit_name = total["unit_name"]
            construction += total["construction_qty"]
            testing += total["testing_qty"]
            class_one += total["class_one_qty"]
            class_two += total["class_two_qty"]
            class_three += total["class_three_qty"]
            class_four += total["class_four_qty"]
            unit_cell = ""
            if index in run_lengths:
                unit_cell = f"<td rowspan='{run_lengths[index]}'>{html.escape(unit_name)}</td>"
            rows.append(
                "<tr>"
                f"{unit_cell}"
                f"<td>{html.escape(total['part_name'])}</td>"
                f"<td>{html.escape(total['length'])}</td>"
                f"<td>{html.escape(total['diameter'])}</td>"
                f"<td>{format_stat_number(total['construction_qty'])}</td>"
                f"<td>{format_stat_number(total['testing_qty'])}</td>"
                f"<td>{unit_ratio_text(total['testing_qty'], total['construction_qty'])}</td>"
                f"<td>{format_stat_number(total['class_one_qty'])}</td>"
                f"<td>{format_stat_number(total['class_two_qty'])}</td>"
                f"<td>{format_stat_number(total['class_three_qty'])}</td>"
                f"<td>{format_stat_number(total['class_four_qty'])}</td>"
                f"<td>{pass_rate_text(total['class_three_qty'], total['class_four_qty'], total['testing_qty'])}</td>"
                "</tr>"
            )
        rows.append(
            "<tr class='stat-total'>"
            "<td colspan='2'>合计</td>"
            "<td>/</td>"
            "<td>/</td>"
            f"<td>{format_stat_number(construction)}</td>"
            f"<td>{format_stat_number(testing)}</td>"
            f"<td>{unit_ratio_text(testing, construction)}</td>"
            f"<td>{format_stat_number(class_one)}</td>"
            f"<td>{format_stat_number(class_two)}</td>"
            f"<td>{format_stat_number(class_three)}</td>"
            f"<td>{format_stat_number(class_four)}</td>"
            f"<td>{pass_rate_text(class_three, class_four, testing)}</td>"
            "</tr>"
        )
        return f"""
        <table class="ledger-table nondestructive-table unit-nondestructive-table">
          <thead>
            <tr>
              <th rowspan="2">单位工程</th>
              <th rowspan="2">分部工程</th>
              <th colspan="2">设计参数</th>
              <th rowspan="2">代表数量<br>(根)</th>
              <th rowspan="2">检测数量<br>(根)</th>
              <th rowspan="2">检测频率<br>(%)</th>
              <th colspan="4">检测成果(根)</th>
              <th rowspan="2">合格率<br>(%)</th>
            </tr>
            <tr>
              <th>杆长<br>(m)</th>
              <th>直径<br>(mm)</th>
              <th>I</th>
              <th>II</th>
              <th>III</th>
              <th>IV</th>
            </tr>
          </thead>
          <tbody>{''.join(rows)}</tbody>
        </table>
        """

    records = filter_records_by_units(records, unit_filters)
    if not records:
        return """
        <div class="panel">
          <h3>单位工程基础汇总</h3>
          <p>当前筛选条件下暂无可统计数据。</p>
        </div>
        """

    totals = {}
    order = []
    for record in records:
        unit_name = (record.get("unit_name") or "").strip()
        if not unit_name or unit_name == "未识别单元工程":
            unit_name = "未填写单位工程"
        section_name = record.get("section_name") or "未填写标段"
        item_name = record.get("report_table_type") or record.get("sheet_name") or "未填写检测项目"
        client_type = record.get("client_type") or "未填写委托类别"
        key = (section_name, item_name, client_type, unit_name)
        if key not in totals:
            totals[key] = {
                "section_name": section_name,
                "item_name": item_name,
                "client_type": client_type,
                "unit_name": unit_name,
                "row_count": 0,
                "construction_qty": 0,
                "testing_qty": 0,
            }
            order.append(key)
        total = totals[key]
        total["row_count"] += record.get("row_count", 1)
        total["construction_qty"] += record.get("construction_qty", 0)
        total["testing_qty"] += record.get("testing_qty", 0)

    grouped_sections = {}
    for key in order:
        total = totals[key]
        section_name = total["section_name"]
        item_name = total["item_name"]
        client_type = total["client_type"]
        grouped_sections.setdefault(section_name, {}).setdefault(item_name, {}).setdefault(client_type, []).append(total)

    client_rank = {"施工": 0, "监理": 1, "业主": 2}
    section_blocks = []
    for section_index, (section_name, item_map) in enumerate(grouped_sections.items(), 1):
        item_blocks = []
        for item_index, (item_name, client_map) in enumerate(item_map.items(), 1):
            client_blocks = []
            for client_type in sorted(client_map, key=lambda value: client_rank.get(value, 99)):
                client_totals = client_map[client_type]
                if is_nondestructive_item(item_name):
                    client_records = [
                        record
                        for record in records
                        if (record.get("section_name") or "未填写标段") == section_name
                        and (record.get("report_table_type") or record.get("sheet_name") or "未填写检测项目") == item_name
                        and (record.get("client_type") or "未填写委托类别") == client_type
                    ]
                    table_html = render_unit_nondestructive_table(client_records)
                else:
                    rows = []
                    row_count = 0
                    construction = 0
                    testing = 0
                    for index, total in enumerate(client_totals, 1):
                        row_count += total["row_count"]
                        construction += total["construction_qty"]
                        testing += total["testing_qty"]
                        rows.append(
                            "<tr>"
                            f"<td>{index}</td>"
                            f"<td>{html.escape(total['unit_name'])}</td>"
                            f"<td>{total['row_count']}</td>"
                            f"<td>{format_stat_number(total['construction_qty'])}</td>"
                            f"<td>{format_stat_number(total['testing_qty'])}</td>"
                            f"<td>{unit_ratio_text(total['testing_qty'], total['construction_qty'])}</td>"
                            "</tr>"
                        )
                    rows.append(
                        "<tr class='stat-total'>"
                        "<td colspan='2'>合计</td>"
                        f"<td>{row_count}</td>"
                        f"<td>{format_stat_number(construction)}</td>"
                        f"<td>{format_stat_number(testing)}</td>"
                        f"<td>{unit_ratio_text(testing, construction)}</td>"
                        "</tr>"
                    )
                    table_html = f"""
                    <table class="ledger-table">
                      <thead><tr><th>序号</th><th>单位工程</th><th>检测组数</th><th>施工总数</th><th>检测数量</th><th>检测比例(%)</th></tr></thead>
                      <tbody>{''.join(rows)}</tbody>
                    </table>
                    """
                client_blocks.append(
                    f"""
                    <h5>{html.escape(client_type)}检测统计</h5>
                    {table_html}
                    """
                )
            item_blocks.append(
                f"""
                <details class="stat-section"{" open" if section_index == 1 and item_index == 1 else ""}>
                  <summary>检测项目：{html.escape(item_name)}</summary>
                  {''.join(client_blocks)}
                </details>
                """
            )
        section_blocks.append(
            f"""
            <details class="stat-section"{" open" if section_index == 1 else ""}>
              <summary>标段：{html.escape(section_name)}</summary>
              {''.join(item_blocks)}
            </details>
            """
        )

    return f"""
    <div class="panel">
      <h3>单位工程基础汇总</h3>
      <p class="muted">按标段、检测项目和委托类别组织，表内按单位工程汇总。</p>
      {''.join(section_blocks)}
    </div>
    """


OPS = {
    ast.Add: operator.add,
    ast.Sub: operator.sub,
    ast.Mult: operator.mul,
    ast.Div: operator.truediv,
    ast.USub: operator.neg,
}


def cell_ref_to_indexes(ref):
    match = re.fullmatch(r"\$?([A-Z]{1,3})\$?(\d+)", ref.upper())
    if not match:
        return None
    return int(match.group(2)), col_to_index(match.group(1))


def numeric_value(value):
    if isinstance(value, (int, float)):
        return value
    if value is None or value == "":
        return 0
    try:
        return float(str(value).replace(",", ""))
    except ValueError:
        return value


def simple_ast_eval(node):
    if isinstance(node, ast.Expression):
        return simple_ast_eval(node.body)
    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return node.value
    if isinstance(node, ast.UnaryOp) and type(node.op) in OPS:
        return OPS[type(node.op)](simple_ast_eval(node.operand))
    if isinstance(node, ast.BinOp) and type(node.op) in OPS:
        return OPS[type(node.op)](simple_ast_eval(node.left), simple_ast_eval(node.right))
    raise ValueError("unsupported formula")


def evaluate_simple_formula(formula, formula_ws, value_ws):
    expr = formula[1:].strip()
    single_ref = cell_ref_to_indexes(expr)
    if single_ref:
        row, col = single_ref
        return value_ws.cell(row, col).value if value_ws.cell(row, col).value is not None else formula_ws.cell(row, col).value

    def replace_ref(match):
        ref = match.group(0)
        row, col = cell_ref_to_indexes(ref)
        value = value_ws.cell(row, col).value
        if value is None:
            value = formula_ws.cell(row, col).value
        value = numeric_value(value)
        if isinstance(value, (int, float)):
            return str(value)
        raise ValueError("non-numeric reference")

    try:
        numeric_expr = re.sub(r"\$?[A-Z]{1,3}\$?\d+", replace_ref, expr)
        if not re.fullmatch(r"[0-9eE\.\+\-\*/\(\) ]+", numeric_expr):
            return None
        return simple_ast_eval(ast.parse(numeric_expr, mode="eval"))
    except Exception:
        return None


def evaluate_db_formula(formula, values):
    expr = str(formula or "")[1:].strip()

    def values_in_range(start_ref, end_ref):
        start = cell_ref_to_indexes(start_ref)
        end = cell_ref_to_indexes(end_ref)
        if not start or not end:
            raise ValueError("invalid range")
        r1, c1 = start
        r2, c2 = end
        result = []
        for row in range(min(r1, r2), max(r1, r2) + 1):
            for col in range(min(c1, c2), max(c1, c2) + 1):
                value = numeric_value(values.get((row, col), 0))
                if isinstance(value, (int, float)):
                    result.append(value)
        return result

    def replace_functions(match):
        fn = match.group(1).upper()
        items = values_in_range(match.group(2), match.group(3))
        if fn == "SUM":
            return str(sum(items))
        if fn in {"AVERAGE", "AVG"}:
            return str(sum(items) / len(items) if items else 0)
        raise ValueError("unsupported function")

    def replace_ref(match):
        ref = match.group(0)
        indexes = cell_ref_to_indexes(ref)
        if not indexes:
            raise ValueError("invalid reference")
        value = numeric_value(values.get(indexes, 0))
        if isinstance(value, (int, float)):
            return str(value)
        raise ValueError("non-numeric reference")

    try:
        expr = re.sub(
            r"\b(SUM|AVERAGE|AVG)\(\s*(\$?[A-Z]{1,3}\$?\d+)\s*:\s*(\$?[A-Z]{1,3}\$?\d+)\s*\)",
            replace_functions,
            expr,
            flags=re.IGNORECASE,
        )
        numeric_expr = re.sub(r"\$?[A-Z]{1,3}\$?\d+", replace_ref, expr)
        if not re.fullmatch(r"[0-9eE\.\+\-\*/\(\) ]+", numeric_expr):
            return None
        result = simple_ast_eval(ast.parse(numeric_expr, mode="eval"))
        return int(result) if isinstance(result, float) and result.is_integer() else result
    except Exception:
        return None


def color_rgb(color):
    return color.rgb if color and color.type == "rgb" else None


def style_to_json(cell):
    style = {
        "number_format": cell.number_format,
        "font": {
            "name": cell.font.name,
            "size": cell.font.sz,
            "bold": cell.font.bold,
            "italic": cell.font.italic,
            "underline": cell.font.underline,
            "color": color_rgb(cell.font.color),
        },
        "fill": {
            "type": cell.fill.fill_type,
            "fgColor": color_rgb(cell.fill.fgColor),
        },
        "alignment": {
            "horizontal": cell.alignment.horizontal,
            "vertical": cell.alignment.vertical,
            "wrap_text": cell.alignment.wrap_text,
            "text_rotation": cell.alignment.text_rotation,
        },
        "border": {
            "left": cell.border.left.style,
            "right": cell.border.right.style,
            "top": cell.border.top.style,
            "bottom": cell.border.bottom.style,
        },
    }
    return json.dumps(style, ensure_ascii=False)


def import_template_to_db(conn, stored_path, meta):
    wb = load_workbook(stored_path, data_only=False)
    value_wb = load_workbook(stored_path, data_only=True)
    cur = conn.execute(
        """
        INSERT INTO template_workbook (
            source_file, file_name, imported_at, project_name, section_name,
            source_type, discipline, remark
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            str(stored_path),
            stored_path.name,
            now_text(),
            meta["project_name"],
            meta["section_name"],
            meta["source_type"],
            meta["discipline"],
            meta["remark"],
        ),
    )
    workbook_id = cur.lastrowid

    for index, ws in enumerate(wb.worksheets, start=1):
        value_ws = value_wb.worksheets[index - 1]
        cur = conn.execute(
            """
            INSERT INTO template_sheet (
                workbook_id, sheet_index, sheet_name, max_row, max_column,
                freeze_panes, sheet_state
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                workbook_id,
                index,
                ws.title,
                ws.max_row,
                ws.max_column,
                str(ws.freeze_panes) if ws.freeze_panes else None,
                ws.sheet_state,
            ),
        )
        sheet_id = cur.lastrowid

        for row in ws.iter_rows():
            for cell in row:
                value_cell = value_ws[cell.coordinate]
                display_value = value_cell.value
                if cell.value is None and display_value is None:
                    continue
                formula = cell.value if isinstance(cell.value, str) and cell.value.startswith("=") else None
                if formula and display_value is None:
                    display_value = evaluate_simple_formula(formula, ws, value_ws)
                conn.execute(
                    """
                    INSERT INTO template_cell (
                        sheet_id, row_index, col_index, cell_ref, raw_value,
                        formula, data_type, number_format, style_json
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        sheet_id,
                        cell.row,
                        cell.column,
                        cell.coordinate,
                        to_text(display_value if formula else cell.value),
                        formula,
                        cell.data_type,
                        cell.number_format,
                        style_to_json(cell),
                    ),
                )

        for merged_range in ws.merged_cells.ranges:
            conn.execute(
                """
                INSERT INTO template_merge (
                    sheet_id, range_ref, min_row, min_col, max_row, max_col
                )
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (
                    sheet_id,
                    str(merged_range),
                    merged_range.min_row,
                    merged_range.min_col,
                    merged_range.max_row,
                    merged_range.max_col,
                ),
            )

        for row_index, dim in ws.row_dimensions.items():
            if dim.height is None and not dim.hidden and not dim.outlineLevel:
                continue
            conn.execute(
                """
                INSERT INTO template_row_dimension (
                    sheet_id, row_index, height, hidden, outline_level
                )
                VALUES (?, ?, ?, ?, ?)
                """,
                (sheet_id, row_index, dim.height, int(bool(dim.hidden)), dim.outlineLevel or 0),
            )

        for col_letter, dim in ws.column_dimensions.items():
            conn.execute(
                """
                INSERT INTO template_column_dimension (
                    sheet_id, col_index, col_letter, width, hidden, outline_level
                )
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (sheet_id, col_to_index(col_letter), col_letter, dim.width, int(bool(dim.hidden)), dim.outlineLevel or 0),
            )
    return workbook_id


def col_to_index(col_letter):
    index = 0
    for char in col_letter:
        if char.isalpha():
            index = index * 26 + ord(char.upper()) - ord("A") + 1
    return index


def upload_file(temp_path, original_filename, meta):
    conn = connect()
    try:
        with conn:
            file_key = "|".join(
                [
                    meta["project_name"],
                    meta["section_name"],
                    meta["source_type"],
                    meta["discipline"],
                    original_filename,
                ]
            )
            row = conn.execute(
                "SELECT id FROM ledger_file WHERE file_key = ?",
                (file_key,),
            ).fetchone()
            if row:
                ledger_file_id = row["id"]
                version_no = conn.execute(
                    "SELECT COALESCE(MAX(version_no), 0) + 1 FROM ledger_file_version WHERE ledger_file_id = ?",
                    (ledger_file_id,),
                ).fetchone()[0]
            else:
                cur = conn.execute(
                    """
                    INSERT INTO ledger_file (
                        project_name, section_name, source_type, discipline,
                        original_filename, file_key, created_at
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        meta["project_name"],
                        meta["section_name"],
                        meta["source_type"],
                        meta["discipline"],
                        original_filename,
                        file_key,
                        now_text(),
                    ),
                )
                ledger_file_id = cur.lastrowid
                version_no = 1

            stem = safe_name(Path(original_filename).stem)
            suffix = Path(original_filename).suffix or ".xlsx"
            version_dir = (
                STORAGE_DIR
                / safe_name(meta["section_name"])
                / safe_name(meta["source_type"])
                / safe_name(meta["discipline"])
                / stem
                / f"v{version_no:03d}"
            )
            version_dir.mkdir(parents=True, exist_ok=True)
            stored_path = version_dir / f"original{suffix}"
            shutil.copy2(temp_path, stored_path)
            file_hash = file_sha256(stored_path)
            workbook_id = import_template_to_db(conn, stored_path, meta)

            cur = conn.execute(
                """
                INSERT INTO ledger_file_version (
                    ledger_file_id, version_no, stored_path, file_size, file_hash,
                    uploaded_at, remark, workbook_id
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    ledger_file_id,
                    version_no,
                    str(stored_path),
                    stored_path.stat().st_size,
                    file_hash,
                    now_text(),
                    meta["remark"],
                    workbook_id,
                ),
            )
            version_id = cur.lastrowid
            conn.execute(
                "UPDATE ledger_file SET current_version_id = ? WHERE id = ?",
                (version_id, ledger_file_id),
            )
        return ledger_file_id, version_no
    finally:
        conn.close()


def reparse_version(version_id):
    conn = connect()
    try:
        with conn:
            row = conn.execute(
                """
                SELECT v.id, v.stored_path, v.workbook_id, f.project_name, f.section_name,
                       f.source_type, f.discipline, v.remark
                FROM ledger_file_version v
                JOIN ledger_file f ON f.id = v.ledger_file_id
                WHERE v.id = ?
                """,
                (version_id,),
            ).fetchone()
            if not row:
                raise ValueError("未找到该版本")
            stored_path = Path(row["stored_path"])
            if not stored_path.exists():
                raise FileNotFoundError(f"原始文件不存在：{stored_path}")
            meta = {
                "project_name": row["project_name"] or "",
                "section_name": row["section_name"] or "",
                "source_type": row["source_type"] or "",
                "discipline": row["discipline"] or "",
                "remark": row["remark"] or "",
            }
            old_workbook_id = row["workbook_id"]
            new_workbook_id = import_template_to_db(conn, stored_path, meta)
            conn.execute(
                "UPDATE ledger_file_version SET workbook_id = ? WHERE id = ?",
                (new_workbook_id, version_id),
            )
            if old_workbook_id:
                conn.execute("DELETE FROM template_workbook WHERE id = ?", (old_workbook_id,))
            return new_workbook_id
    finally:
        conn.close()


def delete_ledger_file(file_id):
    conn = connect()
    paths = []
    workbook_ids = []
    try:
        with conn:
            versions = conn.execute(
                "SELECT stored_path, workbook_id FROM ledger_file_version WHERE ledger_file_id = ?",
                (file_id,),
            ).fetchall()
            for version in versions:
                stored_path = Path(version["stored_path"])
                if stored_path.exists():
                    paths.append(stored_path)
                if version["workbook_id"]:
                    workbook_ids.append(version["workbook_id"])
            for workbook_id in workbook_ids:
                conn.execute("DELETE FROM template_workbook WHERE id = ?", (workbook_id,))
            conn.execute("DELETE FROM ledger_file WHERE id = ?", (file_id,))
    finally:
        conn.close()

    deleted_at = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    for stored_path in paths:
        version_dir = stored_path.parent
        if not version_dir.exists():
            continue
        destination = DELETED_STORAGE_DIR / f"{deleted_at}_{safe_name(version_dir.parent.name)}_{safe_name(version_dir.name)}"
        counter = 1
        while destination.exists():
            destination = DELETED_STORAGE_DIR / f"{deleted_at}_{safe_name(version_dir.parent.name)}_{safe_name(version_dir.name)}_{counter}"
            counter += 1
        try:
            shutil.move(str(version_dir), str(destination))
        except OSError as exc:
            log_error(f"移动删除归档失败：{version_dir} -> {destination}\n{exc}")


def page_layout(title, body, user=None):
    user_nav = ""
    admin_nav = ""
    if user:
        if is_admin(user):
            admin_nav = '<a href="/users">用户管理</a>'
        user_nav = f"<span style='float:right'>用户：{html.escape(user['username'])}　<a href='/logout'>退出</a></span>"
    return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    :root {{ --ink:#172033; --muted:#647084; --line:#d8e0ea; --brand:#145c72; --brand-dark:#0f4658; --accent:#2a8c7c; --table-font:"Times New Roman","方正仿宋GBK","方正仿宋_GBK","FangSong_GB2312","FangSong",serif; --table-font-size:10.5pt; }}
    body {{ margin: 0; font-family: "Microsoft YaHei", "PingFang SC", Arial, sans-serif; background: linear-gradient(180deg,#edf4f8 0,#f7f9fc 300px,#f4f7fb 100%); color: var(--ink); }}
    header {{ background: linear-gradient(115deg,#0f4658 0%,#145c72 56%,#2a8c7c 100%); color: white; padding: 18px 30px 14px; box-shadow: 0 10px 28px rgba(15,70,88,.18); }}
    .brand-title {{ font-size: 21px; font-weight: 700; letter-spacing: .02em; }}
    .brand-subtitle {{ margin-top: 4px; color: rgba(255,255,255,.78); font-size: 13px; }}
    main {{ max-width: 1220px; margin: 24px auto; padding: 0 22px 42px; }}
    main.wide {{ max-width: none; width: 100%; margin: 0; padding: 18px 20px 26px; box-sizing: border-box; }}
    nav {{ margin-top: 14px; display: flex; align-items: center; gap: 8px; flex-wrap: wrap; }}
    nav a {{ color: white; text-decoration: none; padding: 7px 12px; border: 1px solid rgba(255,255,255,.22); border-radius: 4px; background: rgba(255,255,255,.08); }}
    nav a:hover {{ background: rgba(255,255,255,.16); }}
    nav a.active {{ background: rgba(255,255,255,.24); border-color: rgba(255,255,255,.55); }}
    nav span {{ margin-left: auto; color: rgba(255,255,255,.9); }}
    nav span a {{ color: white; padding: 0; border: 0; background: transparent; }}
    .panel {{ background: rgba(255,255,255,.96); border: 1px solid var(--line); border-radius: 6px; padding: 18px; margin-bottom: 16px; box-shadow: 0 8px 24px rgba(23,32,51,.06); }}
    .panel h2, .panel h3 {{ margin-top: 0; color: #12263a; }}
    .panel h4 {{ margin: 18px 0 8px; color: #143d4d; }}
    label {{ display: block; font-weight: 600; margin: 10px 0 5px; color: #27364a; }}
    input, select, textarea {{ width: 100%; box-sizing: border-box; padding: 9px 10px; border: 1px solid #c9d5e3; border-radius: 4px; background: #fbfdff; color: var(--ink); }}
    input:focus, select:focus, textarea:focus {{ outline: 2px solid rgba(42,140,124,.22); border-color: var(--accent); }}
    .toolbar {{ display: grid; grid-template-columns: 2fr 1.2fr auto; gap: 12px; align-items: end; }}
    .toolbar label {{ margin-top: 0; }}
    button, .button {{ display: inline-block; background: var(--brand); color: white; border: 0; padding: 9px 14px; border-radius: 4px; text-decoration: none; cursor: pointer; box-shadow: 0 4px 10px rgba(20,92,114,.18); }}
    button:hover, .button:hover {{ background: var(--brand-dark); }}
    table {{ width: 100%; border-collapse: collapse; background: white; }}
    th, td {{ border: 1px solid var(--line); padding: 8px; text-align: center; vertical-align: middle; font-family: var(--table-font); font-size: var(--table-font-size); }}
    th {{ background: #edf4f6; color: #213446; }}
    .muted {{ color: var(--muted); }}
    .field-note {{ margin: 5px 0 0; color: var(--muted); font-size: 12px; line-height: 1.5; }}
    .multi-select {{ position: relative; width: 100%; }}
    .multi-trigger {{ width: 100%; display: flex; align-items: center; justify-content: space-between; gap: 8px; padding: 9px 10px; border: 1px solid #c9d5e3; border-radius: 4px; background: #fbfdff; color: var(--ink); box-shadow: none; text-align: left; }}
    .multi-trigger:hover {{ background: #f4f8fb; }}
    .multi-trigger span {{ overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }}
    .multi-trigger b {{ color: var(--muted); font-weight: 700; }}
    .multi-menu {{ display: none; position: absolute; z-index: 30; left: 0; right: 0; top: calc(100% + 4px); max-height: 300px; overflow: auto; background: white; border: 1px solid #c9d5e3; border-radius: 4px; box-shadow: 0 12px 28px rgba(23,32,51,.14); padding: 8px; }}
    .multi-select.open .multi-menu {{ display: block; }}
    .multi-actions {{ display: flex; gap: 8px; padding-bottom: 7px; border-bottom: 1px solid var(--line); margin-bottom: 6px; }}
    .multi-actions button {{ width: auto; padding: 5px 9px; font-size: 12px; box-shadow: none; }}
    .multi-option {{ display: flex; align-items: flex-start; gap: 8px; margin: 0; padding: 7px 6px; font-weight: 400; border-radius: 3px; cursor: pointer; }}
    .multi-option:hover {{ background: #eef5f7; }}
    .multi-option input {{ width: auto; margin-top: 2px; }}
    .multi-option span {{ line-height: 1.45; }}
    .inline-edit-form {{ display: grid; grid-template-columns: minmax(220px, 1fr) auto; gap: 6px; align-items: center; }}
    .inline-edit-form select {{ min-width: 220px; padding: 6px 8px; }}
    .inline-edit-form button {{ padding: 6px 9px; box-shadow: none; }}
    .grid {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 12px; }}
    .kpi-grid {{ display: grid; grid-template-columns: repeat(5, minmax(0, 1fr)); gap: 12px; margin-bottom: 16px; }}
    .kpi {{ background: linear-gradient(180deg,#ffffff,#f3f8fa); border: 1px solid var(--line); border-radius: 6px; padding: 16px; box-shadow: 0 8px 24px rgba(23,32,51,.06); }}
    .kpi span {{ display: block; color: var(--muted); font-size: 13px; margin-bottom: 8px; }}
    .kpi strong {{ display: block; color: #12384a; font-size: 24px; line-height: 1.1; }}
    .sheet-wrap {{ overflow: auto; height: calc(100vh - 300px); min-height: 520px; border: 1px solid #c8d6e2; background: white; box-shadow: inset 0 0 0 1px rgba(255,255,255,.5); }}
    main.wide .sheet-wrap {{ height: calc(100vh - 262px); min-height: 620px; }}
    .sheet-grid {{ --zoom: 1; border-collapse: collapse; width: max-content; min-width: 100%; }}
    .sheet-grid th {{ position: sticky; top: 0; z-index: 2; background: #e8f0f3; text-align: center; vertical-align: middle; font-weight: 600; font-family: var(--table-font); font-size: calc(10.5pt * var(--zoom)); }}
    .sheet-grid th.row-head {{ position: sticky; left: 0; z-index: 3; min-width: calc(48px * var(--zoom)); }}
    .sheet-grid td {{ min-width: calc(90px * var(--zoom)); max-width: calc(260px * var(--zoom)); height: calc(28px * var(--zoom)); white-space: pre-wrap; background: #fff; text-align: center; vertical-align: middle; font-family: var(--table-font); font-size: calc(10.5pt * var(--zoom)); }}
    .sheet-grid td[contenteditable] {{ cursor: text; }}
    .sheet-grid td[contenteditable]:focus {{ outline: 2px solid var(--accent); outline-offset: -2px; background: #f4fffb; }}
    .sheet-grid .row-num {{ position: sticky; left: 0; z-index: 1; background: #f1f6f8; text-align: center; vertical-align: middle; font-weight: 600; font-family: var(--table-font); font-size: calc(10.5pt * var(--zoom)); }}
    .zoom-bar, .edit-bar {{ display: flex; align-items: center; gap: 10px; margin: 0 0 10px; flex-wrap: wrap; }}
    .zoom-bar button, .zoom-bar select {{ width: auto; }}
    .zoom-bar input {{ width: 260px; }}
    .zoom-bar select {{ min-width: 96px; padding: 8px 10px; }}
    .filter-btn {{ margin-left: 6px; padding: 2px 6px; font-size: 12px; box-shadow: none; }}
    .cell-text {{ display: inline; }}
    .filter-popover {{ position: fixed; z-index: 20; width: 280px; max-height: 380px; overflow: hidden; background: white; border: 1px solid #b8c7d6; border-radius: 4px; box-shadow: 0 12px 28px rgba(23,32,51,.18); padding: 10px; }}
    .filter-popover input {{ margin-bottom: 8px; }}
    .date-filter {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 6px; margin-bottom: 8px; }}
    .date-filter select {{ width: 100%; padding: 6px 8px; }}
    .filter-actions {{ display: flex; gap: 8px; margin-bottom: 8px; }}
    .filter-actions button {{ padding: 5px 8px; font-size: 12px; }}
    .filter-values {{ max-height: 220px; overflow: auto; border: 1px solid var(--line); padding: 6px; background: #fbfdff; }}
    .filter-values label {{ display: block; margin: 4px 0; font-weight: 400; }}
    .filter-values input {{ width: auto; margin-right: 6px; }}
    .export-list {{ max-height: 360px; overflow: auto; padding: 10px; border: 1px solid var(--line); background: #fbfdff; }}
    .export-project {{ border: 1px solid #d5e0ea; border-radius: 4px; background: white; margin-bottom: 10px; }}
    .export-project > summary {{ cursor: pointer; padding: 9px 11px; font-weight: 700; color: #143d4d; background: #edf4f6; }}
    .export-section {{ margin: 8px 10px; border: 1px solid #e0e7ef; border-radius: 4px; background: #fff; }}
    .export-section > summary {{ cursor: pointer; padding: 7px 9px; font-weight: 600; color: #27364a; background: #f7fafc; }}
    .export-options {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 8px 14px; padding: 9px; }}
    .check-line {{ display: flex; gap: 8px; align-items: flex-start; margin: 0; font-weight: 400; line-height: 1.45; }}
    .check-line input {{ width: auto; margin-top: 3px; }}
    .tabs a {{ display: inline-block; padding: 6px 10px; margin: 0 6px 8px 0; border: 1px solid #cbd5e1; border-radius: 4px; color: var(--brand); text-decoration: none; background: white; }}
    .tabs a.active {{ background: var(--brand); color: white; border-color: var(--brand); }}
    .ledger-table th, .ledger-table td {{ text-align: center; vertical-align: middle; font-family: var(--table-font); }}
    .ledger-table td {{ line-height: 1.45; }}
    .stat-section {{ border: 1px solid var(--line); border-radius: 6px; margin: 12px 0; background: #fbfdff; overflow: auto; }}
    .stat-section summary {{ cursor: pointer; padding: 11px 14px; font-weight: 700; color: #143d4d; background: #edf4f6; }}
    .stat-section table {{ margin: 0; }}
    .nondestructive-table th, .nondestructive-table td {{ text-align: center; vertical-align: middle; }}
    .nondestructive-table th {{ background: #f4f6f8; font-weight: 700; }}
    .nondestructive-table td:first-child {{ white-space: pre-wrap; min-width: 170px; }}
    .nondestructive-table .stat-total td {{ font-weight: 700; background: #fafafa; }}
    .pullout-table th, .pullout-table td {{ text-align: center; vertical-align: middle; white-space: pre-wrap; }}
    .pullout-table th {{ background: #f4f6f8; font-weight: 700; }}
    .pullout-table td:first-child {{ white-space: pre-wrap; min-width: 190px; }}
    .pullout-table .stat-total td {{ font-weight: 700; background: #fafafa; }}
    .grout-table th, .grout-table td {{ text-align: center; vertical-align: middle; white-space: pre-wrap; }}
    .grout-table th {{ background: #f4f6f8; font-weight: 700; }}
    .grout-table td:first-child {{ min-width: 190px; }}
  </style>
</head>
<body>
  <header>
    <div>
      <div class="brand-title">湖南安化抽水蓄能电站施工期物探检测服务台账管理系统</div>
      <div class="brand-subtitle">Ledger Management System for Geophysical Testing Services during Construction of Hunan Anhua Pumped Storage Power Station</div>
    </div>
    <nav style="margin-top:8px">
      <a href="/" class="{'active' if title == '台账管理' else ''}">台账管理</a>
      <a href="/preview" class="{'active' if title == '台账信息查询' else ''}">原表查看</a>
      <a href="/advanced_query" class="{'active' if title == '高级查询' else ''}">台账查询</a>
      <a href="/quality_checks" class="{'active' if title == '异常数据检查' else ''}">数据质检</a>
      <a href="/statistics" class="{'active' if title == '检测数据统计' else ''}">检测数据统计</a>
      <a href="/unit_statistics" class="{'active' if title == '单位工程统计' else ''}">单位工程统计</a>
      {admin_nav}
      {user_nav}
    </nav>
  </header>
  <main class="{'wide' if title == '台账信息查询' else ''}">{body}</main>
</body>
</html>"""


def login_page(error=""):
    error_html = f"<p style='color:#b91c1c'>{html.escape(error)}</p>" if error else ""
    body = f"""
    <div class="panel" style="max-width:420px;margin:60px auto">
      <h2>系统登录</h2>
      {error_html}
      <form method="post" action="/login">
        <label>账号</label><input name="username" autocomplete="username" required>
        <label>密码</label><input type="password" name="password" autocomplete="current-password" required>
        <p><button type="submit">登录</button></p>
      </form>
      <p class="muted">请使用已启用的系统账号登录；旧的 admin/admin123 账号已停用。</p>
    </div>
    """
    return page_layout("系统登录", body)


class AppHandler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        try:
            log_error("%s - %s" % (self.address_string(), format % args))
        except Exception:
            pass

    def handle_one_request(self):
        try:
            super().handle_one_request()
        except Exception:
            log_error("请求处理异常：\n" + traceback.format_exc())
            try:
                self.send_html(page_layout("系统错误", "<div class='panel'>请求处理失败，详情已写入日志。</div>"), status=500)
            except Exception:
                pass

    def current_token(self):
        cookie = SimpleCookie(self.headers.get("Cookie", ""))
        morsel = cookie.get(SESSION_COOKIE)
        return morsel.value if morsel else ""

    def current_user(self):
        return get_session_user(self.current_token())

    def require_user(self):
        user = self.current_user()
        if not user:
            self.redirect("/login")
            return None
        return user

    def do_GET(self):
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path == "/health":
            self.send_text("ok")
        elif parsed.path == "/login":
            self.send_html(login_page())
        elif parsed.path == "/logout":
            delete_session(self.current_token())
            self.redirect("/login")
        elif parsed.path == "/":
            user = self.require_user()
            if user:
                self.send_html(upload_page(user))
        elif parsed.path == "/files":
            user = self.require_user()
            if user:
                self.send_html(files_page(user))
        elif parsed.path == "/preview":
            user = self.require_user()
            if user:
                params = urllib.parse.parse_qs(parsed.query)
                file_id = int(params["file_id"][0]) if "file_id" in params else None
                sheet_id = int(params["sheet_id"][0]) if "sheet_id" in params else None
                self.send_html(preview_page(user, file_id, sheet_id))
        elif parsed.path == "/advanced_query":
            user = self.require_user()
            if user:
                params = urllib.parse.parse_qs(parsed.query)
                self.send_html(advanced_query_page(user, params))
        elif parsed.path == "/advanced_query_export":
            user = self.require_user()
            if user:
                params = urllib.parse.parse_qs(parsed.query)
                query, records = advanced_query_records(params)
                data = build_advanced_query_xlsx(records)
                filename = f"台账高级查询结果_{dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                self.send_xlsx(data, filename)
        elif parsed.path == "/quality_checks":
            user = self.require_user()
            if user:
                self.send_html(quality_checks_page(user))
        elif parsed.path == "/statistics":
            user = self.require_user()
            if user:
                params = urllib.parse.parse_qs(parsed.query)
                self.send_html(statistics_page(user, params))
        elif parsed.path == "/unit_statistics":
            user = self.require_user()
            if user:
                params = urllib.parse.parse_qs(parsed.query)
                self.send_html(unit_statistics_page(user, params))
        elif parsed.path == "/users":
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_html(page_layout("无权访问", "<div class='panel'>仅管理员可访问用户管理。</div>", user), status=403)
                return
            self.send_html(users_page(user))
        elif parsed.path == "/database":
            self.redirect("/")
        elif parsed.path == "/file":
            user = self.require_user()
            if not user:
                return
            params = urllib.parse.parse_qs(parsed.query)
            self.send_html(file_detail_page(int(params["id"][0]), user))
        elif parsed.path == "/sheet":
            user = self.require_user()
            if not user:
                return
            params = urllib.parse.parse_qs(parsed.query)
            self.send_html(sheet_page(int(params["id"][0]), user))
        elif parsed.path == "/download":
            user = self.require_user()
            if not user:
                return
            params = urllib.parse.parse_qs(parsed.query)
            self.send_download(int(params["version_id"][0]))
        else:
            self.send_error(HTTPStatus.NOT_FOUND)

    def do_POST(self):
        if self.path == "/login":
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            username = form_value(form, "username", "")
            password = form_value(form, "password", "")
            user = authenticate_user(username, password)
            if not user:
                self.send_html(login_page("账号或密码错误"), status=401)
                return
            token, expires = create_session(user["id"])
            self.send_response(303)
            self.send_header("Location", "/")
            self.send_header("Set-Cookie", f"{SESSION_COOKIE}={token}; Path=/; HttpOnly; SameSite=Lax")
            self.end_headers()
            return

        if self.path == "/reparse":
            user = self.require_user()
            if not user:
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            version_id = int(form_value(form, "version_id", "0"))
            file_id = int(form_value(form, "file_id", "0"))
            try:
                reparse_version(version_id)
                self.redirect(f"/file?id={file_id}")
            except Exception as exc:
                log_error(traceback.format_exc())
                self.send_html(page_layout("重新解析失败", f"<div class='panel'>重新解析失败：{html.escape(str(exc))}</div>", user), status=500)
            return

        if self.path == "/delete_file":
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_html(page_layout("无权访问", "<div class='panel'>普通用户只能下载台账，不能删除已导入台账。</div>", user), status=403)
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            file_id = int(form_value(form, "file_id", "0"))
            try:
                delete_ledger_file(file_id)
                self.redirect("/")
            except Exception as exc:
                log_error(traceback.format_exc())
                self.send_html(page_layout("删除失败", f"<div class='panel'>删除失败：{html.escape(str(exc))}</div>", user), status=500)
            return

        if self.path == "/create_user":
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_html(page_layout("无权访问", "<div class='panel'>仅管理员可创建用户。</div>", user), status=403)
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            try:
                username = form_value(form, "username", "")
                display_name = form_value(form, "display_name", "")
                organization = form_value(form, "organization", "")
                role = form_value(form, "role", ROLE_USER)
                password = create_app_user(username, display_name, organization, role, user)
                self.send_html(users_page(user, created_username=username, created_password=password, created_role=role))
            except Exception as exc:
                self.send_html(users_page(user, error=str(exc)), status=400)
            return

        if self.path in ("/deactivate_user", "/set_user_active"):
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_html(page_layout("无权访问", "<div class='panel'>仅管理员可操作账号状态。</div>", user), status=403)
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            try:
                target_user_id = int(form_value(form, "user_id", "0"))
                active = int(form_value(form, "active", "0" if self.path == "/deactivate_user" else "1"))
                set_user_active(target_user_id, active, user)
                self.send_html(users_page(user, message="账号已启用" if active else "账号已停用"))
            except Exception as exc:
                self.send_html(users_page(user, error=str(exc)), status=400)
            return

        if self.path == "/update_user_organization":
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_html(page_layout("无权访问", "<div class='panel'>仅管理员可修改用户单位。</div>", user), status=403)
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            try:
                target_user_id = int(form_value(form, "user_id", "0"))
                organization = form_value(form, "organization", "")
                update_user_organization(target_user_id, organization, user)
                self.send_html(users_page(user, message="用户单位已修改"))
            except Exception as exc:
                self.send_html(users_page(user, error=str(exc)), status=400)
            return

        if self.path == "/update_user_role":
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_html(page_layout("无权访问", "<div class='panel'>仅管理员可修改账号角色。</div>", user), status=403)
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            try:
                target_user_id = int(form_value(form, "user_id", "0"))
                role = form_value(form, "role", ROLE_USER)
                update_user_role(target_user_id, role, user)
                self.send_html(users_page(user, message="账号角色已修改"))
            except Exception as exc:
                self.send_html(users_page(user, error=str(exc)), status=400)
            return

        if self.path == "/reset_user_password":
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_html(page_layout("无权访问", "<div class='panel'>仅管理员可重置密码。</div>", user), status=403)
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            try:
                target_user_id = int(form_value(form, "user_id", "0"))
                username, password = reset_user_password(target_user_id, user)
                self.send_html(users_page(user, reset_username=username, reset_password=password))
            except Exception as exc:
                self.send_html(users_page(user, error=str(exc)), status=400)
            return

        if self.path == "/statistics_export":
            user = self.require_user()
            if not user:
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            params = {
                "report_type": [form_value(form, "report_type", "month")],
                "start_date": [form_value(form, "start_date", "")],
                "end_date": [form_value(form, "end_date", "")],
                "source_type": form.getlist("source_type") if "source_type" in form else [],
                "sheet_name": form.getlist("sheet_name") if "sheet_name" in form else [],
            }
            selected_ids = form.getlist("table_id") if "table_id" in form else []
            try:
                stat_params = parse_statistics_params(params)
                records = collect_stat_records(stat_params["start_date"], stat_params["end_date"], stat_params["source_filters"], stat_params["sheet_filters"])
                grouped = aggregate_stat_records(records)
                all_tables = build_statistics_export_tables(grouped)
                selected = [table for table in all_tables if table["id"] in selected_ids]
                data = build_statistics_docx(selected, stat_params)
                self.send_docx(data, f"检测数据统计结果_{stat_params['start_date'].isoformat()}_{stat_params['end_date'].isoformat()}.docx")
            except Exception as exc:
                log_error(traceback.format_exc())
                self.send_html(page_layout("导出失败", f"<div class='panel'>导出失败：{html.escape(str(exc))}</div>", user), status=500)
            return

        if self.path == "/statistics_export_excel":
            user = self.require_user()
            if not user:
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            params = {
                "report_type": [form_value(form, "report_type", "month")],
                "start_date": [form_value(form, "start_date", "")],
                "end_date": [form_value(form, "end_date", "")],
                "source_type": form.getlist("source_type") if "source_type" in form else [],
                "sheet_name": form.getlist("sheet_name") if "sheet_name" in form else [],
            }
            selected_ids = form.getlist("table_id") if "table_id" in form else []
            try:
                stat_params = parse_statistics_params(params)
                records = collect_stat_records(stat_params["start_date"], stat_params["end_date"], stat_params["source_filters"], stat_params["sheet_filters"])
                grouped = aggregate_stat_records(records)
                all_tables = build_statistics_export_tables(grouped)
                selected = [table for table in all_tables if table["id"] in selected_ids]
                data = build_statistics_xlsx(selected, stat_params)
                self.send_xlsx(data, f"检测数据统计结果_{stat_params['start_date'].isoformat()}_{stat_params['end_date'].isoformat()}.xlsx")
            except Exception as exc:
                log_error(traceback.format_exc())
                self.send_html(page_layout("导出失败", f"<div class='panel'>导出失败：{html.escape(str(exc))}</div>", user), status=500)
            return

        if self.path in ("/unit_statistics_export", "/unit_statistics_export_excel"):
            user = self.require_user()
            if not user:
                return
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            params = {
                "report_type": [form_value(form, "report_type", "quarter")],
                "start_date": [form_value(form, "start_date", "")],
                "end_date": [form_value(form, "end_date", "")],
                "source_type": form.getlist("source_type") if "source_type" in form else [],
                "sheet_name": form.getlist("sheet_name") if "sheet_name" in form else [],
                "unit_name": form.getlist("unit_name") if "unit_name" in form else [],
            }
            selected_ids = form.getlist("table_id") if "table_id" in form else []
            try:
                stat_params = parse_statistics_params(params)
                records = collect_stat_records(stat_params["start_date"], stat_params["end_date"], stat_params["source_filters"], stat_params["sheet_filters"])
                filtered_records = filter_records_by_units(records, stat_params["unit_filters"])
                grouped = aggregate_stat_records(filtered_records)
                all_tables = build_period_report_tables(grouped, stat_params, stat_params["source_filters"], stat_params["sheet_filters"], stat_params["unit_filters"])
                selected = [table for table in all_tables if table["id"] in selected_ids]
                filename = f"单位工程季报年报统计_{stat_params['start_date'].isoformat()}_{stat_params['end_date'].isoformat()}"
                if self.path == "/unit_statistics_export_excel":
                    self.send_xlsx(build_statistics_xlsx(selected, stat_params), f"{filename}.xlsx")
                else:
                    self.send_docx(build_statistics_docx(selected, stat_params), f"{filename}.docx")
            except Exception as exc:
                log_error(traceback.format_exc())
                self.send_html(page_layout("导出失败", f"<div class='panel'>导出失败：{html.escape(str(exc))}</div>", user), status=500)
            return

        if self.path == "/save_sheet_cells":
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_text("仅管理员可修改台账信息查询。", status=403)
                return
            content_length = int(self.headers.get("Content-Length", "0") or 0)
            try:
                payload = json.loads(self.rfile.read(content_length).decode("utf-8") or "{}")
                sheet_id = int(payload.get("sheet_id") or 0)
                saved = save_sheet_changes(sheet_id, payload.get("changes", []))
                self.send_text(json.dumps({"ok": True, "saved": saved}, ensure_ascii=False))
            except Exception as exc:
                log_error(traceback.format_exc())
                self.send_text(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False), status=400)
            return

        if self.path == "/delete_sheet_axis":
            user = self.require_user()
            if not user:
                return
            if not is_admin(user):
                self.send_text(json.dumps({"ok": False, "error": "仅管理员可删除行列"}, ensure_ascii=False), status=403)
                return
            content_length = int(self.headers.get("Content-Length", "0") or 0)
            try:
                payload = json.loads(self.rfile.read(content_length).decode("utf-8") or "{}")
                delete_sheet_axis(int(payload.get("sheet_id") or 0), payload.get("axis"), int(payload.get("index") or 0))
                self.send_text(json.dumps({"ok": True}, ensure_ascii=False))
            except Exception as exc:
                log_error(traceback.format_exc())
                self.send_text(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False), status=400)
            return

        if self.path != "/upload":
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        user = self.require_user()
        if not user:
            return
        form = cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ={
                "REQUEST_METHOD": "POST",
                "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
            },
        )
        file_item = form["file"] if "file" in form else None
        if file_item is None or not file_item.filename:
            self.send_html(page_layout("上传失败", "<div class='panel'>未选择文件。</div>", user), status=400)
            return

        temp_dir = DATA_DIR / "tmp"
        temp_dir.mkdir(parents=True, exist_ok=True)
        temp_path = temp_dir / safe_name(file_item.filename)
        with open(temp_path, "wb") as f:
            shutil.copyfileobj(file_item.file, f)

        meta = {
            "project_name": form_value(form, "project_name", PROJECT_OPTIONS[0]),
            "section_name": form_value(form, "section_name", "Q2标"),
            "source_type": form_value(form, "source_type", SOURCE_TYPE_OPTIONS[0]),
            "discipline": form_value(form, "discipline", DETECTION_UNIT_NAME),
            "remark": form_value(form, "remark", ""),
        }
        try:
            ledger_file_id, version_no = upload_file(temp_path, file_item.filename, meta)
            self.redirect(f"/file?id={ledger_file_id}&uploaded=v{version_no:03d}")
        except Exception as exc:
            log_error(traceback.format_exc())
            self.send_html(page_layout("上传失败", f"<div class='panel'>解析或入库失败：{html.escape(str(exc))}</div>", user), status=500)
        finally:
            try:
                temp_path.unlink()
            except OSError:
                pass

    def send_html(self, content, status=200):
        data = content.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_text(self, content, status=200):
        data = content.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "text/plain; charset=utf-8")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def redirect(self, location):
        self.send_response(303)
        self.send_header("Location", location)
        self.end_headers()

    def send_docx(self, data, filename):
        encoded = urllib.parse.quote(filename)
        self.send_response(200)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{encoded}")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_xlsx(self, data, filename):
        encoded = urllib.parse.quote(filename)
        self.send_response(200)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{encoded}")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_download(self, version_id):
        conn = connect()
        try:
            row = conn.execute(
                """
                SELECT f.original_filename, v.stored_path
                FROM ledger_file_version v
                JOIN ledger_file f ON f.id = v.ledger_file_id
                WHERE v.id = ?
                """,
                (version_id,),
            ).fetchone()
        finally:
            conn.close()
        if not row or not Path(row["stored_path"]).exists():
            self.send_error(HTTPStatus.NOT_FOUND)
            return
        path = Path(row["stored_path"])
        data = path.read_bytes()
        filename = urllib.parse.quote(row["original_filename"])
        self.send_response(200)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{filename}")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)


def form_value(form, key, default=""):
    return form[key].value.strip() if key in form and form[key].value is not None else default


def form_hidden_inputs(name, values):
    return "".join(f"<input type='hidden' name='{html.escape(name)}' value='{html.escape(value)}'>" for value in values)


def sql_quote(value):
    return "'" + str(value).replace("'", "''") + "'"


def option_tags(options, selected=""):
    selected_values = set(selected if isinstance(selected, (list, tuple, set)) else [selected])
    return "".join(
        f"<option value='{html.escape(option)}' {'selected' if option in selected_values else ''}>{html.escape(option)}</option>"
        for option in options
    )


def checkbox_dropdown(name, options, selected_values, placeholder):
    selected_set = set(selected_values or [])
    selected_count = len(selected_set)
    summary = f"已选择 {selected_count} 项" if selected_count else placeholder
    items = []
    for value, label in options:
        checked = "checked" if value in selected_set else ""
        items.append(
            "<label class='multi-option'>"
            f"<input type='checkbox' name='{html.escape(name)}' value='{html.escape(value)}' {checked}>"
            f"<span>{html.escape(label)}</span>"
            "</label>"
        )
    return f"""
    <div class="multi-select" data-placeholder="{html.escape(placeholder)}">
      <button type="button" class="multi-trigger"><span>{html.escape(summary)}</span><b>▾</b></button>
      <div class="multi-menu">
        <div class="multi-actions">
          <button type="button" data-action="all">全选</button>
          <button type="button" data-action="none">清空</button>
        </div>
        <div class="multi-options">{''.join(items)}</div>
      </div>
    </div>
    """


def ledger_rows_html(rows, user):
    if not rows:
        return "<tr><td colspan='8'>暂无已导入台账</td></tr>"
    display_ids = {row["id"]: index for index, row in enumerate(rows, start=1)}
    html_rows = []
    current_group = None
    for row in rows:
        group = row["source_type"] or "未填写委托单位"
        detection_unit = row["discipline"] if row["discipline"] and row["discipline"] != "物探" else DETECTION_UNIT_NAME
        if group != current_group:
            current_group = group
            html_rows.append(
                f"<tr><th colspan='8' style='background:#dfecef;color:#143d4d;text-align:left'>委托单位：{html.escape(group)}</th></tr>"
            )
        download = (
            f"<a class='button' href='/download?version_id={row['version_id']}'>下载</a>"
            if row["version_id"]
            else ""
        )
        delete_action = ""
        if is_admin(user):
            delete_action = (
                f"<form method='post' action='/delete_file' style='display:inline' onsubmit=\"return confirm('确认删除该台账及其所有版本？原文件将移入 deleted_storage 备份目录。');\">"
                f"<input type='hidden' name='file_id' value='{row['id']}'>"
                f"<button type='submit' style='background:#9f2d2d'>删除</button>"
                f"</form>"
            )
        html_rows.append(
            f"<tr>"
            f"<td>{display_ids.get(row['id'], '')}</td>"
            f"<td><a href='/file?id={row['id']}'>{html.escape(row['original_filename'])}</a></td>"
            f"<td>{html.escape(row['project_name'] or '')}</td>"
            f"<td>{html.escape(row['section_name'] or '')}</td>"
            f"<td>{html.escape(detection_unit)}</td>"
            f"<td>v{(row['version_no'] or 0):03d}</td>"
            f"<td>{html.escape(row['uploaded_at'] or '')}</td>"
            f"<td>{download} {delete_action}</td>"
            f"</tr>"
        )
    return "".join(html_rows)


def upload_page(user):
    source_order_sql = "CASE " + " ".join(
        f"WHEN f.source_type = {sql_quote(value)} THEN {index}"
        for index, value in enumerate(SOURCE_TYPE_OPTIONS)
    ) + " ELSE 999 END"
    conn = connect()
    try:
        rows = conn.execute(
            f"""
            SELECT f.*, v.version_no, v.uploaded_at, v.id AS version_id
            FROM ledger_file f
            LEFT JOIN ledger_file_version v ON v.id = f.current_version_id
            ORDER BY {source_order_sql}, f.source_type ASC, f.id ASC
            """
        ).fetchall()
    finally:
        conn.close()

    project_options = option_tags(PROJECT_OPTIONS, PROJECT_OPTIONS[0])
    section_options = option_tags(SECTION_OPTIONS, PROJECT_SECTION_MAP[PROJECT_OPTIONS[0]])
    source_type_options = option_tags(SOURCE_TYPE_OPTIONS, PROJECT_CLIENT_MAP[PROJECT_OPTIONS[0]])
    project_section_json = json.dumps(PROJECT_SECTION_MAP, ensure_ascii=False)
    project_client_json = json.dumps(PROJECT_CLIENT_MAP, ensure_ascii=False)
    body = f"""
    <div class="panel">
      <h2>台账管理</h2>
      <form method="post" action="/upload" enctype="multipart/form-data">
        <div class="grid">
          <div><label>项目名称</label><select id="project_name" name="project_name">{project_options}</select></div>
          <div><label>标段</label><select id="section_name" name="section_name">{section_options}</select></div>
          <div><label>委托单位</label><select id="source_type" name="source_type">{source_type_options}</select></div>
          <div><label>检测单位</label><input name="discipline" value="{html.escape(DETECTION_UNIT_NAME)}" readonly></div>
        </div>
        <label>Excel台账文件</label><input type="file" name="file" accept=".xlsx,.xlsm" required>
        <label>上传说明</label><textarea name="remark" rows="3"></textarea>
        <p><button type="submit">上传并原样入库</button></p>
      </form>
    </div>
    <div class="panel muted">系统会保存原始文件，并将工作表、单元格、合并区域、行高列宽和基础样式写入数据库。再次上传同名同标段文件会生成新版本。</div>
    <div class="panel">
      <h3>已导入台账</h3>
      <table class="ledger-table">
        <thead>
          <tr><th>ID</th><th>文件名</th><th>项目</th><th>标段</th><th>检测单位</th><th>当前版本</th><th>更新时间</th><th>操作</th></tr>
        </thead>
        <tbody>
          {ledger_rows_html(rows, user)}
        </tbody>
      </table>
    </div>
    <script>
      const projectSectionMap = {project_section_json};
      const projectClientMap = {project_client_json};
      const projectSelect = document.getElementById("project_name");
      const sectionSelect = document.getElementById("section_name");
      const clientSelect = document.getElementById("source_type");
      function selectValue(select, target) {{
        if (!target) return;
        let matched = false;
        for (const option of select.options) {{
          option.selected = option.value === target;
          if (option.value === target) matched = true;
        }}
        return matched;
      }}
      function syncByProject() {{
        const targetSection = projectSectionMap[projectSelect.value];
        const targetClient = projectClientMap[projectSelect.value];
        selectValue(sectionSelect, targetSection);
        selectValue(clientSelect, targetClient);
      }}
      projectSelect.addEventListener("change", syncByProject);
      syncByProject();
    </script>
    """
    return page_layout("台账管理", body, user)


def users_page(user, created_username="", created_password="", created_role="", reset_username="", reset_password="", error="", message=""):
    conn = connect()
    try:
        rows = conn.execute(
            """
            SELECT u.id, u.username, u.display_name, u.organization, u.role, u.is_active, u.created_at,
                   c.username AS created_by_name
            FROM app_user u
            LEFT JOIN app_user c ON c.id = u.created_by
            ORDER BY u.id ASC
            """
        ).fetchall()
    finally:
        conn.close()

    result_html = ""
    if message:
        result_html = f"<div class='panel' style='border-color:#7ab39f;background:#f2fbf7'>{html.escape(message)}</div>"
    if created_username and created_password:
        created_role_name = {ROLE_ADMIN: "管理员", ROLE_USER: "普通用户"}.get(created_role, "用户")
        result_html = f"""
        <div class="panel" style="border-color:#7ab39f;background:#f2fbf7">
          <h3>{html.escape(created_role_name)}已生成</h3>
          <p>账号：<strong>{html.escape(created_username)}</strong></p>
          <p>初始密码：<strong>{html.escape(created_password)}</strong></p>
          <p class="muted">请记录该密码并交给使用人。系统不会再次明文显示该密码。</p>
        </div>
        """
    if reset_username and reset_password:
        result_html = f"""
        <div class="panel" style="border-color:#7ab39f;background:#f2fbf7">
          <h3>密码已重置</h3>
          <p>账号：<strong>{html.escape(reset_username)}</strong></p>
          <p>新密码：<strong>{html.escape(reset_password)}</strong></p>
          <p class="muted">请立即记录该密码并交给使用人。系统不会长期保存或再次明文显示该密码。</p>
        </div>
        """
    if error:
        result_html = f"<div class='panel' style='border-color:#d08b8b;background:#fff7f7'>创建失败：{html.escape(error)}</div>"

    organization_options = option_tags(USER_ORGANIZATION_OPTIONS, USER_ORGANIZATION_OPTIONS[0])
    create_role_options = option_tags([ROLE_USER, ROLE_ADMIN], ROLE_USER).replace(f">{ROLE_USER}<", ">普通用户<").replace(f">{ROLE_ADMIN}<", ">管理员<")
    row_html = []
    for row in rows:
        status = "启用" if row["is_active"] else "停用"
        row_organization_options = option_tags(USER_ORGANIZATION_OPTIONS, row["organization"] or USER_ORGANIZATION_OPTIONS[0])
        organization_form = (
            "<form method='post' action='/update_user_organization' class='inline-edit-form'>"
            f"<input type='hidden' name='user_id' value='{row['id']}'>"
            f"<select name='organization'>{row_organization_options}</select>"
            "<button type='submit'>保存</button>"
            "</form>"
        )
        row_role_options = option_tags([ROLE_USER, ROLE_ADMIN], row["role"]).replace(f">{ROLE_USER}<", ">普通用户<").replace(f">{ROLE_ADMIN}<", ">管理员<")
        role_form = (
            "<form method='post' action='/update_user_role' class='inline-edit-form'>"
            f"<input type='hidden' name='user_id' value='{row['id']}'>"
            f"<select name='role'>{row_role_options}</select>"
            "<button type='submit'>保存</button>"
            "</form>"
        )
        action = "-"
        if row["id"] != user["id"]:
            target_active = 0 if row["is_active"] else 1
            button_text = "停用" if row["is_active"] else "启用"
            button_style = "background:#9f2d2d" if row["is_active"] else "background:#2a8c7c"
            confirm_text = "确认停用该账号？停用后该账号将不能登录。" if row["is_active"] else "确认启用该账号？启用后该账号可继续登录。"
            action = (
                "<form method='post' action='/set_user_active' style='display:inline' "
                f"onsubmit=\"return confirm('{confirm_text}');\">"
                f"<input type='hidden' name='user_id' value='{row['id']}'>"
                f"<input type='hidden' name='active' value='{target_active}'>"
                f"<button type='submit' style='{button_style}'>{button_text}</button>"
                "</form>"
                "<form method='post' action='/reset_user_password' style='display:inline;margin-left:6px' "
                "onsubmit=\"return confirm('确认重置该账号密码？新密码只会显示一次。');\">"
                f"<input type='hidden' name='user_id' value='{row['id']}'>"
                "<button type='submit'>重置密码</button>"
                "</form>"
            )
        else:
            action = (
                "<form method='post' action='/reset_user_password' style='display:inline' "
                "onsubmit=\"return confirm('确认重置当前账号密码？重置后需要使用新密码登录。');\">"
                f"<input type='hidden' name='user_id' value='{row['id']}'>"
                "<button type='submit'>重置密码</button>"
                "</form>"
            )
        row_html.append(
            "<tr>"
            f"<td>{row['id']}</td>"
            f"<td>{html.escape(row['username'] or '')}</td>"
            f"<td>{html.escape(row['display_name'] or '')}</td>"
            f"<td>{organization_form}</td>"
            f"<td>{role_form}</td>"
            f"<td>{status}</td>"
            f"<td>{html.escape(row['created_by_name'] or '-')}</td>"
            f"<td>{html.escape(row['created_at'] or '')}</td>"
            f"<td>{action}</td>"
            "</tr>"
        )

    body = f"""
    {result_html}
    <div class="panel">
      <h2>用户管理</h2>
      <form method="post" action="/create_user">
        <div class="grid">
          <div>
            <label>登录账号</label>
            <input name="username" placeholder="例如：user01" required>
          </div>
          <div>
            <label>用户名称</label>
            <input name="display_name" placeholder="例如：资料员张三">
          </div>
          <div>
            <label>用户单位</label>
            <select name="organization" required>{organization_options}</select>
          </div>
          <div>
            <label>账号类型</label>
            <select name="role" required>{create_role_options}</select>
          </div>
        </div>
        <p><button type="submit">生成账号及密码</button></p>
      </form>
    </div>
    <div class="panel">
      <h3>账号列表</h3>
      <table class="ledger-table">
        <thead><tr><th>ID</th><th>账号</th><th>用户名称</th><th>用户单位</th><th>角色</th><th>状态</th><th>创建人</th><th>创建时间</th><th>操作</th></tr></thead>
        <tbody>{''.join(row_html) or "<tr><td colspan='9'>暂无用户</td></tr>"}</tbody>
      </table>
    </div>
    """
    return page_layout("用户管理", body, user)


def files_page(user):
    body = """
    <div class="panel">
      <h2>台账列表已合并</h2>
      <p>台账列表和删除功能已合并到“台账管理”页面。</p>
      <p><a class="button" href="/">进入台账管理</a></p>
    </div>
    """
    return page_layout("台账列表", body, user)


def statistics_page_script():
    return f"""
    <script>
      const periodDefaults = {{
        week: {json.dumps([d.isoformat() for d in default_period("week")])},
        month: {json.dumps([d.isoformat() for d in default_period("month")])},
        quarter: {json.dumps([d.isoformat() for d in default_period("quarter")])},
        year: {json.dumps([d.isoformat() for d in default_period("year")])}
      }};
      const reportSelect = document.getElementById("report_type");
      const startInput = document.getElementById("start_date");
      const endInput = document.getElementById("end_date");
      if (reportSelect && startInput && endInput) {{
        reportSelect.addEventListener("change", () => {{
          const value = reportSelect.value;
          if (periodDefaults[value]) {{
            startInput.value = periodDefaults[value][0];
            endInput.value = periodDefaults[value][1];
          }}
        }});
      }}
      document.querySelectorAll(".multi-select").forEach(select => {{
        const trigger = select.querySelector(".multi-trigger");
        const summary = trigger.querySelector("span");
        const placeholder = select.dataset.placeholder || "全部";
        const boxes = Array.from(select.querySelectorAll("input[type='checkbox']"));
        function updateSummary() {{
          const checked = boxes.filter(box => box.checked);
          if (!checked.length) {{
            summary.textContent = placeholder;
          }} else if (checked.length === 1) {{
            const text = checked[0].closest(".multi-option").querySelector("span").textContent.trim();
            summary.textContent = text;
          }} else {{
            summary.textContent = `已选择 ${{checked.length}} 项`;
          }}
        }}
        trigger.addEventListener("click", event => {{
          event.stopPropagation();
          document.querySelectorAll(".multi-select.open").forEach(other => {{
            if (other !== select) other.classList.remove("open");
          }});
          select.classList.toggle("open");
        }});
        select.querySelector(".multi-menu").addEventListener("click", event => event.stopPropagation());
        select.querySelectorAll("[data-action]").forEach(button => {{
          button.addEventListener("click", event => {{
            event.preventDefault();
            const action = button.dataset.action;
            boxes.forEach(box => box.checked = action === "all");
            updateSummary();
          }});
        }});
        boxes.forEach(box => box.addEventListener("change", updateSummary));
        updateSummary();
      }});
      document.addEventListener("click", () => {{
        document.querySelectorAll(".multi-select.open").forEach(select => select.classList.remove("open"));
      }});
    </script>
    """


def statistics_page(user, params):
    stat_params = parse_statistics_params(params)
    report_type = stat_params["report_type"]
    start_date = stat_params["start_date"]
    end_date = stat_params["end_date"]
    source_filters = stat_params["source_filters"]
    sheet_filters = stat_params["sheet_filters"]

    records = collect_stat_records(start_date, end_date, source_filters, sheet_filters)
    grouped = aggregate_stat_records(records)
    total_rows = sum(item["row_count"] for item in grouped)
    total_construction = sum(item["construction_qty"] for item in grouped)
    total_testing = sum(item["testing_qty"] for item in grouped)
    source_count = len({item["source_type"] for item in grouped})
    unit_count = len({item["part_name"] for item in grouped})
    class_one_total = sum(item["class_one_qty"] for item in grouped)
    class_two_total = sum(item["class_two_qty"] for item in grouped)
    class_three_total = sum(item["class_three_qty"] for item in grouped)
    class_four_total = sum(item["class_four_qty"] for item in grouped)

    report_options = [
        ("week", "周报"),
        ("month", "月报"),
        ("quarter", "季报"),
        ("year", "年报"),
        ("custom", "自定义时间段"),
    ]
    report_option_html = "".join(
        f"<option value='{value}' {'selected' if value == report_type else ''}>{label}</option>"
        for value, label in report_options
    )
    source_selector_html = checkbox_dropdown(
        "source_type",
        [(option, option) for option in SOURCE_TYPE_OPTIONS],
        source_filters,
        "全部委托单位",
    )
    sheet_selector_html = checkbox_dropdown(
        "sheet_name",
        statistics_sheet_options(),
        sheet_filters,
        "全部检测项目",
    )

    item_sections_html = render_item_stat_sections(grouped)
    section_totals_html = render_section_total_summary(grouped)
    export_tables = build_statistics_export_tables(grouped)
    export_options = render_export_options(export_tables)
    export_source_inputs = form_hidden_inputs("source_type", source_filters)
    export_sheet_inputs = form_hidden_inputs("sheet_name", sheet_filters)

    body = f"""
    <div class="panel">
      <h2>检测数据统计</h2>
      <form method="get" action="/statistics">
        <div class="grid">
          <div>
            <label>报表类型</label>
            <select id="report_type" name="report_type">{report_option_html}</select>
          </div>
          <div>
            <label>委托单位</label>
            {source_selector_html}
            <p class="field-note">点击下拉后勾选多项；不勾选表示全部。</p>
          </div>
          <div>
            <label>开始日期</label>
            <input id="start_date" type="date" name="start_date" value="{start_date.isoformat()}">
          </div>
          <div>
            <label>结束日期</label>
            <input id="end_date" type="date" name="end_date" value="{end_date.isoformat()}">
          </div>
          <div>
            <label>检测项目</label>
            {sheet_selector_html}
            <p class="field-note">点击下拉后勾选多项；不勾选表示全部。</p>
          </div>
          <div style="align-self:end">
            <button type="submit">统计</button>
          </div>
        </div>
      </form>
    </div>
    <div class="kpi-grid">
      <div class="kpi"><span>检测组数</span><strong>{total_rows}</strong></div>
      <div class="kpi"><span>施工数量合计</span><strong>{format_stat_number(total_construction)}</strong></div>
      <div class="kpi"><span>检测数量合计</span><strong>{format_stat_number(total_testing)}</strong></div>
      <div class="kpi"><span>委托单位</span><strong>{source_count}</strong></div>
      <div class="kpi"><span>工程部位</span><strong>{unit_count}</strong></div>
    </div>
    <div class="panel">
      <h3>统计口径</h3>
      <p class="muted">本页按检测项目组织统计结果；季报、年报正文综合表集中在“单位工程统计”中生成。</p>
    </div>
    <div class="panel">
      <h3>统计结果导出</h3>
      <form method="post" action="/statistics_export">
        <input type="hidden" name="report_type" value="{html.escape(report_type)}">
        <input type="hidden" name="start_date" value="{start_date.isoformat()}">
        <input type="hidden" name="end_date" value="{end_date.isoformat()}">
        {export_source_inputs}
        {export_sheet_inputs}
        <div class="export-list">{export_options}</div>
        <p>
          <button type="submit">导出选中统计表为 Word</button>
          <button type="submit" formaction="/statistics_export_excel">导出选中统计表为 Excel</button>
        </p>
      </form>
    </div>
    {item_sections_html}
    {section_totals_html}
    {statistics_page_script()}
    """
    return page_layout("检测数据统计", body, user)




def advanced_query_page(user, params):
    query, records = advanced_query_records(params)
    display_records = records[:query["limit"]]
    section_selector_html = checkbox_dropdown("section_name", [(option, option) for option in ledger_distinct_options("section_name")], query["section_filters"], "全部标段")
    source_selector_html = checkbox_dropdown("source_type", [(option, option) for option in ledger_distinct_options("source_type")], query["source_filters"], "全部委托单位")
    sheet_selector_html = checkbox_dropdown("sheet_name", statistics_sheet_options(), query["sheet_filters"], "全部检测类型")
    limit_options = "".join(f"<option value='{value}' {'selected' if value == query['limit'] else ''}>{value}</option>" for value in (100, 500, 1000, 2000, 5000))
    export_url = advanced_query_export_url(params)
    body = f"""
    <div class="panel">
      <h2>高级查询</h2>
      <form method="get" action="/advanced_query">
        <div class="grid">
          <div><label>标段</label>{section_selector_html}</div>
          <div><label>检测类型/工作表</label>{sheet_selector_html}</div>
          <div><label>委托单位</label>{source_selector_html}</div>
          <div><label>原文件名</label><input name="filename" value="{html.escape(query['filename'])}" placeholder="例如 TZ-SWQ2"></div>
          <div><label>开始日期</label><input type="date" name="start_date" value="{html.escape(query['start_text'])}"></div>
          <div><label>结束日期</label><input type="date" name="end_date" value="{html.escape(query['end_text'])}"></div>
          <div><label>报告编号</label><input name="report_no" value="{html.escape(query['report_no'])}" placeholder="支持模糊查询"></div>
          <div><label>委托编号</label><input name="entrust_no" value="{html.escape(query['entrust_no'])}" placeholder="支持模糊查询"></div>
          <div><label>单位工程</label><input name="unit_name" value="{html.escape(query['unit_name'])}" placeholder="支持模糊查询"></div>
          <div><label>工程部位/分项</label><input name="location_name" value="{html.escape(query['location_name'])}" placeholder="支持模糊查询"></div>
          <div><label>检测结果</label><input name="result_text" value="{html.escape(query['result_text'])}" placeholder="例如 合格"></div>
          <div><label>页面显示条数</label><select name="limit">{limit_options}</select></div>
          <div style="align-self:end"><button type="submit">查询</button></div>
        </div>
      </form>
    </div>
    <div class="kpi-grid">
      <div class="kpi"><span>符合条件记录</span><strong>{len(records)}</strong></div>
      <div class="kpi"><span>页面显示</span><strong>{len(display_records)}</strong></div>
      <div class="kpi"><span>涉及文件</span><strong>{len({item.get('original_filename') for item in records})}</strong></div>
      <div class="kpi"><span>涉及标段</span><strong>{len({item.get('section_name') for item in records})}</strong></div>
      <div class="kpi"><span>检测类型</span><strong>{len({item.get('sheet_name') for item in records})}</strong></div>
    </div>
    <div class="panel">
      <p><a class="button" href="{html.escape(export_url)}">导出当前查询结果为 Excel</a></p>
      <table class="ledger-table">
        <thead><tr><th>标段</th><th>委托单位</th><th>原文件</th><th>工作表</th><th>行号</th><th>检测日期</th><th>报告编号</th><th>委托编号</th><th>单位工程</th><th>工程部位</th><th>检测结果</th><th>施工数量</th><th>检测数量</th></tr></thead>
        <tbody>{render_advanced_query_table(display_records)}</tbody>
      </table>
    </div>
    {statistics_page_script()}
    """
    return page_layout("高级查询", body, user)


def quality_checks_page(user):
    records, issues = build_quality_issues()
    high_count = sum(1 for issue in issues if issue.get("level") == "高")
    middle_count = sum(1 for issue in issues if issue.get("level") == "中")
    issue_type_count = len({issue.get("issue_type") for issue in issues})
    body = f"""
    <div class="panel">
      <h2>异常数据检查</h2>
      <p class="muted">系统会检查同编号数据冲突、编号缺失、日期缺失、数量异常、检测结果缺失，以及文件名与标段/委托类型不一致等问题。同一编号在多个台账中标段、检测类型、日期和数量一致时不视为异常。</p>
    </div>
    <div class="kpi-grid">
      <div class="kpi"><span>检查记录</span><strong>{len(records)}</strong></div>
      <div class="kpi"><span>异常总数</span><strong>{len(issues)}</strong></div>
      <div class="kpi"><span>高风险</span><strong>{high_count}</strong></div>
      <div class="kpi"><span>中风险</span><strong>{middle_count}</strong></div>
      <div class="kpi"><span>异常类型</span><strong>{issue_type_count}</strong></div>
    </div>
    <div class="panel">{render_quality_issue_table(issues)}</div>
    """
    return page_layout("异常数据检查", body, user)

def unit_statistics_page(user, params):
    stat_params = parse_statistics_params(params)
    report_type = stat_params["report_type"]
    start_date = stat_params["start_date"]
    end_date = stat_params["end_date"]
    source_filters = stat_params["source_filters"]
    sheet_filters = stat_params["sheet_filters"]
    unit_filters = stat_params["unit_filters"]

    records = collect_stat_records(start_date, end_date, source_filters, sheet_filters)
    filtered_records = filter_records_by_units(records, unit_filters)
    total_rows = sum(record.get("row_count", 1) for record in filtered_records)
    total_construction = sum(record.get("construction_qty", 0) for record in filtered_records)
    total_testing = sum(record.get("testing_qty", 0) for record in filtered_records)
    section_count = len({record.get("section_name") for record in filtered_records})
    unit_count = len({record.get("unit_name") for record in filtered_records if record.get("unit_name")})

    report_options = [
        ("week", "周报"),
        ("month", "月报"),
        ("quarter", "季报"),
        ("year", "年报"),
        ("custom", "自定义时间段"),
    ]
    report_option_html = "".join(
        f"<option value='{value}' {'selected' if value == report_type else ''}>{label}</option>"
        for value, label in report_options
    )
    source_selector_html = checkbox_dropdown(
        "source_type",
        [(option, option) for option in SOURCE_TYPE_OPTIONS],
        source_filters,
        "全部委托单位",
    )
    sheet_selector_html = checkbox_dropdown(
        "sheet_name",
        statistics_sheet_options(),
        sheet_filters,
        "全部检测项目",
    )
    unit_selector_html = checkbox_dropdown(
        "unit_name",
        statistics_unit_options(source_filters, sheet_filters),
        unit_filters,
        "全部单位工程",
    )
    unit_project_summary_html = render_unit_project_summary(records, unit_filters)
    filtered_grouped = aggregate_stat_records(filtered_records)
    period_report_tables = build_period_report_tables(
        filtered_grouped,
        stat_params,
        source_filters,
        sheet_filters,
        unit_filters,
    )
    period_report_html = render_period_report_sections(period_report_tables)
    period_export_html = ""
    if period_report_tables:
        period_export_html = f"""
        <div class="panel">
          <h3>季报/年报综合表导出</h3>
          <p class="muted">导出内容跟随当前筛选条件；如已选择单位工程，只导出该单位工程相关综合表。</p>
          <form method="post" action="/unit_statistics_export">
            <input type="hidden" name="report_type" value="{html.escape(report_type)}">
            <input type="hidden" name="start_date" value="{start_date.isoformat()}">
            <input type="hidden" name="end_date" value="{end_date.isoformat()}">
            {form_hidden_inputs("source_type", source_filters)}
            {form_hidden_inputs("sheet_name", sheet_filters)}
            {form_hidden_inputs("unit_name", unit_filters)}
            <p>
              <button type="submit">导出选中综合表为 Word</button>
              <button type="submit" formaction="/unit_statistics_export_excel">导出选中综合表为 Excel</button>
            </p>
            <details class="export-project">
              <summary>选择要导出的综合表</summary>
              <div class="export-list">{render_export_options(period_report_tables)}</div>
            </details>
          </form>
        </div>
        """

    body = f"""
    <div class="panel">
      <h2>单位工程统计</h2>
      <form method="get" action="/unit_statistics">
        <div class="grid">
          <div>
            <label>报表类型</label>
            <select id="report_type" name="report_type">{report_option_html}</select>
          </div>
          <div>
            <label>委托单位</label>
            {source_selector_html}
            <p class="field-note">点击下拉后勾选多项；不勾选表示全部。</p>
          </div>
          <div>
            <label>开始日期</label>
            <input id="start_date" type="date" name="start_date" value="{start_date.isoformat()}">
          </div>
          <div>
            <label>结束日期</label>
            <input id="end_date" type="date" name="end_date" value="{end_date.isoformat()}">
          </div>
          <div>
            <label>检测项目</label>
            {sheet_selector_html}
            <p class="field-note">点击下拉后勾选多项；不勾选表示全部。</p>
          </div>
          <div>
            <label>单位工程</label>
            {unit_selector_html}
            <p class="field-note">点击下拉后勾选多项；不勾选表示全部。</p>
          </div>
          <div style="align-self:end">
            <button type="submit">统计</button>
          </div>
        </div>
      </form>
    </div>
    <div class="kpi-grid">
      <div class="kpi"><span>检测组数</span><strong>{total_rows}</strong></div>
      <div class="kpi"><span>施工数量合计</span><strong>{format_stat_number(total_construction)}</strong></div>
      <div class="kpi"><span>检测数量合计</span><strong>{format_stat_number(total_testing)}</strong></div>
      <div class="kpi"><span>标段数量</span><strong>{section_count}</strong></div>
      <div class="kpi"><span>单位工程</span><strong>{unit_count}</strong></div>
    </div>
    {period_export_html}
    {period_report_html}
    {unit_project_summary_html}
    {statistics_page_script()}
    """
    return page_layout("单位工程统计", body, user)


def col_letter(index):
    letters = ""
    while index:
        index, rem = divmod(index - 1, 26)
        letters = chr(65 + rem) + letters
    return letters or "A"


def render_sheet_grid(sheet, cells, merges, row_limit, col_limit, editable=False):
    filter_row = detect_header_row(cells, row_limit)
    data_start_row = filter_row + 1
    merge_starts = {}
    merge_covered = set()
    for merge in merges:
        min_row = int(merge["min_row"] or 0)
        min_col = int(merge["min_col"] or 0)
        max_row = int(merge["max_row"] or 0)
        max_col = int(merge["max_col"] or 0)
        if min_row <= 0 or min_col <= 0 or max_row < min_row or max_col < min_col:
            continue
        merge_starts[(min_row, min_col)] = (max_row - min_row + 1, max_col - min_col + 1)
        for rr in range(min_row, max_row + 1):
            for cc in range(min_col, max_col + 1):
                if (rr, cc) != (min_row, min_col):
                    merge_covered.add((rr, cc))
    ratio_cols = {
        col
        for col in range(1, col_limit + 1)
        if is_ratio_header(cells.get((filter_row, col), ""))
    }
    header_cells = "<th class='row-head'></th>" + "".join(f"<th>{col_letter(c)}</th>" for c in range(1, col_limit + 1))
    body_rows = []
    for r in range(1, row_limit + 1):
        row_cells = [f"<td class='row-num'>{r}</td>"]
        for c in range(1, col_limit + 1):
            if (r, c) in merge_covered:
                continue
            display_value = display_cell_text(cells.get((r, c), ""))
            if r >= data_start_row and c in ratio_cols:
                display_value = display_ratio_text(display_value)
            value = html.escape(display_value)
            edit_attrs = " contenteditable='plaintext-only' tabindex='0'" if editable else ""
            filter_button = (
                f"<button type='button' class='filter-btn' data-filter-col='{c}' title='筛选' contenteditable='false'>筛选</button>"
                if r == filter_row
                else ""
            )
            span_attrs = ""
            if (r, c) in merge_starts:
                rowspan, colspan = merge_starts[(r, c)]
                span_attrs = f" rowspan='{rowspan}' colspan='{colspan}'"
            row_cells.append(f"<td{edit_attrs}{span_attrs} data-row='{r}' data-col='{c}' data-filter-row='{filter_row}' data-data-start-row='{data_start_row}'><span class='cell-text'>{value}</span>{filter_button}</td>")
        body_rows.append("<tr>" + "".join(row_cells) + "</tr>")
    return f"<div class='sheet-wrap'><table class='sheet-grid'><thead><tr>{header_cells}</tr></thead><tbody>{''.join(body_rows)}</tbody></table></div>"


def clean_sheet_display_name(name):
    text = str(name or "").strip()
    previous = None
    while text and text != previous:
        previous = text
        text = re.sub(r"^\s*[\(\（\[]\s*\d+\s*[\)\）\]]\s*", "", text).strip()
        text = re.sub(r"^\s*\d+\s*[\.．、\-_\s]+\s*", "", text).strip()
    return text or str(name or "")


def preview_page(user, file_id=None, sheet_id=None):
    conn = connect()
    try:
        files = conn.execute(
            """
            SELECT f.id, f.original_filename, f.section_name, f.source_type, f.discipline,
                   v.version_no, v.workbook_id
            FROM ledger_file f
            JOIN ledger_file_version v ON v.id = f.current_version_id
            ORDER BY f.id ASC
            """
        ).fetchall()
        if file_id is None and files:
            file_id = files[0]["id"]

        current_file = None
        sheets = []
        if file_id is not None:
            current_file = conn.execute(
                """
                SELECT f.id, f.original_filename, f.section_name, f.source_type, f.discipline,
                       v.version_no, v.workbook_id
                FROM ledger_file f
                JOIN ledger_file_version v ON v.id = f.current_version_id
                WHERE f.id = ?
                """,
                (file_id,),
            ).fetchone()
            if current_file:
                sheets = conn.execute(
                    "SELECT * FROM template_sheet WHERE workbook_id = ? ORDER BY sheet_index",
                    (current_file["workbook_id"],),
                ).fetchall()

        if sheet_id is None and sheets:
            sheet_id = sheets[0]["id"]

        current_sheet = None
        cells = {}
        merges = []
        row_limit = 0
        col_limit = 0
        if sheet_id is not None:
            current_sheet = conn.execute("SELECT * FROM template_sheet WHERE id = ?", (sheet_id,)).fetchone()
            if current_sheet:
                row_limit = current_sheet["max_row"] or 0
                col_limit = current_sheet["max_column"] or 0
                for cell in conn.execute(
                    """
                    SELECT row_index, col_index, raw_value
                    FROM template_cell
                    WHERE sheet_id = ? AND row_index <= ? AND col_index <= ?
                    """,
                    (sheet_id, row_limit, col_limit),
                ):
                    cells[(cell["row_index"], cell["col_index"])] = cell["raw_value"] or ""
                merges = conn.execute(
                    """
                    SELECT range_ref, min_row, min_col, max_row, max_col
                    FROM template_merge
                    WHERE sheet_id = ?
                    ORDER BY id
                    """,
                    (sheet_id,),
                ).fetchall()
    finally:
        conn.close()

    file_options = "".join(
        f"<option value='{f['id']}' {'selected' if f['id'] == file_id else ''}>{html.escape(f['section_name'] or '')} / {html.escape(f['source_type'] or '')}</option>"
        for f in files
    )
    if not files:
        body = "<div class='panel'><h2>台账信息查询</h2><p>暂无已上传台账，请先上传 Excel 文件。</p></div>"
        return page_layout("台账信息查询", body, user)

    sheet_options = "".join(
        f"<option value='{s['id']}' {'selected' if s['id'] == sheet_id else ''}>{html.escape(clean_sheet_display_name(s['sheet_name']))}</option>"
        for s in sheets
    )
    can_edit = is_admin(user)
    grid_html = ""
    sheet_meta = ""
    merge_text = "无"
    if current_sheet:
        grid_html = render_sheet_grid(current_sheet, cells, merges, row_limit, col_limit, editable=can_edit)
        sheet_meta = f"当前工作表：{html.escape(current_sheet['sheet_name'])}；原始范围：{current_sheet['max_row']} 行 x {current_sheet['max_column']} 列；当前显示完整工作表内容。"
        merge_text = "、".join(html.escape(m["range_ref"]) for m in merges) or "无"
    zoom_toolbar = ""
    zoom_script = ""
    if current_sheet:
        zoom_toolbar = """
        <div class="zoom-bar">
          <input id="sheetFilterInput" type="search" placeholder="筛选当前表格内容">
          <button type="button" id="zoomOutBtn">-</button>
          <select id="zoomSelect" aria-label="缩放比例">
            <option value="0.5">50%</option>
            <option value="0.75">75%</option>
            <option value="0.9">90%</option>
            <option value="1" selected>100%</option>
            <option value="1.1">110%</option>
            <option value="1.25">125%</option>
            <option value="1.5">150%</option>
            <option value="2">200%</option>
          </select>
          <button type="button" id="zoomInBtn">+</button>
        </div>
        """
        zoom_script = """
        <script>
        (() => {
          const grid = document.querySelector(".sheet-grid");
          const select = document.getElementById("zoomSelect");
          const zoomOut = document.getElementById("zoomOutBtn");
          const zoomIn = document.getElementById("zoomInBtn");
          const filterInput = document.getElementById("sheetFilterInput");
          if (!grid || !select) return;
          const columnFilters = new Map();
          let activePopover = null;
          const levels = Array.from(select.options).map(option => Number(option.value));
          function cellText(row, col) {
            const cell = row.querySelector(`[data-col="${col}"]`);
            const textNode = cell ? cell.querySelector(".cell-text") : null;
            return (textNode ? textNode.innerText : "").trim();
          }
          function applyFilters() {
            const keyword = filterInput ? filterInput.value.trim().toLowerCase() : "";
            grid.querySelectorAll("tbody tr").forEach(row => {
              const rowIndex = Number(row.querySelector(".row-num")?.innerText || 0);
              const dataStart = Number(row.querySelector("[data-data-start-row]")?.dataset.dataStartRow || 5);
              if (rowIndex && rowIndex < dataStart) {
                row.style.display = "";
                return;
              }
              let visible = !keyword || Array.from(row.querySelectorAll(".cell-text")).some(cell => cell.innerText.toLowerCase().includes(keyword));
              if (visible) {
                for (const [col, allowed] of columnFilters.entries()) {
                  if (allowed.size && !allowed.has(cellText(row, col))) {
                    visible = false;
                    break;
                  }
                }
              }
              row.style.display = visible ? "" : "none";
            });
          }
          function closePopover() {
            if (activePopover) activePopover.remove();
            activePopover = null;
          }
          function uniqueColumnValues(col) {
            const values = new Set();
            grid.querySelectorAll("tbody tr").forEach(row => {
              const rowIndex = Number(row.querySelector(".row-num")?.innerText || 0);
              const dataStart = Number(row.querySelector("[data-data-start-row]")?.dataset.dataStartRow || 5);
              if (!rowIndex || rowIndex < dataStart) return;
              values.add(cellText(row, col));
            });
            return Array.from(values).sort((a, b) => a.localeCompare(b, "zh-CN"));
          }
          function dateParts(value) {
            const match = String(value || "").trim().match(/^(\\d{4})-(\\d{1,2})-(\\d{1,2})$/);
            if (!match) return null;
            return { year: match[1], month: match[2].padStart(2, "0"), day: match[3].padStart(2, "0") };
          }
          function fillDateSelect(select, values, label) {
            select.innerHTML = `<option value="">${label}</option>` + values.map(value => `<option value="${value}">${value}</option>`).join("");
          }
          function openFilter(col, anchor) {
            closePopover();
            const values = uniqueColumnValues(col);
            const selected = columnFilters.get(col) || new Set(values);
            const dates = values.map(dateParts).filter(Boolean);
            const hasDateFilter = dates.length > 0;
            const pop = document.createElement("div");
            pop.className = "filter-popover";
            const rect = anchor.getBoundingClientRect();
            pop.style.left = Math.min(rect.left, window.innerWidth - 300) + "px";
            pop.style.top = Math.min(rect.bottom + 6, window.innerHeight - 390) + "px";
            pop.innerHTML = `
              <input type="search" class="filter-search" placeholder="搜索当前列">
              <div class="date-filter" style="${hasDateFilter ? "" : "display:none"}">
                <select class="filter-year"></select>
                <select class="filter-month"></select>
                <select class="filter-day"></select>
              </div>
              <div class="filter-actions">
                <button type="button" data-action="all">全选</button>
                <button type="button" data-action="none">清空</button>
                <button type="button" data-action="apply">确定</button>
                <button type="button" data-action="cancel">取消</button>
              </div>
              <div class="filter-values"></div>
            `;
            const list = pop.querySelector(".filter-values");
            const yearSelect = pop.querySelector(".filter-year");
            const monthSelect = pop.querySelector(".filter-month");
            const daySelect = pop.querySelector(".filter-day");
            function renderList(keyword = "") {
              list.innerHTML = "";
              values.filter(value => !keyword || value.toLowerCase().includes(keyword.toLowerCase())).forEach(value => {
                const label = document.createElement("label");
                const checkbox = document.createElement("input");
                checkbox.type = "checkbox";
                checkbox.value = value;
                checkbox.checked = selected.has(value);
                label.appendChild(checkbox);
                label.appendChild(document.createTextNode(value || "(空白)"));
                list.appendChild(label);
              });
            }
            renderList();
            function applyDatePick() {
              const year = yearSelect.value;
              const month = monthSelect.value;
              const day = daySelect.value;
              if (!year && !month && !day) return;
              selected.clear();
              values.forEach(value => {
                const parts = dateParts(value);
                if (!parts) return;
                if (year && parts.year !== year) return;
                if (month && parts.month !== month) return;
                if (day && parts.day !== day) return;
                selected.add(value);
              });
              renderList(pop.querySelector(".filter-search").value);
            }
            if (hasDateFilter) {
              fillDateSelect(yearSelect, Array.from(new Set(dates.map(item => item.year))).sort(), "年");
              fillDateSelect(monthSelect, Array.from(new Set(dates.map(item => item.month))).sort(), "月");
              fillDateSelect(daySelect, Array.from(new Set(dates.map(item => item.day))).sort(), "日");
              yearSelect.addEventListener("change", applyDatePick);
              monthSelect.addEventListener("change", applyDatePick);
              daySelect.addEventListener("change", applyDatePick);
            }
            pop.querySelector(".filter-search").addEventListener("input", event => renderList(event.target.value));
            pop.addEventListener("click", event => event.stopPropagation());
            pop.querySelector("[data-action='all']").addEventListener("click", () => {
              values.forEach(value => selected.add(value));
              renderList(pop.querySelector(".filter-search").value);
            });
            pop.querySelector("[data-action='none']").addEventListener("click", () => {
              selected.clear();
              renderList(pop.querySelector(".filter-search").value);
            });
            pop.querySelector("[data-action='cancel']").addEventListener("click", closePopover);
            pop.querySelector("[data-action='apply']").addEventListener("click", () => {
              selected.clear();
              list.querySelectorAll("input[type='checkbox']").forEach(box => {
                if (box.checked) selected.add(box.value);
              });
              if (selected.size === values.length) columnFilters.delete(col);
              else columnFilters.set(col, selected);
              anchor.style.background = columnFilters.has(col) ? "#2a8c7c" : "";
              applyFilters();
              closePopover();
            });
            document.body.appendChild(pop);
            activePopover = pop;
          }
          function applyZoom(value) {
            const zoom = Math.min(2, Math.max(0.5, Number(value) || 1));
            grid.style.setProperty("--zoom", zoom);
            let nearest = levels.reduce((best, item) => Math.abs(item - zoom) < Math.abs(best - zoom) ? item : best, levels[0]);
            select.value = String(nearest);
            localStorage.setItem("ledger_preview_zoom", String(nearest));
          }
          function step(delta) {
            const current = levels.indexOf(Number(select.value));
            const next = Math.min(levels.length - 1, Math.max(0, current + delta));
            applyZoom(levels[next]);
          }
          select.addEventListener("change", () => applyZoom(select.value));
          zoomOut.addEventListener("click", () => step(-1));
          zoomIn.addEventListener("click", () => step(1));
          if (filterInput) {
            filterInput.addEventListener("input", applyFilters);
          }
          grid.querySelectorAll(".filter-btn").forEach(button => {
            button.addEventListener("click", event => {
              event.stopPropagation();
              openFilter(Number(button.dataset.filterCol), button);
            });
          });
          document.addEventListener("click", closePopover);
          applyZoom(localStorage.getItem("ledger_preview_zoom") || select.value);
        })();
        </script>
        """
    edit_toolbar = ""
    edit_script = ""
    if can_edit and current_sheet:
        edit_toolbar = """
        <div class="edit-bar">
          <button type="button" id="saveSheetBtn">保存修改</button>
          <input id="deleteRowInput" type="number" min="1" placeholder="行号" style="width:96px">
          <button type="button" id="deleteRowBtn" style="background:#9f2d2d">删除行</button>
          <input id="deleteColInput" type="text" placeholder="列号/字母" style="width:96px">
          <button type="button" id="deleteColBtn" style="background:#9f2d2d">删除列</button>
          <span id="saveStatus" class="muted">管理员编辑模式：可直接修改单元格，支持粘贴 Excel 区域。</span>
        </div>
        """
        edit_script = f"""
        <script>
        (() => {{
          const sheetId = {int(sheet_id or 0)};
          const changed = new Map();
          const status = document.getElementById("saveStatus");
          const saveBtn = document.getElementById("saveSheetBtn");
          const grid = document.querySelector(".sheet-grid");
          const deleteRowBtn = document.getElementById("deleteRowBtn");
          const deleteColBtn = document.getElementById("deleteColBtn");
          function keyOf(cell) {{ return cell.dataset.row + ":" + cell.dataset.col; }}
          function plainText(value) {{
            return String(value || "").replace(/\\u00a0/g, " ").replace(/\\r\\n/g, "\\n").replace(/\\r/g, "\\n");
          }}
          function colTextToIndex(text) {{
            const value = String(text || "").trim().toUpperCase();
            if (/^\\d+$/.test(value)) return Number(value);
            let total = 0;
            for (const ch of value) {{
              if (ch < "A" || ch > "Z") return 0;
              total = total * 26 + (ch.charCodeAt(0) - 64);
            }}
            return total;
          }}
          function markChanged(cell) {{
            if (!cell || !cell.dataset.row || !cell.dataset.col) return;
            const textNode = cell.querySelector(".cell-text") || cell;
            const text = plainText(textNode.innerText).replace(/\\n$/, "");
            if (textNode.innerText !== text) textNode.innerText = text;
            changed.set(keyOf(cell), {{
              row: Number(cell.dataset.row),
              col: Number(cell.dataset.col),
              value: text
            }});
            status.textContent = "有未保存修改：" + changed.size + " 个单元格";
          }}
          function focusCell(row, col) {{
            const next = grid.querySelector(`[data-row="${{row}}"][data-col="${{col}}"]`);
            if (next) {{
              next.focus();
              const range = document.createRange();
              range.selectNodeContents(next);
              range.collapse(false);
              const sel = window.getSelection();
              sel.removeAllRanges();
              sel.addRange(range);
            }}
          }}
          grid.addEventListener("input", event => markChanged(event.target));
          grid.addEventListener("click", event => {{
            if (event.target.classList.contains("filter-btn")) return;
          }});
          grid.addEventListener("drop", event => event.preventDefault());
          grid.addEventListener("keydown", event => {{
            const cell = event.target.closest("td[data-row]");
            if (!cell) return;
            const row = Number(cell.dataset.row);
            const col = Number(cell.dataset.col);
            if (event.key === "Tab") {{
              event.preventDefault();
              focusCell(row, col + (event.shiftKey ? -1 : 1));
            }} else if (event.key === "Enter") {{
              event.preventDefault();
              focusCell(row + (event.shiftKey ? -1 : 1), col);
            }} else if (event.altKey && event.key === "ArrowUp") {{
              event.preventDefault();
              focusCell(row - 1, col);
            }} else if (event.altKey && event.key === "ArrowDown") {{
              event.preventDefault();
              focusCell(row + 1, col);
            }} else if (event.altKey && event.key === "ArrowLeft") {{
              event.preventDefault();
              focusCell(row, col - 1);
            }} else if (event.altKey && event.key === "ArrowRight") {{
              event.preventDefault();
              focusCell(row, col + 1);
            }}
          }});
          grid.addEventListener("paste", event => {{
            const start = event.target.closest("td[data-row]");
            if (!start) return;
            event.preventDefault();
            const text = plainText(event.clipboardData.getData("text/plain"));
            const startRow = Number(start.dataset.row);
            const startCol = Number(start.dataset.col);
            if (!text.includes("\\t") && !text.includes("\\n")) {{
              const textNode = start.querySelector(".cell-text") || start;
              textNode.innerText = text;
              markChanged(start);
              return;
            }}
            text.replace(/\\r$/, "").split(/\\r?\\n/).forEach((line, rOffset) => {{
              if (line === "" && rOffset === text.split(/\\r?\\n/).length - 1) return;
              line.split("\\t").forEach((value, cOffset) => {{
                const cell = grid.querySelector(`[data-row="${{startRow + rOffset}}"][data-col="${{startCol + cOffset}}"]`);
                if (cell) {{
                  const textNode = cell.querySelector(".cell-text") || cell;
                  textNode.innerText = value;
                  markChanged(cell);
                }}
              }});
            }});
          }});
          saveBtn.addEventListener("click", async () => {{
            if (!changed.size) {{
              status.textContent = "没有需要保存的修改";
              return;
            }}
            saveBtn.disabled = true;
            status.textContent = "正在保存...";
            try {{
              const response = await fetch("/save_sheet_cells", {{
                method: "POST",
                headers: {{ "Content-Type": "application/json" }},
                body: JSON.stringify({{ sheet_id: sheetId, changes: Array.from(changed.values()) }})
              }});
              const result = await response.json();
              if (!response.ok || !result.ok) throw new Error(result.error || "保存失败");
              changed.clear();
              status.textContent = "已保存 " + result.saved + " 个单元格";
            }} catch (error) {{
              status.textContent = error.message;
            }} finally {{
              saveBtn.disabled = false;
            }}
          }});
          async function deleteAxis(axis, index) {{
            if (!index || index < 1) {{
              status.textContent = "请输入有效的" + (axis === "row" ? "行号" : "列号");
              return;
            }}
            if (!confirm("确认删除该" + (axis === "row" ? "行" : "列") + "？此操作会移动后续数据。")) return;
            status.textContent = "正在删除...";
            try {{
              const response = await fetch("/delete_sheet_axis", {{
                method: "POST",
                headers: {{ "Content-Type": "application/json" }},
                body: JSON.stringify({{ sheet_id: sheetId, axis, index }})
              }});
              const result = await response.json();
              if (!response.ok || !result.ok) throw new Error(result.error || "删除失败");
              location.reload();
            }} catch (error) {{
              status.textContent = error.message;
            }}
          }}
          deleteRowBtn.addEventListener("click", () => {{
            deleteAxis("row", Number(document.getElementById("deleteRowInput").value));
          }});
          deleteColBtn.addEventListener("click", () => {{
            deleteAxis("col", colTextToIndex(document.getElementById("deleteColInput").value));
          }});
        }})();
        </script>
        """

    body = f"""
    <div class="panel">
      <h2>台账信息查询</h2>
      <form method="get" action="/preview">
        <div class="toolbar">
          <div><label>选择台账</label><select name="file_id" onchange="this.form.submit()">{file_options}</select></div>
          <div><label>选择工作表</label><select name="sheet_id" onchange="this.form.submit()">{sheet_options}</select></div>
          <div><button type="submit">查看</button></div>
        </div>
      </form>
      <p class="muted">选择文件和工作表后，下方显示该工作表的完整原始范围。公式单元格显示计算结果，公式本身仍保存在数据库中。</p>
    </div>
    <div class="panel">
      <h3>{html.escape(current_sheet['sheet_name']) if current_sheet else '无工作表'}</h3>
      {zoom_toolbar}
      {edit_toolbar}
      <p class="muted">{sheet_meta}</p>
      <p class="muted">合并区域：{merge_text}</p>
      {grid_html or '该文件暂无工作表数据。'}
    </div>
    {zoom_script}
    {edit_script}
    """
    return page_layout("台账信息查询", body, user)


def file_detail_page(file_id, user):
    conn = connect()
    try:
        file_row = conn.execute("SELECT * FROM ledger_file WHERE id = ?", (file_id,)).fetchone()
        versions = conn.execute(
            "SELECT * FROM ledger_file_version WHERE ledger_file_id = ? ORDER BY version_no DESC",
            (file_id,),
        ).fetchall()
        current = versions[0] if versions else None
        sheets = []
        if current:
            sheets = conn.execute(
                "SELECT * FROM template_sheet WHERE workbook_id = ? ORDER BY sheet_index",
                (current["workbook_id"],),
            ).fetchall()
    finally:
        conn.close()
    if not file_row:
        return page_layout("未找到", "<div class='panel'>未找到该台账。</div>", user)

    version_rows = "".join(
        f"<tr><td>v{v['version_no']:03d}</td><td>{html.escape(v['uploaded_at'])}</td><td>{v['file_size']}</td>"
        f"<td>{html.escape((v['file_hash'] or '')[:16])}...</td><td>{html.escape(v['remark'] or '')}</td>"
        f"<td><a class='button' href='/download?version_id={v['id']}'>下载原文件</a> "
        f"<form method='post' action='/reparse' style='display:inline'>"
        f"<input type='hidden' name='version_id' value='{v['id']}'>"
        f"<input type='hidden' name='file_id' value='{file_id}'>"
        f"<button type='submit'>重新解析预览数据</button></form></td></tr>"
        for v in versions
    )
    sheet_rows = "".join(
        f"<tr><td>{s['sheet_index']}</td><td><a href='/sheet?id={s['id']}'>{html.escape(s['sheet_name'])}</a></td>"
        f"<td>{s['max_row']}</td><td>{s['max_column']}</td></tr>"
        for s in sheets
    )
    body = f"""
    <div class="panel">
      <h2>{html.escape(file_row['original_filename'])}</h2>
      <p>项目：{html.escape(file_row['project_name'] or '')}　标段：{html.escape(file_row['section_name'] or '')}　委托单位：{html.escape(file_row['source_type'] or '')}　检测单位：{html.escape(file_row['discipline'] or '')}</p>
    </div>
    <div class="panel"><h3>版本记录</h3><table><thead><tr><th>版本</th><th>上传时间</th><th>大小</th><th>哈希</th><th>说明</th><th>操作</th></tr></thead><tbody>{version_rows}</tbody></table></div>
    <div class="panel"><h3>当前版本工作表</h3><table><thead><tr><th>序号</th><th>工作表</th><th>行数</th><th>列数</th></tr></thead><tbody>{sheet_rows}</tbody></table></div>
    """
    return page_layout("台账详情", body, user)


def sheet_page(sheet_id, user):
    conn = connect()
    try:
        sheet = conn.execute("SELECT * FROM template_sheet WHERE id = ?", (sheet_id,)).fetchone()
        cells = conn.execute(
            """
            SELECT cell_ref, raw_value, number_format
            FROM template_cell
            WHERE sheet_id = ?
            ORDER BY row_index, col_index
            LIMIT 300
            """,
            (sheet_id,),
        ).fetchall()
        merges = conn.execute(
            "SELECT range_ref FROM template_merge WHERE sheet_id = ? ORDER BY id LIMIT 80",
            (sheet_id,),
        ).fetchall()
    finally:
        conn.close()
    if not sheet:
        return page_layout("未找到", "<div class='panel'>未找到该工作表。</div>", user)
    cell_rows = "".join(
        f"<tr><td>{html.escape(c['cell_ref'])}</td><td>{html.escape(display_cell_text(c['raw_value'] or ''))}</td><td>{html.escape(c['number_format'] or '')}</td></tr>"
        for c in cells
    )
    merge_text = "、".join(html.escape(m["range_ref"]) for m in merges) or "无"
    body = f"""
    <div class="panel"><h2>{html.escape(sheet['sheet_name'])}</h2><p>行数：{sheet['max_row']}　列数：{sheet['max_column']}</p><p>合并区域：{merge_text}</p></div>
    <div class="panel"><h3>单元格预览 前300个有值单元格</h3><table><thead><tr><th>单元格</th><th>原始值</th><th>格式</th></tr></thead><tbody>{cell_rows}</tbody></table></div>
    """
    return page_layout("工作表预览", body, user)


def main():
    init_db()
    host = os.environ.get("LEDGER_WEB_HOST", "0.0.0.0")
    port = int(os.environ.get("LEDGER_WEB_PORT", "8765"))
    backup_port = int(os.environ.get("LEDGER_WEB_BACKUP_PORT", "8766"))
    server = ThreadingHTTPServer((host, port), AppHandler)
    backup_server = ThreadingHTTPServer((host, backup_port), AppHandler)
    threading.Thread(target=backup_server.serve_forever, daemon=True).start()
    log_error(f"台账系统已启动：http://{host}:{port}")
    log_error(f"台账系统备用端口已启动：http://{host}:{backup_port}")
    log_error(f"数据库：{DB_PATH}")
    log_error(f"文件归档：{STORAGE_DIR}")
    server.serve_forever()


if __name__ == "__main__":
    try:
        main()
    except Exception:
        log_error("服务启动或运行异常：\n" + traceback.format_exc())
        raise
